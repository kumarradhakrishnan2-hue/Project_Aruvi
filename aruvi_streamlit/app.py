from pathlib import Path
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

import base64
import csv
import io
import json
import re
from datetime import datetime, date
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import math
import threading
import queue

import uuid
import streamlit as st
import streamlit.components.v1 as components
from streamlit.runtime.scriptrunner import add_script_run_ctx
import anthropic
import os
# ── Prompt caching toggle ─────────────────────────────────────────────────────
# Set USE_PROMPT_CACHE = True  → cache_control blocks active (1h TTL)
#                                cache_write costs 2× input rate per token
#                                cache_read  costs 0.1× input rate per token
#                                benefit: repeated chapters in same session
#                                hit the cache and save ~90% on static tokens
# Set USE_PROMPT_CACHE = False → no cache_control sent; all tokens billed at
#                                standard input rate (1× — no surcharge)
#                                use during development / single-chapter runs
USE_PROMPT_CACHE = False

def _cache_ctrl() -> dict:
    """Return cache_control block if caching is enabled, else empty dict."""
    return {"cache_control": {"type": "ephemeral", "ttl": "1h"}} if USE_PROMPT_CACHE else {}

# ── Ask Aruvi backend toggle ──────────────────────────────────────────────────
# Set USE_MANAGED_AGENT = True  → new managed-agent path (ask_aruvi_agent.py)
# Set USE_MANAGED_AGENT = False → original Haiku path  (ask_aruvi_qa.py)
# The old module is NOT deleted — flip the flag to revert instantly.
USE_MANAGED_AGENT = False

if USE_MANAGED_AGENT:
    from ask_aruvi_agent import ask as aruvi_ask          # ← managed agent
else:
    from ask_aruvi_qa import ask as aruvi_ask             # ← original Haiku (immobilised)

from ask_aruvi_feedback import write_thumbs_feedback, write_general_feedback
from prompt_redaction import redact_summary_for_assessment, redact_coverage_handoff

# ── Project root (needed by helper functions below) ───────────────────────────

PROJECT_ROOT = Path(__file__).parent.parent
MISC_DIR     = PROJECT_ROOT / "miscellaneous"

# ── Stage derivation ──────────────────────────────────────────────────────────

def get_stage(grade: str) -> str:
    preparatory = {"Grade III", "Grade IV", "Grade V"}
    middle       = {"Grade VI", "Grade VII", "Grade VIII"}
    if grade in preparatory: return "preparatory"
    if grade in middle:      return "middle"
    return "secondary"

def grade_to_folder(grade: str) -> str:
    """Return the folder name for a grade — matches the roman-numeral dirs in mirror/."""
    _mapping = {
        "Grade I":    "i",    "Grade II":   "ii",   "Grade III": "iii",
        "Grade IV":   "iv",   "Grade V":    "v",    "Grade VI":  "vi",
        "Grade VII":  "vii",  "Grade VIII": "viii",
        "Grade IX":   "ix",
    }
    return _mapping.get(grade, grade.lower().replace("grade ", ""))

def subject_to_folder(subject: str) -> str:
    mapping = {
        "Social Science":       "social_sciences",
        "Mathematics":          "mathematics",
        "Science":              "science",
        "English":              "english",
        "The World Around Us":  "the_world_around_us",
    }
    return mapping.get(subject, subject.lower().replace(" ", "_"))

# Subjects whose chapter summaries are JSON (structured for downstream LP/A
# constitutions). All others are plain .txt.
_JSON_SUMMARY_SUBJECTS = {"mathematics", "english", "the_world_around_us"}

# ── Path resolver ─────────────────────────────────────────────────────────────

def resolve_paths(grade: str, subject: str, chapter_number: int) -> dict:
    stage  = get_stage(grade)
    grade_f = grade_to_folder(grade)
    subj_f  = subject_to_folder(subject)
    mirror  = PROJECT_ROOT / "mirror"
    nn      = f"{chapter_number:02d}"
    # Prefer stage-routed LP and assessment constitutions
    # (`{subject}/{stage}/...txt`); fall back to the flat path for subjects
    # that haven't been split by stage yet.
    _lp_staged = mirror / f"constitutions/lesson_plan/{subj_f}/{stage}/lesson_plan_constitution.txt"
    _lp_flat   = mirror / f"constitutions/lesson_plan/{subj_f}/lesson_plan_constitution.txt"
    _ac_staged = mirror / f"constitutions/assessment/{subj_f}/{stage}/assessment_constitution.txt"
    _ac_flat   = mirror / f"constitutions/assessment/{subj_f}/assessment_constitution.txt"
    return {
        "lp_constitution":  _lp_staged if _lp_staged.exists() else _lp_flat,
        "assessment_const": _ac_staged if _ac_staged.exists() else _ac_flat,
        "pedagogy":         mirror / f"framework/{subj_f}/{stage}/pedagogy_{stage}_{subj_f}.txt",
        # Mathematics and English summaries are .json (structured for LP/A
        # constitutions); all others are plain .txt.
        "chapter_summary":  (
            mirror / f"chapters/{subj_f}/{grade_f}/summaries/ch_{nn}_summary.json"
            if subj_f in _JSON_SUMMARY_SUBJECTS
            else mirror / f"chapters/{subj_f}/{grade_f}/summaries/ch_{nn}_summary.txt"
        ),
        "chapter_mapping":  mirror / f"chapters/{subj_f}/{grade_f}/mappings/ch_{nn}_mapping.json",
    }

def read_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except Exception as e:
        return f"[FILE NOT FOUND: {path}]"

# ── API rates and token logging ───────────────────────────────────────────────

API_RATES_PATH = PROJECT_ROOT / "knowledge_commons/evaluation_mappings/api_rates.json"
TOKEN_LOG_PATH = PROJECT_ROOT / "knowledge_commons/evaluation_mappings/token_log.csv"

@st.cache_data
def load_api_rates() -> dict:
    try:
        return json.loads(API_RATES_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}

def calculate_cost_inr(
    model: str,
    input_tokens: int,
    output_tokens: int,
    cache_write_tokens: int = 0,
    cache_read_tokens:  int = 0,
) -> float:
    rates       = load_api_rates()
    usd_to_inr  = rates.get("usd_to_inr", 84.0)
    model_rates = rates.get("models", {}).get(model, {})
    input_rate  = model_rates.get("input_per_1k_usd",  0.003)
    output_rate = model_rates.get("output_per_1k_usd", 0.015)
    # Prompt caching pricing (1h TTL, verified against Anthropic Console 2026-05-22):
    #   cache write  = 2.00× input rate  (100% surcharge for 1h TTL)
    #   cache read   = 0.10× input rate  (90% discount vs normal)
    cache_write_rate = input_rate * 2.00
    cache_read_rate  = input_rate * 0.10
    cost_usd = (
        (input_tokens       / 1000) * input_rate        +
        (output_tokens      / 1000) * output_rate       +
        (cache_write_tokens / 1000) * cache_write_rate  +
        (cache_read_tokens  / 1000) * cache_read_rate
    )
    return round(cost_usd * usd_to_inr, 4)

def grade_to_roman(grade: str) -> str:
    mapping = {
        "Grade III": "iii", "Grade IV": "iv",  "Grade V":   "v",
        "Grade VI":  "vi",  "Grade VII": "vii", "Grade VIII":"viii",
        "Grade IX":  "ix",
    }
    return mapping.get(grade, grade.lower().replace("grade ", ""))

def log_tokens(
    call_type:              str,
    grade:                  str,
    subject:                str,
    chapter_number:         int,
    chapter_title:          str,
    input_tokens:           int,
    output_tokens:          int,
    model:                  str = "claude-sonnet-4-6",
    cache_write_tokens:     int = 0,
    cache_read_tokens:      int = 0,
):
    cost_inr = calculate_cost_inr(model, input_tokens, output_tokens, cache_write_tokens, cache_read_tokens)
    row = [
        datetime.now().isoformat(timespec="seconds"),
        call_type,
        subject_to_folder(subject),
        grade_to_roman(grade),
        chapter_number,
        chapter_title,
        input_tokens,
        output_tokens,
        input_tokens + output_tokens,
        cost_inr,
        cache_write_tokens,
        cache_read_tokens,
    ]
    try:
        with open(TOKEN_LOG_PATH, "a", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(row)
    except Exception:
        pass  # never crash the app over a logging failure

ASK_ARUVI_LOG_PATH = PROJECT_ROOT / "knowledge_commons/evaluation_mappings/ask_aruvi.csv"

def log_ask_aruvi_tokens(
    session_id:    str,
    query:         str,
    category:      str,
    tab:           str,
    subject:       str,
    grade:         str,
    input_tokens:  int,
    output_tokens: int,
) -> None:
    try:
        cost_inr      = calculate_cost_inr("claude-haiku-4-5-20251001", input_tokens, output_tokens)
        query_snippet = query[:60]
        category_val  = category if category else "none"
        write_header  = not ASK_ARUVI_LOG_PATH.exists()
        with open(ASK_ARUVI_LOG_PATH, "a", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if write_header:
                writer.writerow([
                    "timestamp", "session_id", "tab", "subject", "grade",
                    "category", "query_snippet",
                    "input_tokens", "output_tokens", "total_tokens", "cost_inr",
                ])
            writer.writerow([
                datetime.now().isoformat(timespec="seconds"),
                session_id,
                tab,
                subject,
                grade,
                category_val,
                query_snippet,
                input_tokens,
                output_tokens,
                input_tokens + output_tokens,
                cost_inr,
            ])
    except Exception:
        pass  # never crash the app over a logging failure

# ── Export helpers ────────────────────────────────────────────────────────────

# ── Export helpers ────────────────────────────────────────────────────────────

def add_markdown_content(doc, text):
    """Add markdown text to a python-docx Document — handles headings, bullets, bold."""
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(
                re.sub(r"^\d+\.\s", "", stripped),
                style="List Number"
            )
        else:
            para = doc.add_paragraph()
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for j, part in enumerate(parts):
                run = para.add_run(part)
                if j % 2 == 1:
                    run.bold = True


def generate_docx_bytes_lp(result: dict, chapter: dict, grade: str, subject: str) -> bytes:
    """DOCX export — Lesson Plan only, no Assessment."""
    doc = Document()

    title = doc.add_heading("Aruvi · Lesson Plan", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for i, (lbl, val) in enumerate(zip(
        ["Grade", "Subject", "Chapter", "Chapter Weight"],
        [grade, subject, chapter.get("chapter_title", ""), str(chapter.get("chapter_weight", ""))]
    )):
        meta.rows[i].cells[0].text = lbl
        meta.rows[i].cells[1].text = val
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    lp = ""  # temporarily stubbed — export not yet updated for JSON shape
    # lp = result.get("lesson_plan", "")
    add_markdown_content(doc, lp)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf_bytes_lp(result: dict, chapter: dict, grade: str, subject: str) -> bytes:
    """PDF export — Lesson Plan only, no Assessment."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(44, 62, 80)
    pdf.set_x(10)
    pdf.cell(190, 10, "Aruvi - Lesson Plan", ln=True)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(80, 80, 80)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Grade: {grade}   Subject: {subject}", ln=True)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Chapter: {chapter.get('chapter_title', '')}", ln=True)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Chapter Weight: {chapter.get('chapter_weight', '')}", ln=True)
    pdf.ln(4)
    pdf.set_draw_color(44, 62, 80)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    lp = ""  # temporarily stubbed — export not yet updated for JSON shape
    # lp = result.get("lesson_plan", "")

    pdf.set_text_color(30, 30, 30)
    for line in lp.splitlines():
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        stripped = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        stripped = re.sub(r"^#{1,3}\s+", "", stripped)
        stripped = re.sub(r"^[-*]\s+", "- ", stripped)
        if line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(44, 62, 80)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(44, 62, 80)
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
        pdf.set_x(10)
        try:
            pdf.multi_cell(190, 5, stripped)
        except Exception:
            try:
                pdf.set_x(10)
                pdf.multi_cell(190, 5, stripped.encode("latin-1", "replace").decode("latin-1"))
            except Exception:
                pass  # skip lines that cannot render

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def generate_docx_bytes_assess(result: dict, chapter: dict, grade: str, subject: str) -> bytes:
    """DOCX export — Assessment only, no Lesson Plan."""
    doc = Document()

    title = doc.add_heading("Aruvi · Chapter Assessment", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    for i, (lbl, val) in enumerate(zip(
        ["Grade", "Subject", "Chapter", "Chapter Weight"],
        [grade, subject, chapter.get("chapter_title", ""), str(chapter.get("chapter_weight", ""))]
    )):
        meta.rows[i].cells[0].text = lbl
        meta.rows[i].cells[1].text = val
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    asmt = ""  # temporarily stubbed — export not yet updated for JSON shape
    # asmt = result.get("assessment", "")
    add_markdown_content(doc, asmt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf_bytes_assess(result: dict, chapter: dict, grade: str, subject: str) -> bytes:
    """PDF export — Assessment only, no Lesson Plan."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(44, 62, 80)
    pdf.set_x(10)
    pdf.cell(190, 10, "Aruvi - Chapter Assessment", ln=True)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(80, 80, 80)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Grade: {grade}   Subject: {subject}", ln=True)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Chapter: {chapter.get('chapter_title', '')}", ln=True)
    pdf.set_x(10)
    pdf.cell(190, 6, f"Chapter Weight: {chapter.get('chapter_weight', '')}", ln=True)
    pdf.ln(4)
    pdf.set_draw_color(44, 62, 80)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    asmt = ""  # temporarily stubbed — export not yet updated for JSON shape
    # asmt = result.get("assessment", "")
    pdf.set_text_color(30, 30, 30)
    for line in asmt.splitlines():
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        stripped = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        stripped = re.sub(r"^#{1,3}\s+", "", stripped)
        stripped = re.sub(r"^[-*]\s+", "- ", stripped)
        if line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(44, 62, 80)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(44, 62, 80)
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
        pdf.set_x(10)
        try:
            pdf.multi_cell(190, 5, stripped)
        except Exception:
            try:
                pdf.set_x(10)
                pdf.multi_cell(190, 5, stripped.encode("latin-1", "replace").decode("latin-1"))
            except Exception:
                pass  # skip lines that cannot render

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def format_period_schedule(period_rows: list, session: dict) -> str:
    lines = []
    total_periods = 0
    total_minutes = 0
    for i, rid in enumerate(period_rows):
        dur = session.get(f"dur_sel_{rid}", 40)
        cnt = session.get(f"cnt_{rid}", 1)
        total_periods += cnt
        total_minutes += dur * cnt
        lines.append(
            f"  Row {i+1}: {dur} minutes × {cnt} period{'s' if cnt != 1 else ''} = {dur*cnt} minutes"
        )
    h, m = divmod(total_minutes, 60)
    time_str = f"{h}h {m}min" if h > 0 else f"{m} minutes"
    return (
        f"Period schedule:\n" + "\n".join(lines) +
        f"\nTotal: {total_periods} periods · {time_str}"
    )

# ── Saved plans — local file storage ─────────────────────────────────────────

def _saved_plans_dir(grade: str, subject: str) -> Path:
    subj_f  = subject_to_folder(subject)
    grade_f = grade_to_folder(grade)
    d = PROJECT_ROOT / "mirror" / "saved_plans" / subj_f / grade_f
    d.mkdir(parents=True, exist_ok=True)
    return d

def save_plan(
    grade:       str,
    subject:     str,
    chapter:     dict,
    period_rows: list,
    session:     dict,
    result:      dict,
) -> None:
    d        = _saved_plans_dir(grade, subject)
    nn       = f"{chapter['chapter_number']:02d}"
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"ch_{nn}_{ts}.json"
    sched    = format_period_schedule(period_rows, session)
    payload  = {
        "filename":               filename,
        "saved_at":               datetime.now().isoformat(timespec="seconds"),
        "grade":                  grade,
        "subject":                subject,
        "chapter_number":         chapter["chapter_number"],
        "chapter_title":          chapter.get("chapter_title", ""),
        "period_schedule_display": sched,
        "period_rows_snapshot": [
            {
                "id":       r,
                "duration": session.get(f"dur_sel_{r}", 40),
                "count":    session.get(f"cnt_{r}", 1),
            }
            for r in (period_rows or [])
        ],
        # plan_status is a top-level key (not inside result) so My Plans can
        # read it cheaply without unpacking the result dict. lp_only plans
        # have empty assessment_items; deferred assessment fills them later.
        "plan_status":            result.get("plan_status", "full_lpa"),
        "result": {
            "lesson_plan":      result.get("lesson_plan", {}),
            "coverage_handoff": result.get("coverage_handoff", {}),
            "assessment_items": result.get("assessment_items", []),
            "input_tokens":     result.get("input_tokens", 0),
            "output_tokens":    result.get("output_tokens", 0),
            "cost_inr":         result.get("cost_inr", 0),
            # Assessment-run tokens stored separately so LP and A costs are
            # individually auditable when assessment is deferred.
            "assess_input_tokens":  result.get("assess_input_tokens", 0),
            "assess_output_tokens": result.get("assess_output_tokens", 0),
            "assess_cost_inr":      result.get("assess_cost_inr", 0.0),
        },
    }
    (d / filename).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def update_saved_plan_with_assessment(
    filename:      str,
    grade:         str,
    subject:       str,
    assess_result: dict,
) -> None:
    """Merge deferred assessment items into an existing lp_only saved plan.

    Overwrites the file in place. Idempotent — safe to call twice.
    """
    d    = _saved_plans_dir(grade, subject)
    path = d / filename
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return

    payload["plan_status"] = "full_lpa"
    payload["result"]["assessment_items"]     = assess_result.get("assessment_items", [])
    payload["result"]["assess_input_tokens"]  = assess_result.get("assess_input_tokens", 0)
    payload["result"]["assess_output_tokens"] = assess_result.get("assess_output_tokens", 0)
    payload["result"]["assess_cost_inr"]      = assess_result.get("assess_cost_inr", 0.0)

    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

def load_saved_plans(grade: str, subject: str) -> list:
    d = _saved_plans_dir(grade, subject)
    plans = []
    for f in sorted(d.glob("ch_*.json")):
        try:
            plans.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            pass
    return plans

def delete_saved_plan(grade: str, subject: str, filename: str) -> None:
    d = _saved_plans_dir(grade, subject)
    target = d / filename
    try:
        target.unlink(missing_ok=True)
    except Exception:
        pass

def _build_lpa_prompts_english(
    grade: str,
    subject: str,
    chapter: dict,
    period_sched: str,
    paths: dict,
    include_assessment: bool = False,
) -> tuple[str, str]:
    """Build (system_prompt, user_prompt) for English LP (and optionally A).

    English uses a two-axis schema (main_section × spine). The chapter
    summary is JSON (produced by the cowork prompt
    `chapter_summary_competency_mapping_english.md`) and is the source
    of truth for the LP and assessment. C-codes do not appear in LP/A
    output; the Allocate page reads `spine_to_cg.json` separately.

    When include_assessment=False the assessment constitution is dropped
    from the system prompt and the output schema instruction omits
    assessment_items — the LP-only run produces lesson_plan plus
    coverage_handoff that a later deferred run will consume.
    """
    stage = get_stage(grade)

    lp_const     = read_file(paths["lp_constitution"])
    assess_const = read_file(paths["assessment_const"]) if include_assessment else ""
    pedagogy     = read_file(paths["pedagogy"])
    summary      = read_file(paths["chapter_summary"])

    # Stage-aware rubric depth (Assessment Constitution Rule 10).
    rubric_bullets = (
        "3"   if stage == "preparatory"
        else "3-4" if stage == "middle"
        else "4-5"
    )

    # Stage-aware prompt template fragments. Prep uses the 5-spine prep-native
    # vocabulary; middle uses the 6-spine model. Prep bans ECR, adds the
    # picture_narrative section type, and folds listening into oracy.
    if stage == "preparatory":
        _section_type_enum = "prose|poem|narrative|dialogue|informational|picture_narrative"
        _coverage_handoff_block = (
            '"coverage_handoff": {\n'
            '    "reading":                      { "section_contributions": [<contribution>] },\n'
            '    "oracy":                        { "section_contributions": [...] },\n'
            '    "writing":                      { "section_contributions": [...] },\n'
            '    "word_work": { "section_contributions": [...] },\n'
            '    "beyond_text":                  { "section_contributions": [...] }\n'
            '  }'
        )
        _spine_code_enum  = "reading|oracy|writing|word_work|beyond_text"
        _spine_title_enum = "Reading|Oracy|Writing|Word Work|Beyond the Text"
        _qtype_enum       = "MCQ|SCR|MATCH|FILL_IN|TRUE_FALSE|ORAL_PROMPT|WRITING_TASK|PROJECT"
        _open_types_list  = "ORAL_PROMPT, WRITING_TASK, PROJECT, reflective SCR"
        _bullet_word_cap  = "8"
        _content_sources_long = (
            "prose_summary (prose/narrative/dialogue/informational sections), "
            "poem_text + poem_appreciation_summary (poem sections), or "
            "picture_story_summary + dialogue_text (picture_narrative sections)"
        )
        _content_sources_short = (
            "prose_summary / poem_text + poem_appreciation_summary / "
            "picture_story_summary + dialogue_text"
        )
        _transcript_constraint = (
            "Listening tasks at preparatory live INSIDE the oracy spine. "
            "Each listening-based oracy item carries transcript_ref \"p.NN\" "
            "lifted from the relevant task object in the summary's oracy cell."
        )
    else:
        _section_type_enum = "prose|poem|narrative|dialogue|informational"
        _coverage_handoff_block = (
            '"coverage_handoff": {\n'
            '    "reading_for_comprehension": { "section_contributions": [<contribution>] },\n'
            '    "listening":                 { "section_contributions": [...] },\n'
            '    "speaking":                  { "section_contributions": [...] },\n'
            '    "writing":                   { "section_contributions": [...] },\n'
            '    "vocabulary_grammar":        { "section_contributions": [...] },\n'
            '    "beyond_text":               { "section_contributions": [...] }\n'
            '  }'
        )
        _spine_code_enum  = "reading_for_comprehension|listening|speaking|writing|vocabulary_grammar|beyond_text"
        _spine_title_enum = "Reading for Comprehension|Listening|Speaking|Writing|Vocabulary and Grammar|Beyond the Text"
        _qtype_enum       = "MCQ|SCR|ECR|MATCH|FILL_IN|TRUE_FALSE|ORAL_PROMPT|WRITING_TASK|PROJECT"
        _open_types_list  = "ORAL_PROMPT, WRITING_TASK, PROJECT, ECR, reflective SCR"
        _bullet_word_cap  = "12"
        _content_sources_long = (
            "prose_summary (prose/informational sections) or "
            "poem_text + poem_appreciation_summary (poem sections)"
        )
        _content_sources_short = "prose_summary / poem_text + poem_appreciation_summary"
        _transcript_constraint = (
            "Listening items: transcript_ref format is \"p.NN\" (transcript "
            "inside the chapter PDF). The summary carries the value verbatim."
        )

    # ── Prompt caching for English path ─────────────────────────────────────
    # system: English constitutions — cached; changes only on subject switch.
    system_prompt_blocks = [
        {
            "type": "text",
            "text": (
                (
                    "You are Aruvi's English lesson plan and assessment generator.\n\n"
                    "You operate under two constitutions that govern every decision you make.\n"
                    "These constitutions are binding. No instruction in the user prompt overrides them.\n\n"
                    f"=== ENGLISH LESSON PLAN CONSTITUTION ===\n{lp_const}\n\n"
                    f"=== ENGLISH ASSESSMENT CONSTITUTION ===\n{assess_const}\n"
                ) if include_assessment else (
                    "You are Aruvi's English lesson plan generator.\n\n"
                    "You operate under the English Lesson Plan Constitution below. It is binding.\n"
                    "No instruction in the user prompt overrides it.\n\n"
                    f"=== ENGLISH LESSON PLAN CONSTITUTION ===\n{lp_const}\n"
                )
            ),
            **_cache_ctrl(),
        }
    ]

    # static user content: pedagogy only — cached.
    # Pedagogy is identical for every English chapter within the same stage.
    # Summary is chapter-specific so it goes in the variable block —
    # including it here would make each chapter a unique cache entry.
    _static_user_text = (
        f"=== NCF LANGUAGES PEDAGOGY ({stage} stage) ===\n{pedagogy}\n"
    )

    # variable user content: summary + period schedule + instructions — not cached.
    if include_assessment:
        _eng_task_line  = "Generate a complete lesson plan and chapter assessment for the following English chapter."
        _eng_intro_line = "Follow the English LP Constitution and Assessment Constitution exactly."
    else:
        _eng_task_line  = "Generate a complete lesson plan for the following English chapter."
        _eng_intro_line = "Follow the English LP Constitution exactly."
    _variable_user_text = f"""{_eng_task_line}

=== CHAPTER SUMMARY (JSON, two-axis: main_sections × spines) ===
{summary}

=== TEACHER PERIOD SCHEDULE ===
{period_sched}

=== INSTRUCTIONS ===
{_eng_intro_line}
Produce a SINGLE valid JSON object with this top-level structure:

{{
  "grade": "{grade}",
  "subject": "{subject}",
  "stage": "{stage}",
  "chapter_number": {chapter["chapter_number"]},
  "chapter_title": "{chapter.get('chapter_title', '')}",
  "period_schedule": <derived from teacher period schedule above>,

  "main_sections_inventory": [
    {{ "section_id": "A|B|C", "title": "...", "type": "{_section_type_enum}" }}
  ],

  "periods_allocated": <integer = total period count from the teacher schedule>,

  "lesson_plan": {{
    "periods": [
      <one object per period per LP Constitution Rule 1+2 — each period
       anchors to ONE main_section + 1-2 spines within it; periods walk
       main_sections in textbook order then spines within each section.
       Required fields: period_number, period_duration_minutes,
       section_id, section_title, spines_taught, activity_title,
       pedagogical_methods (object keyed by each spine in spines_taught;
       each value is one method drawn from that spine's permitted list
       in LP Rule 4 for the stage — keys MUST equal spines_taught
       exactly), tasks_in_class (each {{spine, task_index, task_brief}}),
       homework, phases (tile 0..duration with no gaps), teacher_notes
       (2-3 sentences max, grounded in main_section's prose_summary or
       poem_appreciation_summary), materials.>
    ]
  }},

  {_coverage_handoff_block},

  "assessment_items": [
    {{
      "spine_code":  "{_spine_code_enum}",
      "spine_title": "{_spine_title_enum}",
      "note":        "",
      "items": [
        <one item per section_contribution in coverage_handoff for this
         spine (Assessment Constitution Rule 2). Each item tests the
         cell's implied_lo.

         STRICT GENERATION RULES — these override everything else:
         - DO NOT read summary.<spine>.tasks_verbatim[] or question_bank[].
           These fields are FORBIDDEN inputs to the assessment generator.
           Reading either field is a constitution violation regardless of
           intent. The implied_lo in coverage_handoff already encodes
           what was taught.
         - Derive every item solely from the section's content sources
           ({_content_sources_long}) plus the implied_lo from
           coverage_handoff.
         - The item MUST be original — it must not reproduce, paraphrase,
           or structurally echo any textbook exercise wording.
         - The item MUST be visibly grounded in the section's actual
           content: name a character, scene, specific line, grammar
           concept, or writing context drawn from prose_summary /
           poem_text. Generic questions that could apply to any chapter
           are prohibited (Assessment Rule 3).

         Required fields per item: id (e.g. "Q-RFC-A-1"),
         source_section_id, source_section_title, source_section_type,
         source_spine, source_lo (implied_lo copied verbatim from
         coverage_handoff), item_stem (original question grounded in
         section content), question_type (from Assessment Rule 4 set:
         {_qtype_enum}), options ([] unless MCQ or TRUE_FALSE),
         visual_stimulus ("" or pipe-table only), transcript_ref
         (Listening items only; "" otherwise), teacher_guide
         {{suggested_answer (CLOSED non-MCQ items; "" otherwise),
         expected_elements (OPEN items: {rubric_bullets} bullets ≤ {_bullet_word_cap}
         words each; [] otherwise), note ("" unless fallback)}},
         verified (true for open items; true for closed only when
         answer is unambiguously supported by the section's content
         sources).>
      ]
    }}
  ]
}}

CRITICAL CONSTRAINTS:
- Total LP period count = the teacher schedule's period_count. Distribute
  across (section × spine) cells in textbook order (LP Rule 1+2), with
  per-section period share roughly proportional to the section's
  page_count (±1 period tolerance).
- Total assessment item count = number of section_contributions across
  all spines in coverage_handoff that have at least one anchored task
  (one item per spine-cell implied_lo, per Assessment Rule 2). Spines
  with no section_contributions are omitted entirely. For each item,
  read ONLY the cell's implied_lo from coverage_handoff and the
  section's content sources ({_content_sources_short}).
  DO NOT read tasks_verbatim[] or question_bank[] for any purpose —
  these fields are forbidden inputs to the assessment generator.
  Generate one original item per cell grounded in the section content.
- C-codes MUST NOT appear anywhere in the LP or assessment JSON.
- `pedagogical_methods` per period MUST be an object whose keys equal
  `spines_taught` exactly. Each value MUST be drawn from that spine's
  permitted method list in LP Constitution Rule 4 for the {stage}
  stage. Do NOT invent methods. Do NOT collapse multiple spines onto
  a single method.
- {_transcript_constraint}
- The answer layer applies per item. A closed item (MCQ, FILL_IN,
  MATCH, TRUE_FALSE, factual SCR) carries
  `teacher_guide.suggested_answer` (verified against the section's
  content sources; omitted for MCQ — correct option is flagged in
  options[].is_correct). An open item ({_open_types_list}) carries
  `teacher_guide.expected_elements` ({rubric_bullets} short bullets,
  each ≤ {_bullet_word_cap} words). No item carries both fields.

LENGTH CONSTRAINTS:
- Each phase `description`: 2-3 sentences maximum.
- Each `teacher_notes`: 2-3 sentences maximum.
- Each `suggested_answer`: 1-2 sentences plain prose.
- Each `expected_elements` bullet: ≤ {_bullet_word_cap} words.

Output only the raw JSON object. No markdown. No prose. No headers. No ```json fences.
"""

    # LP-only path: strip the assessment_items schema block and the assessment-
    # related critical constraints from the variable user text. The full prompt
    # above stays the source of truth for the LPA path; the surgery below is
    # bounded by unique anchor strings so a constitution edit upstream would
    # surface as a clear failure rather than silent drift.
    if not include_assessment:
        _vt = _variable_user_text
        # Remove the comma after coverage_handoff's closing brace and the entire
        # assessment_items array (from `,\n\n  "assessment_items": [` through
        # the next `]\n}`).
        # Anchors below match the *rendered* f-string output — i.e. each `{{`
        # in source becomes `{` and each `}}` becomes `}`. Do NOT add escapes.
        _start_marker = '  },\n\n  "assessment_items": ['
        _end_marker   = ']\n}\n\nCRITICAL CONSTRAINTS:'
        _si = _vt.find(_start_marker)
        _ei = _vt.find(_end_marker)
        if _si != -1 and _ei != -1:
            _vt = _vt[:_si] + '  }\n}\n\nCRITICAL CONSTRAINTS:' + _vt[_ei + len(_end_marker):]
        # Drop assessment-only critical constraints (the "Total assessment item
        # count …" paragraph through "Generate one original item …").
        _ac_start = '- Total assessment item count'
        _ac_end   = '- C-codes MUST NOT'
        _si2 = _vt.find(_ac_start)
        _ei2 = _vt.find(_ac_end)
        if _si2 != -1 and _ei2 != -1:
            _vt = _vt[:_si2] + _vt[_ei2:]
        # Drop the answer-layer paragraph (closed/open items) — only relevant
        # to assessment items.
        _al_start = '- The answer layer applies per item.'
        _al_end   = '\nLENGTH CONSTRAINTS:'
        _si3 = _vt.find(_al_start)
        _ei3 = _vt.find(_al_end)
        if _si3 != -1 and _ei3 != -1:
            _vt = _vt[:_si3] + _vt[_ei3 + 1:]  # +1 to consume the leading \n of LENGTH
        # Drop assessment-only length constraints (suggested_answer,
        # expected_elements bullets) — they have no effect for LP-only output.
        for _line in (
            "- Each `suggested_answer`: 1-2 sentences plain prose.\n",
            "- Each `expected_elements` bullet: ≤ 12 words.\n",
        ):
            _vt = _vt.replace(_line, "")
        _variable_user_text = _vt

    user_message_blocks = [
        {
            "type": "text",
            "text": _static_user_text,
            **_cache_ctrl(),
        },
        {
            "type": "text",
            "text": _variable_user_text,
        },
    ]

    return system_prompt_blocks, user_message_blocks


def generate_lp_only(
    grade: str,
    subject: str,
    chapter: dict,
    period_rows: list,
    session: dict,
    include_assessment: bool = False,
    result_queue: "queue.Queue | None" = None,
    stop_event:   "threading.Event | None" = None,
) -> dict:
    """Generate lesson plan (and optionally assessment) in one call.

    When include_assessment=False: the system prompt carries only the LP
    Constitution, the output schema instruction omits assessment_items, and
    the returned dict has plan_status="lp_only" with an empty assessment_items
    list. This is the default for the LP/A split (assessment deferred).

    When include_assessment=True: behaves like the original generate_lpa() —
    LP + Assessment in a single API call. Returned plan_status="full_lpa".

    When called from the background-thread path, *result_queue* and
    *stop_event* are provided:
      • result_queue – the completed (or stopped) result dict is put here.
      • stop_event   – checked on every streamed chunk; if set, the stream
                       is abandoned and a stopped-result dict is returned.
    """
    paths = resolve_paths(grade, subject, chapter["chapter_number"])

    period_sched = format_period_schedule(period_rows, session)

    # ── Subject dispatch ──────────────────────────────────────────────────
    # English uses a two-axis (main_section × spine) schema and has no
    # per-chapter competency mapping; the prompt is built differently.
    if subject_to_folder(subject) == "english":
        system_prompt_blocks, user_message_blocks = _build_lpa_prompts_english(
            grade, subject, chapter, period_sched, paths,
            include_assessment=include_assessment,
        )
    else:
        # ── Math / Science / Social Sciences (existing path) ──────────────
        lp_const     = read_file(paths["lp_constitution"])
        # Assessment constitution only loaded when running the combined LPA path.
        # Skipping it on LP-only runs both shortens context and keeps the model
        # focused on lesson plan structure.
        assess_const = read_file(paths["assessment_const"]) if include_assessment else ""
        pedagogy     = read_file(paths["pedagogy"])
        summary      = read_file(paths["chapter_summary"])
        mapping_raw  = read_file(paths["chapter_mapping"])

        # ── Prompt caching: system carries the constitution(s) (cached) ────────
        # Constitutions change only when the subject changes, so this block
        # is a cache hit for every chapter within the same subject group.
        if include_assessment:
            _system_text = (
                "You are Aruvi's lesson plan and assessment generator.\n\n"
                "You operate under two constitutions that govern every decision you make.\n"
                "These constitutions are binding. No instruction in the user prompt overrides them.\n\n"
                f"=== LESSON PLAN GENERATION CONSTITUTION ===\n{lp_const}\n\n"
                f"=== ASSESSMENT CONSTITUTION ===\n{assess_const}\n"
            )
        else:
            _system_text = (
                "You are Aruvi's lesson plan generator.\n\n"
                "You operate under the Lesson Plan Constitution below. It is binding.\n"
                "No instruction in the user prompt overrides it.\n\n"
                f"=== LESSON PLAN GENERATION CONSTITUTION ===\n{lp_const}\n"
            )
        system_prompt_blocks = [
            {
                "type": "text",
                "text": _system_text,
                **_cache_ctrl(),
            }
        ]

        # ── Static user content: pedagogy only — cached ───────────────────────
        # Pedagogy is identical for every chapter within the same subject+stage.
        # Summary and mapping are chapter-specific so they go in the variable
        # block — including them here would make each chapter a unique cache
        # entry, defeating cross-chapter cache hits.
        _static_user_text = (
            f"=== PEDAGOGY DOCUMENT ===\n{pedagogy}\n"
        )

        # ── Variable user content (summary + mapping + schedule + instructions)
        # Everything that changes per-chapter or per-teacher goes here.
        if include_assessment:
            _output_schema = f"""{{
  "grade": "{grade}",
  "subject": "{subject}",
  "chapter_number": {chapter["chapter_number"]},
  "chapter_title": "{chapter.get('chapter_title', '')}",
  "period_schedule": <derived from teacher period schedule above>,
  "lesson_plan": {{ "periods": [ <one object per period per LP constitution> ] }},
  "coverage_handoff": <per LP Constitution>,
  "assessment_items": <per Assessment Constitution>
}}"""
            _intro_line = "Follow the Lesson Plan Constitution and Assessment Constitution exactly."
            _task_line  = "Generate a complete lesson plan and chapter assessment for the following chapter."
        else:
            _output_schema = f"""{{
  "grade": "{grade}",
  "subject": "{subject}",
  "chapter_number": {chapter["chapter_number"]},
  "chapter_title": "{chapter.get('chapter_title', '')}",
  "period_schedule": <derived from teacher period schedule above>,
  "lesson_plan": {{ "periods": [ <one object per period per LP constitution> ] }},
  "coverage_handoff": <per LP Constitution>
}}"""
            _intro_line = "Follow the Lesson Plan Constitution exactly."
            _task_line  = "Generate a complete lesson plan for the following chapter."

        _variable_user_text = f"""{_task_line}

=== CHAPTER SUMMARY ===
{summary}

=== CHAPTER MAPPING JSON ===
{mapping_raw}

=== TEACHER PERIOD SCHEDULE ===
{period_sched}

=== INSTRUCTIONS ===
{_intro_line}
Produce your entire output as a single valid JSON object with this top-level structure:

{_output_schema}

LENGTH CONSTRAINTS (strictly enforced to keep output compact):
- Each phase `description`: 2–3 sentences maximum.
- Each `teacher_notes` field: 2–3 sentences maximum.

Output only the raw JSON object. No markdown. No prose. No section headers. No ```json fences.
"""

        user_message_blocks = [
            {
                "type": "text",
                "text": _static_user_text,
                **_cache_ctrl(),
            },
            {
                "type": "text",
                "text": _variable_user_text,
            },
        ]

    try:
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

        full_output = ""
        input_tokens = 0
        output_tokens = 0

        import time as _time
        progress_placeholder = st.empty()
        # Separate placeholder for the live timer badge — rendered via
        # components.html so that JS actually executes (st.markdown strips
        # <script> tags, which is why the badge previously stayed at 00:00).
        timer_placeholder = st.empty()

        # ── Record start time (ms epoch) for the timer ────────────────────────
        _gen_start_ms = int(_time.time() * 1000)

        # ── Shared CSS (animations + Deploy button suppression) ──────────────
        # The Deploy button is hidden for the entire duration of generation by
        # embedding the hide rule in _pcss, which is prepended to every progress
        # markdown update. When progress_placeholder.empty() is called at the end
        # (or on error / stop), the <style> tag disappears and the button returns.
        # Animations only — safe to keep in the DOM both during and after a run.
        _pcss_anim = (
            "<style>"
            "@keyframes aruviPulse{0%,100%{opacity:1}50%{opacity:.3}}"
            "@keyframes spin{to{transform:rotate(360deg)}}"
            "</style>"
        )
        # Deploy/toolbar hide rule — only emitted while generation is ACTIVE.
        # Multiple selectors cover the rename across Streamlit versions
        # (stDeployButton -> stAppDeployButton) and the wrapping stToolbar.
        # This MUST NOT be included in the "Generation complete" markdown — once
        # the run ends and the placeholder is not cleared (success path), any
        # lingering hide rule would leave Deploy permanently hidden/inert.
        _pcss_hide = (
            "<style>"
            "body [data-testid='stDeployButton'],"
            "body [data-testid='stAppDeployButton'],"
            "body [data-testid='stToolbar']"
            "{display:none!important;visibility:hidden!important;}"
            "</style>"
        )
        _pcss = _pcss_anim + _pcss_hide

        # ── Live timer badge (separate components.html block) ────────────────
        # Streamlit's components.html mounts in a sandboxed iframe and DOES run
        # scripts, but st.markdown(unsafe_allow_html=True) silently strips
        # <script> — that's why the badge previously stayed at 00:00. We give
        # the iframe a real height so its content is visible, and use Streamlit
        # CSS to float the iframe wrapper at the bottom-right of the viewport.
        _timer_widget_html = (
            '<!doctype html><html><head><style>'
            'html,body{margin:0;padding:0;background:transparent;text-align:right;}'
            '#aruvi-timer{'
            '  font-family:monospace;font-size:10px;color:#000000;'
            '  background:transparent;border:none;border-radius:0;'
            '  padding:3px 8px;display:inline-block;margin:4px 6px 0 0;'
            '}'
            '</style></head><body>'
            '<div id="aruvi-timer">00:00</div>'
            '<script>'
            '(function(){'
            f'  var _start={_gen_start_ms};'
            '  function _tick(){'
            '    var el=document.getElementById("aruvi-timer");'
            '    if(!el){return;}'
            '    var s=Math.floor((Date.now()-_start)/1000);'
            '    var m=Math.floor(s/60);var sc=s%60;'
            '    el.textContent=(m<10?"0"+m:m)+":"+(sc<10?"0"+sc:sc);'
            '  }'
            '  _tick();'
            '  setInterval(_tick,1000);'
            '})();'
            '</script>'
            '</body></html>'
        )
        # Float the iframe over the popup at bottom-right of the viewport.
        # We rely on the fact that Streamlit wraps each components.html call
        # in <iframe srcdoc="…">, so we can match by the unique 'aruvi-timer'
        # string in the srcdoc attribute.
        st.markdown(
            "<style>"
            "iframe[srcdoc*='aruvi-timer']{"
            "  position:fixed !important;"
            "  top:340px !important;"
            "  right:24px !important;"
            "  width:96px !important;"
            "  height:30px !important;"
            "  border:0 !important;"
            "  z-index:10000 !important;"
            "  background:transparent !important;"
            "}"
            "</style>",
            unsafe_allow_html=True,
        )
        with timer_placeholder.container():
            components.html(_timer_widget_html, height=30, scrolling=False)
        # _timer_js / _timer_badge kept as no-ops so the existing markdown
        # blocks (which append them) continue to work without the JS injection.
        _timer_js = ""
        _timer_badge = (
            '<div style="display:flex;justify-content:flex-end;padding:4px 12px 8px 0;">'
            '<span style="font-family:monospace;font-size:10px;color:transparent;'
            'padding:2px 6px;">&nbsp;</span></div>'
        )
        # ── Icon snippets ─────────────────────────────────────────────────────
        _tick_icon = (
            '<div style="width:12px;height:12px;border-radius:50%;background:#d4f0e4;'
            'display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px;">'
            '<div style="width:5px;height:3px;border-left:1.5px solid #2d8a5e;'
            'border-bottom:1.5px solid #2d8a5e;transform:rotate(-45deg);'
            'margin-top:-1px;"></div></div>'
        )
        _spin_icon = (
            '<div style="width:12px;height:12px;border-radius:50%;'
            'border:1.5px solid #e8a83e;border-top-color:transparent;'
            'animation:spin 0.7s linear infinite;flex-shrink:0;'
            'margin-top:1px;box-sizing:border-box;"></div>'
        )
        _dot_icon = (
            '<div style="width:12px;height:12px;display:flex;align-items:center;'
            'justify-content:center;flex-shrink:0;margin-top:1px;">'
            '<div style="width:6px;height:6px;background:#d9d6d0;border-radius:50%;"></div></div>'
        )

        def _row_done(text):
            return (
                f'<div style="display:flex;align-items:flex-start;gap:8px;'
                f'font-size:12px;color:#9c9895;">{_tick_icon}<span>{text}</span></div>'
            )
        def _row_active(text):
            return (
                f'<div style="display:flex;align-items:flex-start;gap:8px;'
                f'font-size:12px;color:#3d3b38;font-weight:500;">'
                f'{_spin_icon}<span>{text}</span></div>'
            )
        def _row_pending(text):
            return (
                f'<div style="display:flex;align-items:flex-start;gap:8px;'
                f'font-size:12px;opacity:0.45;">{_dot_icon}<span>{text}</span></div>'
            )

        # ── Pre-compute totals for tick-based proximity tracking ─────────────
        # Total LP periods: sum of counts across all period_rows.
        _total_periods = sum(session.get(f"cnt_{rid}", 1) for rid in period_rows) if period_rows else 0

        # Assessment group totals per subject (known before the API call):
        #   Science   → unique progression_stages (3 or 4 — read from LP stream)
        #   SS        → competency count from mapping JSON
        #   Maths     → always 3 sections (A, B, C)
        #   English   → always 6 spines
        _subj_folder = subject_to_folder(subject)
        if _subj_folder == "social_sciences":
            try:
                _mapping_obj = json.loads(mapping_raw) if mapping_raw else {}
                _total_assess_groups = len(_mapping_obj.get("competencies", []))
            except Exception:
                _total_assess_groups = 0
        elif _subj_folder == "mathematics":
            # Sections A/B/C are structurally fixed today, but derive from
            # coverage_handoff at the phase-2 transition to stay future-proof.
            # Start at 0; set when "assessment_items" is detected in the stream.
            _total_assess_groups = 0   # will be set at phase-2 transition
        elif _subj_folder == "english":
            # Spine count varies by chapter — only spines with LP contributions
            # are emitted. Derived from coverage_handoff at the LP→assessment
            # transition (coverage_handoff keys = present spines). Start at 0;
            # set when "assessment_items" is detected in the stream.
            _total_assess_groups = 0   # will be set at phase-2 transition
        elif _subj_folder == "the_world_around_us":
            # TWAU: one item per period — total known upfront from period schedule.
            _total_assess_groups = _total_periods
        else:
            # Science: stage count not in mapping — discovered from LP stream.
            # Start at 0; updated when "progression_stage" tokens appear in LP portion.
            _total_assess_groups = 0   # will be set once LP portion completes

        # ── Tick-row helpers ──────────────────────────────────────────────────
        # Renders a compact row of green tick pills: e.g. "✓ ✓ ✓ · 3 of 8 done"
        def _ticks_row(done: int, total: int, label: str) -> str:
            if total <= 0:
                return (
                    f'<div style="display:flex;align-items:flex-start;gap:8px;'
                    f'font-size:12px;color:#3d3b38;font-weight:500;">'
                    f'{_spin_icon}<span>{label}</span></div>'
                )
            # Tick pills — green filled for done, grey outline for remaining
            _pill_done = (
                'display:inline-flex;align-items:center;justify-content:center;'
                'width:14px;height:14px;border-radius:3px;background:#d4f0e4;'
                'font-size:8px;color:#2d8a5e;font-weight:700;margin:0 1px;flex-shrink:0;'
            )
            _pill_todo = (
                'display:inline-flex;align-items:center;justify-content:center;'
                'width:14px;height:14px;border-radius:3px;background:#f2f0ec;'
                'border:1px solid #d9d6d0;margin:0 1px;flex-shrink:0;'
            )
            pills = "".join(
                f'<span style="{_pill_done}">✓</span>' if i < done
                else f'<span style="{_pill_todo}"></span>'
                for i in range(total)
            )
            fraction = f'<span style="font-size:10px;color:#5c5a56;margin-left:4px;">{done}&thinsp;/&thinsp;{total}</span>'
            spin = _spin_icon if done < total else _tick_icon
            return (
                f'<div style="display:flex;align-items:flex-start;gap:8px;'
                f'font-size:12px;color:#3d3b38;font-weight:500;">'
                f'{spin}'
                f'<div style="display:flex;flex-direction:column;gap:3px;">'
                f'<span>{label}</span>'
                f'<div style="display:flex;align-items:center;flex-wrap:wrap;">'
                f'{pills}{fraction}'
                f'</div></div></div>'
            )

        # LP/A split — choose the 6-step (LPA) or 5-step (LP-only) progress list.
        # Branch off `include_assessment` so the LP-only path doesn't display
        # an Assessment step that never fires.
        _steps_lpa = [
            "Reading LP &amp; Assessment Constitutions",
            "Reading chapter summary",
            "Loading matched competencies",
            "Loading stage pedagogy",
            "Building period-by-period activities&#8230;",
            "Writing assessment questions",
        ]
        _steps_lp = [
            "Reading LP Constitution",
            "Reading chapter summary",
            "Loading matched competencies",
            "Loading stage pedagogy",
            "Building period-by-period activities&#8230;",
        ]
        _steps = _steps_lpa if include_assessment else _steps_lp
        _note_html = (
            '<div style="display:flex;flex-direction:column;align-items:center;'
            'gap:0.5rem;padding:8px 0 4px 0;">'
            f'<img src="{_rotate_logo_src}" style="width:48px;height:48px;'
            'animation:aruviSpin 3.5s linear infinite;" alt="Aruvi">'
            '<style>@keyframes aruviSpin{from{transform:rotate(0deg);}to{transform:rotate(360deg);}}</style>'
            '<span style="font-size:12px;color:#5c5a56;font-style:italic;">'
            'Running in the background \u2014 keep this tab open until complete.'
            '</span></div>'
        )
        _box_open = (
            '<div class="aruvi-progress-box" style="position:fixed;top:80px;right:24px;'
            'width:280px;z-index:9999;background:white;border:1px solid #d9d6d0;'
            'border-radius:10px;overflow:hidden;">'
        )
        _hdr_working = (
            '<div style="padding:8px 12px;border-bottom:1px solid #ece9e4;'
            'display:flex;gap:8px;align-items:center;">'
            '<div style="width:10px;height:10px;border-radius:50%;background:#e8a83e;'
            'animation:aruviPulse 1.4s infinite;flex-shrink:0;"></div>'
            '<span style="font-size:11px;color:#7a776f;font-weight:500;flex:1;">'
            'Aruvi is working&#8230;</span>'
            # Directly click the hidden Streamlit stop button (wrapper carries the
            # auto-assigned 'st-key-btn_stop_generation' class). postMessage to
            # window.parent did not work because the only listener lived inside
            # a components.html child iframe, which never receives messages
            # posted to its parent window.
            '<button onclick="(function(){'
            'var w=document.querySelector(\'[class*=st-key-btn_stop_generation]\');'
            'if(w){var b=w.querySelector(\'button\');if(b)b.click();}'
            '})()" '
            'style="display:inline-flex;align-items:center;gap:4px;font-size:10px;'
            'color:#9c9895;background:#f2f0ec;border:1px solid #dddad5;'
            'border-radius:4px;padding:3px 7px;cursor:pointer;white-space:nowrap;'
            'font-family:inherit;line-height:1;">'
            '<span style="width:7px;height:7px;background:#9c9895;border-radius:1px;'
            'display:inline-block;flex-shrink:0;"></span>'
            'stop</button>'
            '</div>'
        )
        _body_open = (
            '<div style="padding:10px 12px 12px;display:flex;flex-direction:column;gap:2px;">'
        )

        # Build a step-row block where indices < active_idx are done, exactly
        # one (active_idx) is spinning, and the rest are pending. Works for
        # both the 6-step LPA list and the 5-step LP-only list.
        # lp_ticks_row / assess_ticks_row: when provided, replace the plain
        # row text for step index 4 (LP) and 5 (assessment) with tick rows.
        def _steps_block(active_idx: int,
                         lp_ticks_row: str = "",
                         assess_ticks_row: str = "") -> str:
            parts = []
            for _i, _s in enumerate(_steps):
                if _i < active_idx:
                    parts.append(_row_done(_s))
                elif _i == active_idx:
                    # Use tick-row override if provided for LP (4) or assess (5)
                    if _i == 4 and lp_ticks_row:
                        parts.append(lp_ticks_row)
                    elif _i == 5 and assess_ticks_row:
                        parts.append(assess_ticks_row)
                    else:
                        parts.append(_row_active(_s))
                else:
                    parts.append(_row_pending(_s))
            return "".join(parts)

        def _progress_html(active_idx: int,
                           lp_ticks_row: str = "",
                           assess_ticks_row: str = "") -> str:
            return (
                _pcss + _box_open + _hdr_working + _body_open
                + _steps_block(active_idx,
                               lp_ticks_row=lp_ticks_row,
                               assess_ticks_row=assess_ticks_row)
                + _note_html
                + '</div>'
                + _timer_badge
                + '</div>'
                + _timer_js
            )

        # Phase 1: 4 ticked · step 5 (activities) active · step 6 (assessment)
        # pending. For LP-only runs (5 steps), this also functions as the
        # terminal "working" view.
        PROGRESS_HTML_WORKING = _progress_html(4)
        # Phase 2: only fires when include_assessment is True.
        PROGRESS_HTML_ASSESSMENT_ACTIVE = (
            _progress_html(5) if include_assessment else PROGRESS_HTML_WORKING
        )

        # Stage 0 — step 1 active
        progress_placeholder.markdown(_progress_html(0), unsafe_allow_html=True)
        _time.sleep(5)

        # Stage 1 — step 2 active
        progress_placeholder.markdown(_progress_html(1), unsafe_allow_html=True)
        _time.sleep(5)

        # Stage 2 — step 3 active
        progress_placeholder.markdown(_progress_html(2), unsafe_allow_html=True)
        _time.sleep(5)

        # Stage 3 — step 4 active
        progress_placeholder.markdown(_progress_html(3), unsafe_allow_html=True)
        _time.sleep(5)

        # Stage 4 — step 5 (activities) active
        progress_placeholder.markdown(PROGRESS_HTML_WORKING, unsafe_allow_html=True)

        # ── Stream loop ───────────────────────────────────────────────────────
        streamed_text = ""
        _assessment_triggered = False
        _stopped_by_user = False

        # Tick-tracking counters — updated each time a new period or assessment
        # group is detected in the accumulating stream text.
        _lp_periods_seen    = 0   # count of "period_number" tokens in LP portion
        _assess_groups_seen = 0   # count of assessment group tokens post-trigger

        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=32000,
            system=system_prompt_blocks,
            messages=[{"role": "user", "content": user_message_blocks}],
            extra_headers={
                "anthropic-beta": "prompt-caching-2024-07-31"
            },
        ) as stream:
            for text in stream.text_stream:
                # ── Stop-button check ─────────────────────────────────────────
                if stop_event is not None and stop_event.is_set():
                    _stopped_by_user = True
                    break
                streamed_text += text

                # ── Phase-2 transition: LP done → assessment starting ─────────
                if (include_assessment
                        and not _assessment_triggered
                        and '"assessment_items"' in streamed_text):
                    _assessment_triggered = True
                    import re as _re_trans
                    # Science: unique progression_stage values in the LP portion.
                    if _subj_folder == "science":
                        _stage_hits = _re_trans.findall(
                            r'"progression_stage"\s*:\s*(\d+)', streamed_text
                        )
                        _total_assess_groups = len(set(int(x) for x in _stage_hits))
                    # SS: 3 fixed weight bands (Central / Substantive / Present).
                    elif _subj_folder == "social_sciences":
                        _total_assess_groups = 3
                    # English: distinct spine_code values in LP/coverage_handoff
                    # portion — only present spines appear there.
                    elif _subj_folder == "english":
                        _spine_hits = _re_trans.findall(
                            r'"spine_code"\s*:\s*"([^"]+)"', streamed_text
                        )
                        _total_assess_groups = len(set(_spine_hits)) if _spine_hits else 6
                    # Maths: distinct single-letter section keys (section_a/b/c)
                    # in coverage_handoff — middle uses section_a/b/c keys;
                    # prep uses intent keys (explore/reason/practice/solve).
                    elif _subj_folder == "mathematics":
                        _sect_hits = _re_trans.findall(
                            r'"(section_[a-z])"(?:\s*:)', streamed_text
                        )
                        _intent_hits = _re_trans.findall(
                            r'"(explore|reason|practice|solve)"(?:\s*:)', streamed_text
                        )
                        if _sect_hits:
                            _total_assess_groups = len(set(_sect_hits))
                        elif _intent_hits:
                            _total_assess_groups = len(set(_intent_hits))
                        else:
                            _total_assess_groups = 3
                    progress_placeholder.markdown(
                        _progress_html(5, assess_ticks_row=_ticks_row(
                            0, _total_assess_groups, "Writing assessment questions"
                        )),
                        unsafe_allow_html=True,
                    )

                elif not _assessment_triggered:
                    # ── LP tick tracking ──────────────────────────────────────
                    # Tick period N when period N+1 starts — i.e. done =
                    # (occurrences of "period_number") - 1, so the final period
                    # only ticks after the stream completes (see post-loop below).
                    _new_lp  = streamed_text.count('"period_number"')
                    _capped  = min(max(_new_lp - 1, 0), _total_periods) if _total_periods else max(_new_lp - 1, 0)
                    if _capped != _lp_periods_seen:
                        _lp_periods_seen = _capped
                        progress_placeholder.markdown(
                            _progress_html(4, lp_ticks_row=_ticks_row(
                                _lp_periods_seen, _total_periods,
                                "Building period-by-period activities"
                            )),
                            unsafe_allow_html=True,
                        )

                else:
                    # ── Assessment tick tracking ──────────────────────────────
                    # Tick group N when group N+1 starts (done = raw_count - 1),
                    # so the final group only ticks after the stream ends.
                    # Token varies by subject:
                    #   Science   → "progression_stage" (unique int values)
                    #   SS        → "weight_label"      (unique string values: 3 bands)
                    #   Maths     → "section_code"      (occurrence count)
                    #   English   → "spine_code"        (occurrence count)
                    _assess_text = streamed_text[
                        streamed_text.index('"assessment_items"'):
                    ]
                    if _subj_folder == "science":
                        import re as _re2
                        _ag = len(set(int(x) for x in _re2.findall(
                            r'"progression_stage"\s*:\s*(\d+)', _assess_text
                        )))
                    elif _subj_folder == "social_sciences":
                        import re as _re3
                        _ag = len(set(_re3.findall(
                            r'"weight_label"\s*:\s*"([^"]+)"', _assess_text
                        )))
                    elif _subj_folder == "mathematics":
                        _ag = _assess_text.count('"section_code"')
                    elif _subj_folder == "the_world_around_us":
                        # One item per period — count implied_lo occurrences
                        _ag = _assess_text.count('"implied_lo"')
                    else:  # english
                        _ag = _assess_text.count('"spine_code"')
                    # N-1: tick N fires when group N+1 starts
                    _ag_done   = max(_ag - 1, 0)
                    _ag_capped = min(_ag_done, _total_assess_groups) if _total_assess_groups else _ag_done
                    if _ag_capped != _assess_groups_seen:
                        _assess_groups_seen = _ag_capped
                        progress_placeholder.markdown(
                            _progress_html(5, assess_ticks_row=_ticks_row(
                                _assess_groups_seen, _total_assess_groups,
                                "Writing assessment questions"
                            )),
                            unsafe_allow_html=True,
                        )

        full_output = streamed_text

        # ── Final LP tick: mark all periods done once stream completes ────────
        # The last period's tick is intentionally deferred until here — it only
        # fires when the full LP content has been streamed (not when its opening
        # "period_number" token first appeared).
        if not _stopped_by_user and not _assessment_triggered and _total_periods > 0:
            if _lp_periods_seen < _total_periods:
                _lp_periods_seen = _total_periods
                progress_placeholder.markdown(
                    _progress_html(4, lp_ticks_row=_ticks_row(
                        _lp_periods_seen, _total_periods,
                        "Building period-by-period activities"
                    )),
                    unsafe_allow_html=True,
                )

        # ── Final assessment tick: mark all groups done once stream completes ──
        # Same deferred-last logic as LP — the final group's tick fires here,
        # not when its opening token first appeared mid-stream.
        if not _stopped_by_user and _assessment_triggered and _total_assess_groups > 0:
            if _assess_groups_seen < _total_assess_groups:
                _assess_groups_seen = _total_assess_groups
                progress_placeholder.markdown(
                    _progress_html(5, assess_ticks_row=_ticks_row(
                        _assess_groups_seen, _total_assess_groups,
                        "Writing assessment questions"
                    )),
                    unsafe_allow_html=True,
                )

        # ── If user stopped, record partial tokens and return a stopped result ─
        if _stopped_by_user:
            try:
                _partial_usage = stream.get_current_message_snapshot().usage
                input_tokens        = getattr(_partial_usage, "input_tokens",                0)
                output_tokens       = getattr(_partial_usage, "output_tokens",               0)
                _cache_write_stopped = getattr(_partial_usage, "cache_creation_input_tokens", 0)
                _cache_read_stopped  = getattr(_partial_usage, "cache_read_input_tokens",     0)
            except Exception:
                input_tokens         = 0
                output_tokens        = 0
                _cache_write_stopped = 0
                _cache_read_stopped  = 0
            log_tokens(
                call_type          = ("lpa_generation_stopped"
                                       if include_assessment
                                       else "lp_generation_stopped"),
                grade              = grade,
                subject            = subject,
                chapter_number     = chapter["chapter_number"],
                chapter_title      = chapter.get("chapter_title", ""),
                input_tokens       = input_tokens,
                output_tokens      = output_tokens,
                model              = "claude-sonnet-4-6",
                cache_write_tokens = _cache_write_stopped,
                cache_read_tokens  = _cache_read_stopped,
            )
            progress_placeholder.empty()
            timer_placeholder.empty()
            _stopped_result = {
                "grade":          grade,
                "subject":        subject,
                "chapter_number": chapter["chapter_number"],
                "chapter_title":  chapter.get("chapter_title", ""),
                "lesson_plan":    {},
                "coverage_handoff": {},
                "assessment_items": [],
                "input_tokens":   input_tokens,
                "output_tokens":  output_tokens,
                "cost_inr":       calculate_cost_inr("claude-sonnet-4-6", input_tokens, output_tokens),
                "plan_status":    "full_lpa" if include_assessment else "lp_only",
                "stopped":        True,
            }
            if result_queue is not None:
                result_queue.put(_stopped_result)
            return _stopped_result

        usage = stream.get_final_message().usage
        input_tokens       = usage.input_tokens
        output_tokens      = usage.output_tokens
        cache_write_tokens = getattr(usage, "cache_creation_input_tokens", 0)
        cache_read_tokens  = getattr(usage, "cache_read_input_tokens",     0)

        log_tokens(
            call_type          = "lpa_generation" if include_assessment else "lp_generation",
            grade              = grade,
            subject            = subject,
            chapter_number     = chapter["chapter_number"],
            chapter_title      = chapter.get("chapter_title", ""),
            input_tokens       = input_tokens,
            output_tokens      = output_tokens,
            model              = "claude-sonnet-4-6",
            cache_write_tokens = cache_write_tokens,
            cache_read_tokens  = cache_read_tokens,
        )

        parsed = {}
        _raw = full_output.strip()
        # Strip ```json ... ``` fences if the model wrapped output despite instructions
        if _raw.startswith("```"):
            _fence_end = _raw.find("```", 3)
            _raw = (_raw[_raw.index("\n") + 1 : _fence_end] if _fence_end > 3 else _raw).strip()
        # Strip any prose preamble before the opening brace (model sometimes
        # reasons aloud before emitting JSON despite instructions)
        _brace_pos = _raw.find("{")
        if _brace_pos > 0:
            _raw = _raw[_brace_pos:]
        try:
            parsed = json.loads(_raw)
            # ── Final elapsed time (computed server-side, baked into HTML) ──
            _elapsed_s = max(0, int(_time.time() * 1000 - _gen_start_ms) // 1000)
            _final_mmss = f"{_elapsed_s // 60:02d}:{_elapsed_s % 60:02d}"
            # ── Phase 3: completion box (built after parse so numbers are real) ─
            _n_periods = len(parsed.get("lesson_plan", {}).get("periods", []))
            _n_acts    = sum(
                len(p.get("activities", []))
                for p in parsed.get("lesson_plan", {}).get("periods", [])
            )
            _c_codes   = {
                (_c.get("c_code", "") if isinstance(_c, dict) else str(_c))
                for _p2 in parsed.get("lesson_plan", {}).get("periods", [])
                for _c  in _p2.get("competencies", [])
            }
            _n_comps   = len(_c_codes - {""})
            _n_qs      = len(parsed.get("assessment_items", []))
            _sv        = "color:#2d8a5e;font-weight:500;"
            _assess_line = (
                f'Assessment: <span style="{_sv}">{_n_qs}</span> questions'
                if include_assessment
                else 'Assessment: <span style="color:#9c9895;">deferred</span>'
            )
            _summary   = (
                '<div style="font-family:monospace;font-size:10.5px;background:#f7f5f2;'
                'border-radius:6px;padding:8px 10px;margin-bottom:8px;">'
                f'Lesson plan: <span style="{_sv}">{_n_periods}</span> periods'
                f' &#183; <span style="{_sv}">{_n_acts}</span> activities<br>'
                f'Competencies: <span style="{_sv}">{_n_comps}</span> mapped<br>'
                f'{_assess_line}'
                '</div>'
            )
            _tick_sm   = (
                '<div style="width:10px;height:10px;border-radius:50%;background:#d4f0e4;'
                'display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px;">'
                '<div style="width:4px;height:2.5px;border-left:1.5px solid #2d8a5e;'
                'border-bottom:1.5px solid #2d8a5e;transform:rotate(-45deg);'
                'margin-top:-0.5px;"></div></div>'
            )
            def _row_sm(text):
                return (
                    f'<div style="display:flex;align-items:flex-start;gap:8px;'
                    f'font-size:11px;color:#b0ada8;">'
                    f'{_tick_sm}<span>{text}</span></div>'
                )
            PROGRESS_HTML_DONE = (
                # Use animations-only style here — NOT _pcss — so the Deploy
                # button + toolbar reappear as soon as generation completes.
                _pcss_anim
                + '<div class="aruvi-progress-box" style="position:fixed;top:80px;right:24px;'
                'width:280px;z-index:9999;background:white;border:1px solid #d9d6d0;'
                'border-radius:10px;overflow:hidden;">'
                '<div style="padding:8px 12px;border-bottom:1px solid #ece9e4;'
                'display:flex;gap:8px;align-items:center;">'
                '<div style="width:10px;height:10px;border-radius:50%;background:#2d8a5e;'
                'flex-shrink:0;"></div>'
                '<span style="font-size:11px;color:#7a776f;font-weight:500;flex:1;">'
                'Generation complete</span>'
                '<button onclick="this.closest(\'.aruvi-progress-box\').style.display=\'none\'"'
                ' style="font-size:10px;background:#f2f0ec;border:1px solid #dddad5;'
                'border-radius:4px;padding:3px 7px;cursor:pointer;color:#9c9895;">'
                'collapse &#8250;</button>'
                '</div>'
                '<div style="padding:10px 12px 12px;display:flex;flex-direction:column;gap:3px;">'
                + _summary
                + "".join(_row_sm(_s) for _s in _steps)
                + '<div style="display:flex;justify-content:flex-end;padding:4px 12px 8px 0;">'
                '<span style="font-family:monospace;font-size:10px;color:#2d8a5e;'
                'background:#f0faf5;border:1px solid #b8e8d0;border-radius:4px;'
                f'padding:2px 6px;" id="aruvi-timer-final">{_final_mmss}</span>'
                '</div>'
                + '</div></div>'
            )
            progress_placeholder.markdown(PROGRESS_HTML_DONE, unsafe_allow_html=True)
            # Stop the live ticker iframe and remove the floating badge it
            # injected into the parent document, so the only visible time is
            # the static green final-time inside the completion box.
            # Clear the live ticker iframe — the static green final-time inside
            # the completion box now shows the elapsed time.
            timer_placeholder.empty()
        except Exception as _je:
            progress_placeholder.empty()
            # Tear down the live timer + floating badge on the error path too,
            # so the ticker doesn't keep running after generation has aborted.
            try:
                timer_placeholder.empty()
                components.html(
                    '<script>try{'
                    'if(window.parent && window.parent.document){'
                    '  var f=window.parent.document.getElementById("aruvi-timer-fixed");'
                    '  if(f){f.remove();}'
                    '}'
                    '}catch(e){}</script>',
                    height=0,
                    scrolling=False,
                )
            except Exception:
                pass
            # ── DEBUG: dump full raw output to file so we can inspect it ──────
            try:
                import datetime as _dt
                _debug_dir = Path(__file__).parent.parent / "mirror" / "debug"
                _debug_dir.mkdir(parents=True, exist_ok=True)
                _debug_file = _debug_dir / f"raw_output_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                _debug_file.write_text(_raw, encoding="utf-8")
                _debug_path_str = str(_debug_file)
            except Exception as _dfe:
                _debug_path_str = f"(could not write debug file: {_dfe})"
            _preview = _raw[:500] + (" … [truncated] … " + _raw[-200:] if len(_raw) > 700 else "")
            st.warning(
                f"⚠️ JSON parse failed ({_je}). "
                f"output_tokens={output_tokens}. "
                f"Full raw output saved to: {_debug_path_str}\n\n"
                f"Raw output preview:\n\n```\n{_preview}\n```"
            )
            parsed = {}
        lp_part = ""
        assess_part = ""
        lo_block_part = ""

        _final_result = {
            "grade":            grade,
            "subject":          subject,
            "chapter_number":   chapter["chapter_number"],
            "chapter_title":    chapter.get("chapter_title", ""),
            "lesson_plan":      parsed.get("lesson_plan", {}),
            "coverage_handoff": parsed.get("coverage_handoff", {}),
            "assessment_items": parsed.get("assessment_items", []) if include_assessment else [],
            "input_tokens":     input_tokens,
            "output_tokens":    output_tokens,
            "cost_inr":         calculate_cost_inr("claude-sonnet-4-6", input_tokens, output_tokens),
            "plan_status":      "full_lpa" if include_assessment else "lp_only",
        }
        if result_queue is not None:
            result_queue.put(_final_result)
        return _final_result

    except Exception as e:
        _err_result = {
            "grade":            grade,
            "subject":          subject,
            "chapter_number":   chapter["chapter_number"],
            "chapter_title":    chapter.get("chapter_title", ""),
            "lesson_plan":      {},
            "coverage_handoff": {},
            "assessment_items": [],
            "plan_status":      "full_lpa" if include_assessment else "lp_only",
            "error":            str(e),
        }
        if result_queue is not None:
            result_queue.put(_err_result)
        return _err_result


# Backwards-compat alias — old name still resolves for any external callers
# during the migration. New code should call generate_lp_only() directly.
generate_lpa = generate_lp_only


def generate_assessment_only(
    saved_plan:   dict,
    result_queue: "queue.Queue | None" = None,
    stop_event:   "threading.Event | None" = None,
) -> dict:
    """Generate assessment items for a previously saved LP-only plan.

    Inputs are sourced from the saved plan and mirror — the teacher already
    committed all decisions (grade, subject, chapter, period schedule) at
    LP generation time, so no parameter entry is required.

    Returns a result dict carrying assessment_items plus the run's token
    accounting under assess_* keys (kept separate from the LP run's tokens
    so per-stage costs remain auditable in the saved plan).
    """
    grade          = saved_plan.get("grade", "")
    subject        = saved_plan.get("subject", "")
    chapter_number = saved_plan.get("chapter_number", 0)
    chapter_title  = saved_plan.get("chapter_title", "")
    coverage_handoff = (saved_plan.get("result") or {}).get("coverage_handoff", {}) or {}

    # Edge case: pre-split saved plans may have an empty coverage_handoff.
    # The Assessment Constitution treats coverage_handoff as the sole
    # structural input — without it we cannot ground items, so bail early.
    if not coverage_handoff:
        _empty = {
            "grade":                grade,
            "subject":              subject,
            "chapter_number":       chapter_number,
            "chapter_title":        chapter_title,
            "assessment_items":     [],
            "assess_input_tokens":  0,
            "assess_output_tokens": 0,
            "assess_cost_inr":      0.0,
            "error":                ("coverage_handoff is empty — this plan was saved before "
                                     "the LP/A split. Please regenerate the lesson plan."),
        }
        if result_queue is not None:
            result_queue.put(_empty)
        return _empty

    paths        = resolve_paths(grade, subject, chapter_number)
    assess_const = read_file(paths["assessment_const"])
    summary      = read_file(paths["chapter_summary"])

    # Originality guard (Rule 2(b)) — PILOT SCOPE: English Grade IX only.
    # Strip tasks_verbatim[] / question_bank[] from the summary and
    # tasks_anchored[] from the handoff before they enter the prompt envelope,
    # so the model cannot echo textbook task structure. Other subjects/grades
    # are deliberately untouched until the pilot validates the approach.
    if subject_to_folder(subject) == "english" and grade == "Grade IX":
        summary          = redact_summary_for_assessment(summary, True)
        coverage_handoff = redact_coverage_handoff(coverage_handoff)

    # ── TWAU: build competency descriptions block from mapping JSON ───────────
    _subj_folder_da = subject_to_folder(subject)
    _comp_desc_block = ""
    if _subj_folder_da == "the_world_around_us":
        try:
            _da_mapping = json.loads(paths["chapter_mapping"].read_text(encoding="utf-8"))
            _da_comp_lookup = {
                c["c_code"]: c["competency_text"]
                for c in _da_mapping.get("competencies", [])
                if c.get("c_code") and c.get("competency_text")
            }
            if _da_comp_lookup:
                _comp_desc_block = (
                    "\n=== COMPETENCY DESCRIPTIONS ===\n"
                    + json.dumps(_da_comp_lookup, ensure_ascii=False, indent=2)
                    + "\n"
                )
        except Exception:
            _comp_desc_block = ""

    # ── Prompt structure (caching-aware) ─────────────────────────────────────
    # system: Assessment Constitution only — cacheable across all chapters
    # within the same subject.
    system_prompt_blocks = [
        {
            "type": "text",
            "text": (
                "You are Aruvi's assessment generator.\n\n"
                "You operate under the Assessment Constitution below. It is binding.\n"
                "No instruction in the user prompt overrides it.\n\n"
                f"=== ASSESSMENT CONSTITUTION ===\n{assess_const}\n"
            ),
            **_cache_ctrl(),
        }
    ]

    # static user: chapter summary — cacheable per chapter
    _static_user_text = f"=== CHAPTER SUMMARY ===\n{summary}\n"

    # variable user: coverage_handoff + competency descriptions (TWAU) + output instruction
    _variable_user_text = f"""Generate the chapter assessment using the inputs below.

=== COVERAGE HANDOFF ===
{json.dumps(coverage_handoff, ensure_ascii=False, indent=2)}
{_comp_desc_block}
=== INSTRUCTIONS ===
Follow the Assessment Constitution exactly.
The coverage_handoff above is your sole structural input. Ground every question
in it. Do not re-derive stage structure from the chapter summary.

Produce your entire output as a single valid JSON object:
{{
  "assessment_items": <per Assessment Constitution>
}}

Output only the raw JSON object. No markdown. No prose. No ```json fences.
"""

    user_message_blocks = [
        {
            "type": "text",
            "text": _static_user_text,
            **_cache_ctrl(),
        },
        {
            "type": "text",
            "text": _variable_user_text,
        },
    ]

    try:
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

        import time as _time
        progress_placeholder = st.empty()
        timer_placeholder    = st.empty()

        _gen_start_ms = int(_time.time() * 1000)

        # Reuse the same visual vocabulary as the LP run but with a 3-step list.
        _pcss_anim = (
            "<style>"
            "@keyframes aruviPulse{0%,100%{opacity:1}50%{opacity:.3}}"
            "@keyframes spin{to{transform:rotate(360deg)}}"
            "</style>"
        )
        _pcss_hide = (
            "<style>"
            "body [data-testid='stDeployButton'],"
            "body [data-testid='stAppDeployButton'],"
            "body [data-testid='stToolbar']"
            "{display:none!important;visibility:hidden!important;}"
            "</style>"
        )
        _pcss = _pcss_anim + _pcss_hide

        # Hidden 0-height driver iframe that reaches into the parent document
        # every second and updates a visible <span id="aruvi-da-timer"> that
        # lives INSIDE the popup body. Putting the time text inline (instead
        # of in a separately-floated iframe) keeps it visually tied to the
        # popup regardless of how tall the popup is.
        _timer_widget_html = (
            '<!doctype html><html><body><script>'
            '(function(){'
            f'  var _start={_gen_start_ms};'
            '  function _tick(){'
            '    try {'
            '      var el=window.parent.document.getElementById("aruvi-da-timer");'
            '      if(!el){return;}'
            '      var s=Math.floor((Date.now()-_start)/1000);'
            '      var m=Math.floor(s/60);var sc=s%60;'
            '      el.textContent=(m<10?"0"+m:m)+":"+(sc<10?"0"+sc:sc);'
            '    } catch(e) {}'
            '  }'
            '  _tick();'
            '  setInterval(_tick,1000);'
            '})();'
            '</script></body></html>'
        )
        st.markdown(
            "<style>"
            "iframe[srcdoc*='aruvi-da-timer']{"
            "  width:0!important;height:0!important;border:0!important;"
            "  position:absolute!important;left:-9999px!important;"
            "}"
            "</style>",
            unsafe_allow_html=True,
        )
        with timer_placeholder.container():
            components.html(_timer_widget_html, height=0, scrolling=False)

        _tick_icon = (
            '<div style="width:12px;height:12px;border-radius:50%;background:#d4f0e4;'
            'display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px;">'
            '<div style="width:5px;height:3px;border-left:1.5px solid #2d8a5e;'
            'border-bottom:1.5px solid #2d8a5e;transform:rotate(-45deg);'
            'margin-top:-1px;"></div></div>'
        )
        _spin_icon = (
            '<div style="width:12px;height:12px;border-radius:50%;'
            'border:1.5px solid #e8a83e;border-top-color:transparent;'
            'animation:spin 0.7s linear infinite;flex-shrink:0;'
            'margin-top:1px;box-sizing:border-box;"></div>'
        )
        _dot_icon = (
            '<div style="width:12px;height:12px;display:flex;align-items:center;'
            'justify-content:center;flex-shrink:0;margin-top:1px;">'
            '<div style="width:6px;height:6px;background:#d9d6d0;border-radius:50%;"></div></div>'
        )
        def _row_done(t):    return f'<div style="display:flex;align-items:flex-start;gap:8px;font-size:12px;color:#9c9895;">{_tick_icon}<span>{t}</span></div>'
        def _row_active(t):  return f'<div style="display:flex;align-items:flex-start;gap:8px;font-size:12px;color:#3d3b38;font-weight:500;">{_spin_icon}<span>{t}</span></div>'
        def _row_pending(t): return f'<div style="display:flex;align-items:flex-start;gap:8px;font-size:12px;opacity:0.45;">{_dot_icon}<span>{t}</span></div>'

        _steps_assess = [
            "Reading Assessment Constitution",
            "Reading chapter summary",
            "Writing assessment questions&#8230;",
        ]

        # ── Pre-compute assessment group totals for tick tracking ─────────────
        _da_subj_folder = subject_to_folder(subject)
        if _da_subj_folder == "social_sciences":
            try:
                _da_mapping_raw = paths["chapter_mapping"].read_text(encoding="utf-8")
                _da_mapping_obj = json.loads(_da_mapping_raw)
                _da_total_groups = len(_da_mapping_obj.get("competencies", []))
            except Exception:
                _da_total_groups = 0
        elif _da_subj_folder == "mathematics":
            # coverage_handoff has section_a/b/c keys — count them directly.
            # Future-proof: if a Section D is added, the count auto-adjusts.
            try:
                import re as _re_da_m
                _da_sect = _re_da_m.findall(
                    r'"(section_[a-z])"(?:\s*:)',
                    json.dumps(coverage_handoff)
                )
                _da_total_groups = len(set(_da_sect)) if _da_sect else 3
            except Exception:
                _da_total_groups = 3
        elif _da_subj_folder == "english":
            # coverage_handoff is a dict keyed by spine_code — only present
            # spines appear. len() gives the exact spine count for this chapter.
            _da_total_groups = len(coverage_handoff) if isinstance(coverage_handoff, dict) else 0
        else:  # science — coverage_handoff is a list of stage objects
            try:
                # coverage_handoff for science is a list of stage dicts,
                # one per progression stage. Its length is exactly the stage count.
                _da_total_groups = len(coverage_handoff) if isinstance(coverage_handoff, list) else 0
            except Exception:
                _da_total_groups = 0

        def _da_ticks_row(done: int, total: int) -> str:
            """Compact tick-pill row for assessment popup (deferred path)."""
            if total <= 0:
                return (
                    f'<div style="display:flex;align-items:flex-start;gap:8px;'
                    f'font-size:12px;color:#3d3b38;font-weight:500;">'
                    f'{_spin_icon}<span>Writing assessment questions&#8230;</span></div>'
                )
            _pill_done = (
                'display:inline-flex;align-items:center;justify-content:center;'
                'width:14px;height:14px;border-radius:3px;background:#d4f0e4;'
                'font-size:8px;color:#2d8a5e;font-weight:700;margin:0 1px;flex-shrink:0;'
            )
            _pill_todo = (
                'display:inline-flex;align-items:center;justify-content:center;'
                'width:14px;height:14px;border-radius:3px;background:#f2f0ec;'
                'border:1px solid #d9d6d0;margin:0 1px;flex-shrink:0;'
            )
            pills = "".join(
                f'<span style="{_pill_done}">✓</span>' if i < done
                else f'<span style="{_pill_todo}"></span>'
                for i in range(total)
            )
            fraction = f'<span style="font-size:10px;color:#5c5a56;margin-left:4px;">{done}&thinsp;/&thinsp;{total}</span>'
            spin = _spin_icon if done < total else _tick_icon
            return (
                f'<div style="display:flex;align-items:flex-start;gap:8px;'
                f'font-size:12px;color:#3d3b38;font-weight:500;">'
                f'{spin}'
                f'<div style="display:flex;flex-direction:column;gap:3px;">'
                f'<span>Writing assessment questions</span>'
                f'<div style="display:flex;align-items:center;flex-wrap:wrap;">'
                f'{pills}{fraction}'
                f'</div></div></div>'
            )

        def _da_progress(active_idx: int, ticks_row: str = "") -> str:
            parts = []
            for _i, _s in enumerate(_steps_assess):
                if _i < active_idx:
                    parts.append(_row_done(_s))
                elif _i == active_idx:
                    if _i == 2 and ticks_row:
                        parts.append(ticks_row)
                    else:
                        parts.append(_row_active(_s))
                else:
                    parts.append(_row_pending(_s))
            return (
                _pcss + _box_open + _hdr_working + _body_open
                + "".join(parts)
                + '</div>'
                + _timer_badge
                + '</div>'
            )

        _box_open = (
            '<div class="aruvi-progress-box" style="position:fixed;top:80px;right:24px;'
            'width:280px;z-index:9999;background:white;border:1px solid #d9d6d0;'
            'border-radius:10px;overflow:hidden;">'
        )
        _hdr_working = (
            '<div style="padding:8px 12px;border-bottom:1px solid #ece9e4;'
            'display:flex;gap:8px;align-items:center;">'
            '<div style="width:10px;height:10px;border-radius:50%;background:#e8a83e;'
            'animation:aruviPulse 1.4s infinite;flex-shrink:0;"></div>'
            '<span style="font-size:11px;color:#7a776f;font-weight:500;flex:1;">'
            'Generating assessment&#8230;</span>'
            '<button onclick="(function(){'
            'var w=document.querySelector(\'[class*=st-key-btn_stop_da_generation]\');'
            'if(w){var b=w.querySelector(\'button\');if(b)b.click();}'
            '})()" '
            'style="display:inline-flex;align-items:center;gap:4px;font-size:10px;'
            'color:#9c9895;background:#f2f0ec;border:1px solid #dddad5;'
            'border-radius:4px;padding:3px 7px;cursor:pointer;white-space:nowrap;'
            'font-family:inherit;line-height:1;">'
            '<span style="width:7px;height:7px;background:#9c9895;border-radius:1px;'
            'display:inline-block;flex-shrink:0;"></span>'
            'stop</button>'
            '</div>'
        )
        _body_open = (
            '<div style="padding:10px 12px 12px;display:flex;flex-direction:column;gap:2px;">'
        )
        # Visible live-time pill inside the popup. The hidden driver iframe
        # rewrites this span's textContent every second via
        # window.parent.document.getElementById("aruvi-da-timer").
        _timer_badge = (
            '<div style="display:flex;justify-content:flex-end;padding:4px 12px 8px 0;">'
            '<span id="aruvi-da-timer" style="font-family:monospace;font-size:10px;'
            'color:#5a5754;background:#f7f5f2;border-radius:4px;'
            'padding:2px 6px;">00:00</span></div>'
        )

        progress_placeholder.markdown(_da_progress(0), unsafe_allow_html=True)
        _time.sleep(3)
        progress_placeholder.markdown(_da_progress(1), unsafe_allow_html=True)
        _time.sleep(3)
        progress_placeholder.markdown(
            _da_progress(2, ticks_row=_da_ticks_row(0, _da_total_groups)),
            unsafe_allow_html=True,
        )

        streamed_text   = ""
        _stopped_by_user = False
        _da_groups_seen = 0

        with client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=32000,
            system=system_prompt_blocks,
            messages=[{"role": "user", "content": user_message_blocks}],
            extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
        ) as stream:
            for text in stream.text_stream:
                if stop_event is not None and stop_event.is_set():
                    _stopped_by_user = True
                    break
                streamed_text += text

                # ── Assessment tick tracking (deferred path) ──────────────────
                if _da_subj_folder == "science":
                    import re as _re_da
                    _dag = len(set(int(x) for x in _re_da.findall(
                        r'"progression_stage"\s*:\s*(\d+)', streamed_text
                    )))
                elif _da_subj_folder == "social_sciences":
                    import re as _re_da2
                    _dag = len(set(_re_da2.findall(
                        r'"c_code"\s*:\s*"([^"]+)"', streamed_text
                    )))
                elif _da_subj_folder == "mathematics":
                    _dag = streamed_text.count('"section_code"')
                elif _da_subj_folder == "the_world_around_us":
                    # One item per period — count completed items via implied_lo
                    # field occurrences (one per assessment item in the schema).
                    _dag = streamed_text.count('"implied_lo"')
                else:  # english
                    _dag = streamed_text.count('"spine_code"')
                _dag_capped = min(_dag, _da_total_groups) if _da_total_groups else _dag
                if _dag_capped != _da_groups_seen:
                    _da_groups_seen = _dag_capped
                    progress_placeholder.markdown(
                        _da_progress(2, ticks_row=_da_ticks_row(
                            _da_groups_seen, _da_total_groups
                        )),
                        unsafe_allow_html=True,
                    )

        if _stopped_by_user:
            try:
                _u = stream.get_current_message_snapshot().usage
                _it = getattr(_u, "input_tokens", 0)
                _ot = getattr(_u, "output_tokens", 0)
                _cw = getattr(_u, "cache_creation_input_tokens", 0)
                _cr = getattr(_u, "cache_read_input_tokens", 0)
            except Exception:
                _it = _ot = _cw = _cr = 0
            log_tokens(
                call_type          = "assessment_generation_deferred_stopped",
                grade              = grade,
                subject            = subject,
                chapter_number     = chapter_number,
                chapter_title      = chapter_title,
                input_tokens       = _it,
                output_tokens      = _ot,
                model              = "claude-sonnet-4-6",
                cache_write_tokens = _cw,
                cache_read_tokens  = _cr,
            )
            progress_placeholder.empty()
            timer_placeholder.empty()
            _stopped = {
                "grade":                grade,
                "subject":              subject,
                "chapter_number":       chapter_number,
                "chapter_title":        chapter_title,
                "assessment_items":     [],
                "assess_input_tokens":  _it,
                "assess_output_tokens": _ot,
                "assess_cost_inr":      calculate_cost_inr("claude-sonnet-4-6", _it, _ot),
                "stopped":              True,
            }
            if result_queue is not None:
                result_queue.put(_stopped)
            return _stopped

        usage = stream.get_final_message().usage
        input_tokens       = usage.input_tokens
        output_tokens      = usage.output_tokens
        cache_write_tokens = getattr(usage, "cache_creation_input_tokens", 0)
        cache_read_tokens  = getattr(usage, "cache_read_input_tokens",     0)
        log_tokens(
            call_type          = "assessment_generation_deferred",
            grade              = grade,
            subject            = subject,
            chapter_number     = chapter_number,
            chapter_title      = chapter_title,
            input_tokens       = input_tokens,
            output_tokens      = output_tokens,
            model              = "claude-sonnet-4-6",
            cache_write_tokens = cache_write_tokens,
            cache_read_tokens  = cache_read_tokens,
        )

        _raw = streamed_text.strip()
        if _raw.startswith("```"):
            _fence_end = _raw.find("```", 3)
            _raw = (_raw[_raw.index("\n") + 1 : _fence_end] if _fence_end > 3 else _raw).strip()
        _brace = _raw.find("{")
        if _brace > 0:
            _raw = _raw[_brace:]

        parsed = {}
        try:
            parsed = json.loads(_raw)
        except Exception as _je:
            try:
                _debug_dir = Path(__file__).parent.parent / "mirror" / "debug"
                _debug_dir.mkdir(parents=True, exist_ok=True)
                _debug_file = _debug_dir / f"raw_da_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                _debug_file.write_text(_raw, encoding="utf-8")
                _debug_path_str = str(_debug_file)
            except Exception as _dfe:
                _debug_path_str = f"(could not write debug file: {_dfe})"
            _preview = _raw[:500] + (" … [truncated] … " + _raw[-200:] if len(_raw) > 700 else "")
            st.warning(
                f"⚠️ Assessment JSON parse failed ({_je}). "
                f"Raw output saved to: {_debug_path_str}\n\n"
                f"Preview:\n\n```\n{_preview}\n```"
            )
            progress_placeholder.empty()
            timer_placeholder.empty()
            _err = {
                "grade":                grade,
                "subject":              subject,
                "chapter_number":       chapter_number,
                "chapter_title":        chapter_title,
                "assessment_items":     [],
                "assess_input_tokens":  input_tokens,
                "assess_output_tokens": output_tokens,
                "assess_cost_inr":      calculate_cost_inr("claude-sonnet-4-6", input_tokens, output_tokens),
                "error":                f"JSON parse failed: {_je}",
            }
            if result_queue is not None:
                result_queue.put(_err)
            return _err

        progress_placeholder.empty()
        timer_placeholder.empty()

        _final = {
            "grade":                grade,
            "subject":              subject,
            "chapter_number":       chapter_number,
            "chapter_title":        chapter_title,
            "assessment_items":     parsed.get("assessment_items", []),
            "assess_input_tokens":  input_tokens,
            "assess_output_tokens": output_tokens,
            "assess_cost_inr":      calculate_cost_inr("claude-sonnet-4-6", input_tokens, output_tokens),
        }
        if result_queue is not None:
            result_queue.put(_final)
        return _final

    except Exception as e:
        _err = {
            "grade":                grade,
            "subject":              subject,
            "chapter_number":       chapter_number,
            "chapter_title":        chapter_title,
            "assessment_items":     [],
            "assess_input_tokens":  0,
            "assess_output_tokens": 0,
            "assess_cost_inr":      0.0,
            "error":                str(e),
        }
        if result_queue is not None:
            result_queue.put(_err)
        return _err

# ── LPA normalisation helpers ─────────────────────────────────────────────────
# These bridge the old (lo_handoff flat list) and new (A3 lesson_plan.periods)
# JSON shapes so that lpa_page.html always receives the same field names.

def _normalise_lo_handoff(result: dict, comp_descs: dict) -> list:
    """
    Return per-period dicts in the shape lpa_page.html lo_handoff expects.

    New A3 format:  result["lesson_plan"]["periods"]  — nested competency{},
                    time_bands[{minutes,activity}], material as list,
                    section_anchor, visual_representation{}
    Old format:     result["lo_handoff"]              — flat per-period objects
    """
    lp = result.get("lesson_plan")
    if isinstance(lp, dict) and lp.get("periods"):
        # English: build section_id → section_type lookup from
        # main_sections_inventory if present (the field lives at result-level,
        # not on each period). Saved plans pre-inventory have no type info,
        # so the renderer simply omits the type pill.
        _eng_type_by_sec = {}
        _inv = result.get("main_sections_inventory") or []
        if isinstance(_inv, list):
            for _s in _inv:
                if isinstance(_s, dict) and _s.get("section_id"):
                    _eng_type_by_sec[_s["section_id"]] = _s.get("type", "") or _s.get("section_type", "") or ""
        out = []
        for p in lp["periods"]:
            # ── Mathematics format detection (v2.1 shape) ─────────────────────
            # Only Maths LP carries `textbook_segments` (array of §-locators) +
            # `textbook_items_in_class` (typed item pointers). These two together
            # are unique to Maths and absent from Science / Social Sciences.
            _is_maths = (
                isinstance(p.get("textbook_segments"), list)
                and (
                    "textbook_items_in_class" in p
                    or "section_goal" in p
                )
            )
            if _is_maths:
                mat = p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                # phases [{minutes (range string), description}] → time_slots
                time_slots = [
                    {"time": ph.get("minutes", ""), "desc": ph.get("description", "")}
                    for ph in (p.get("phases") or [])
                ]
                # Anchor display: §-locators joined ("§5.1" or "§5.4, §5.5").
                # Constitution shape is [{"ref": "§5.1", "title": "..."}], but
                # older saved plans may still hold plain strings — handle both.
                _segs = p.get("textbook_segments") or []
                if isinstance(_segs, list):
                    _seg_refs = [
                        (s.get("ref") or "").strip() if isinstance(s, dict)
                        else str(s).strip()
                        for s in _segs
                    ]
                    _anchor = ", ".join(r for r in _seg_refs if r)
                    # Collect section titles for display (prefer title over ref)
                    _seg_titles = [
                        (s.get("title") or "").strip() if isinstance(s, dict)
                        else ""
                        for s in _segs
                    ]
                    _section_title = ", ".join(t for t in _seg_titles if t)
                else:
                    _anchor = str(_segs)
                    _section_title = ""
                # Build a teacher-facing list of textbook items used in class,
                # rendered by book_ref (NEVER by internal id) per LP Rule 10.
                _items_inclass = p.get("textbook_items_in_class") or []
                _items_homework = p.get("homework") or []
                _ic_lines = "; ".join(
                    (it.get("book_ref") or "").strip()
                    for it in _items_inclass
                    if it.get("book_ref")
                )
                _hw_lines = "; ".join(
                    (it.get("book_ref") or "").strip()
                    for it in _items_homework
                    if it.get("book_ref")
                )
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         _anchor,
                    "activity_name":           p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    # Maths has no per-period implied LO; show pedagogical method
                    # in this slot so teachers see SOMETHING informative.
                    "implied_lo":              p.get("pedagogical_method", ""),
                    "c_code":                  "",
                    "cg":                      "",
                    "weight":                  0,
                    "competency_text":         "",
                    "visual_representation":   None,
                    # ── Maths-specific fields surfaced to lpa_page.html ─────────
                    # The HTML may safely ignore unknown keys; new renderers can
                    # use these for richer display.
                    "is_mathematics":          True,
                    "section_title":           _section_title,
                    "section_goal":            p.get("section_goal", ""),
                    "pedagogical_method":      p.get("pedagogical_method", ""),
                    "textbook_segments":       _segs,
                    "textbook_items_in_class": _items_inclass,
                    "homework":                _items_homework,
                    "items_in_class_book_refs": _ic_lines,
                    "homework_book_refs":      _hw_lines,
                    "teacher_notes":           p.get("teacher_notes", ""),
                    # NOTE: deliberately do NOT set `activity_title` or
                    # `stage_label` at the top level — lpa_page.html uses
                    # `activity_title !== undefined || stage_label !== undefined`
                    # to detect Science. Maths uses `activity_name` (SS field)
                    # so the HTML routes Maths through the SS render path,
                    # which displays activity_name + time_slots correctly.
                })
                continue
            # ── English format detection (section × spine schema) ───────────
            # Unique signals: section_id (A/B/C) + spines_taught (list).
            # Maths uses textbook_segments; Science uses stage_label;
            # SS uses competency. None of these collide with English.
            _is_english = (
                isinstance(p.get("spines_taught"), list)
                and p.get("section_id") is not None
            )
            if _is_english:
                mat = p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                time_slots = [
                    {"time": ph.get("minutes", ""), "desc": ph.get("description", "")}
                    for ph in (p.get("phases") or [])
                ]
                _sec_id    = p.get("section_id", "") or ""
                _sec_title = p.get("section_title", "") or ""
                # section_type prefers the per-period field (newer schema);
                # falls back to inventory lookup; finally empty.
                _sec_type  = (
                    p.get("section_type", "")
                    or _eng_type_by_sec.get(_sec_id, "")
                    or ""
                )
                _chapter_section = (
                    f"Section {_sec_id} · {_sec_title}"
                    if _sec_id and _sec_title else (_sec_title or _sec_id or "")
                )
                # pedagogical_methods is a dict spine→method per the
                # constitution. Tolerate the singular `pedagogical_method`
                # string used by saved plans pre-dict — broadcast it to
                # every spine in spines_taught.
                _ped = p.get("pedagogical_methods")
                if not isinstance(_ped, dict) or not _ped:
                    _single = p.get("pedagogical_method") or ""
                    _spines = p.get("spines_taught") or []
                    _ped = {s: _single for s in _spines if isinstance(s, str)} if _single else {}
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         _chapter_section,
                    "activity_name":           p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    "implied_lo":              "",
                    "c_code":                  "",
                    "cg":                      "",
                    "weight":                  0,
                    "competency_text":         "",
                    "visual_representation":   None,
                    # ── English-specific fields surfaced to lpa_page.html ────
                    # The HTML's English render branch reads these by name.
                    # Do NOT set stage_label / activity_title at top level —
                    # the HTML now dispatches on d.subject, but belt-and-braces
                    # keeps any residual data-shape detection from misrouting.
                    "is_english":              True,
                    "section_id":              _sec_id,
                    "section_title":           _sec_title,
                    "section_type":            _sec_type,
                    "spines_taught":           p.get("spines_taught") or [],
                    "pedagogical_methods":     _ped,
                    "tasks_in_class":          p.get("tasks_in_class") or [],
                    "homework":                p.get("homework") or [],
                    "teacher_notes":           p.get("teacher_notes", ""),
                })
                continue
            # ── The World Around Us (TWAU) format detection ─────────────────
            # Primary signal: dominant_mode (one of O&R|HI|D&C|C&E|R&A).
            # cg_codes is no longer emitted from v1.2 of the LP constitution
            # so detection must not depend on it.
            _TWAU_MODES = {"O&R", "HI", "D&C", "C&E", "R&A"}
            _is_twau = (
                p.get("dominant_mode") in _TWAU_MODES
                and p.get("stage_label") is None
                and p.get("progression_stage") is None
            )
            if _is_twau:
                mat = p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                # time_bands [{minutes, activity}] → time_slots [{time, desc}]
                time_slots = [
                    {"time": tb.get("minutes", ""), "desc": tb.get("activity", "")}
                    for tb in (p.get("time_bands") or [])
                ]
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         p.get("section_ref", "") or p.get("textbook_anchor", ""),
                    "activity_name":           p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    "implied_lo":              p.get("implied_lo", ""),
                    "c_code":                  "",
                    "cg":                      "",
                    "weight":                  0,
                    "competency_text":         "",
                    "visual_representation":   None,
                    # ── TWAU-specific fields surfaced to lpa_page.html ──────────
                    "is_twau":                 True,
                    "cg_codes":                p.get("cg_codes") or [],
                    "dominant_mode":           p.get("dominant_mode", ""),
                    "textbook_anchor":         p.get("textbook_anchor", ""),
                    "teacher_notes":           p.get("teacher_facilitation_note", ""),
                })
                continue
            # ── Mathematics v2 format detection (Grade IV / preparatory-stage schema)
            # Signal: section_refs (list) + section_titles (list) + tasks_in_class +
            # phases; no textbook_segments, no stage_label, no spines_taught,
            # no competency dict, no dominant_mode. activity_title IS present here.
            _is_maths_v2 = (
                isinstance(p.get("section_refs"), list)
                and isinstance(p.get("section_titles"), list)
                and isinstance(p.get("tasks_in_class"), list)
                and p.get("stage_label") is None
                and p.get("progression_stage") is None
                and not isinstance(p.get("spines_taught"), list)
                and p.get("dominant_mode") not in {"O&R", "HI", "D&C", "C&E", "R&A"}
                and not (p.get("competency") or {})
            )
            if _is_maths_v2:
                mat = p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                # phases [{minutes, description}] → time_slots
                time_slots = [
                    {"time": ph.get("minutes", ""), "desc": ph.get("description", "")}
                    for ph in (p.get("phases") or [])
                ]
                # Section display: join section_titles (fall back to section_refs)
                _sec_titles = [t for t in (p.get("section_titles") or []) if t]
                _sec_refs   = [r for r in (p.get("section_refs")   or []) if r]
                _section_title  = ", ".join(_sec_titles) if _sec_titles else ", ".join(_sec_refs)
                _anchor         = ", ".join(_sec_refs)
                # Homework items from tasks_in_class flagged as homework, or a
                # separate homework list if present
                _items_homework = p.get("homework") or []
                _items_inclass  = p.get("tasks_in_class") or []
                _hw_lines = "; ".join(
                    (it.get("book_ref") or "").strip()
                    for it in _items_homework if it.get("book_ref")
                )
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         _anchor,
                    "activity_name":           p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    "implied_lo":              p.get("pedagogical_method", ""),
                    "c_code":                  "",
                    "cg":                      "",
                    "weight":                  0,
                    "competency_text":         "",
                    "visual_representation":   None,
                    # Maths-specific fields
                    "is_mathematics":          True,
                    "section_title":           _section_title,
                    "section_goal":            p.get("section_goal", ""),
                    "pedagogical_method":      p.get("pedagogical_method", ""),
                    "textbook_segments":       [{"ref": r, "title": t} for r, t in zip(_sec_refs, _sec_titles)],
                    "textbook_items_in_class": _items_inclass,
                    "homework":                _items_homework,
                    "items_in_class_book_refs": "; ".join(
                        (it.get("book_ref") or "").strip()
                        for it in _items_inclass if it.get("book_ref")
                    ),
                    "homework_book_refs":      _hw_lines,
                    "teacher_notes":           p.get("teacher_notes", ""),
                })
                continue

            # ── Science format detection ────────────────────────────────────
            # Only use truly Science-specific fields (stage_label / progression_stage).
            # activity_title is NOT a reliable Science signal — Social Sciences plans
            # may also use that field name (e.g. ch_04 generated with activity_title
            # instead of activity_name), which would wrongly set c_code="" and break
            # competency-based collapsible grouping for Social Sciences.
            if p.get("stage_label") is not None or p.get("progression_stage") is not None:
                mat = p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                time_slots = [
                    {"time": ph.get("minutes", ""), "desc": ph.get("description", "")}
                    for ph in (p.get("phases") or [])
                ]
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         p.get("stage_label", ""),
                    "activity_name":           p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    "implied_lo":              (p.get("activity_description") or "")[:200],
                    "c_code":                  "",
                    "cg":                      "",
                    "weight":                  p.get("progression_stage", 1),
                    "competency_text":         p.get("pedagogical_approach", ""),
                    "visual_representation":   None,
                    # ── Science-detection fields for lpa_page.html ───────────────
                    # The HTML checks periods[0].stage_label / activity_title to
                    # detect Science; these must be present as top-level keys.
                    "stage_label":             p.get("stage_label", ""),
                    "activity_title":          p.get("activity_title", ""),
                    "progression_stage":       p.get("progression_stage", 0),
                    "description":             p.get("description", ""),
                    # ── Lesson view panel fields for lpa_page.html ───────────────
                    "activity_description":    p.get("activity_description", ""),
                    "actors":                  p.get("roles", []),
                    "pedagogical_approach":    p.get("pedagogical_approach", ""),
                })
            # ── Secondary-stage Science (section-anchored, flat) ─────────────
            # Distinguisher vs Social Sciences A3: secondary Science periods
            # carry `section_context` (and, once regenerated, a per-period
            # `pedagogical_approach`); Social Sciences periods carry neither.
            # Shape: section_anchor + time_bands + competency{C-code}, NO stage
            # fields. Rendered FLAT (no stage layer), approach column, LO at end,
            # NO materials row.
            _is_science_secondary = (
                p.get("section_context") is not None
                and p.get("section_anchor") is not None
                and isinstance(p.get("time_bands"), list)
                and p.get("stage_label") is None
                and p.get("progression_stage") is None
                and not isinstance(p.get("spines_taught"), list)
                and not isinstance(p.get("textbook_segments"), list)
                and p.get("dominant_mode") not in {"O&R", "HI", "D&C", "C&E", "R&A"}
            )
            if _is_science_secondary:
                comp = p.get("competency") or {}
                time_slots = [
                    {"time": tb.get("minutes", ""), "desc": tb.get("activity", "")}
                    for tb in (p.get("time_bands") or [])
                ]
                c_code = comp.get("c_code", "")
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         p.get("section_anchor", ""),
                    "activity_name":           p.get("activity_title") or p.get("activity_name", ""),
                    "activity_summary":        p.get("activity_title") or p.get("activity_name", ""),
                    "time_slots":              time_slots,
                    # No materials row for secondary Science — deliberately blank.
                    "material":                "",
                    "implied_lo":              p.get("implied_lo", ""),
                    "c_code":                  c_code,
                    "cg":                      comp.get("cg", ""),
                    "weight":                  0,
                    "competency_text":         comp_descs.get(c_code, "") or comp.get("competency_text", ""),
                    "visual_representation":   None,
                    # ── Secondary-Science render fields for lpa_page.html ────────
                    "science_secondary":       True,
                    "pedagogical_approach":    p.get("pedagogical_approach", ""),
                    "visual_aids":             p.get("visual_aids", ""),
                })
                continue
            else:
                # ── Social Sciences A3 format ────────────────────────────────
                comp = p.get("competency") or {}
                # time_bands [{minutes, activity}] → time_slots [{time, desc}]
                time_slots = [
                    {"time": tb.get("minutes", ""), "desc": tb.get("activity", "")}
                    for tb in (p.get("time_bands") or [])
                ]
                # material / materials (plural alias used by some generated plans)
                # → comma-joined string
                mat = p.get("material") if p.get("material") is not None else p.get("materials", "")
                if isinstance(mat, list):
                    mat = ", ".join(mat)
                c_code = comp.get("c_code", "")
                out.append({
                    "period_number":           p.get("period_number"),
                    "period_duration_minutes": p.get("period_duration_minutes"),
                    "chapter_section":         p.get("section_anchor", ""),
                    # activity_name is the canonical SS field; activity_title is an
                    # accepted alias used by some generated plans (e.g. ch_04).
                    "activity_name":           p.get("activity_name") or p.get("activity_title", ""),
                    "activity_summary":        p.get("activity_name") or p.get("activity_title", ""),
                    "time_slots":              time_slots,
                    "material":                mat,
                    "implied_lo":              p.get("implied_lo", ""),
                    "c_code":                  c_code,
                    "cg":                      comp.get("cg", ""),
                    "weight":                  comp.get("weight", 1),
                    "competency_text":         comp_descs.get(c_code, "") or comp.get("competency_text", ""),
                    "visual_representation":   p.get("visual_representation"),
                })
        return out

    # Old flat lo_handoff — enrich competency_text from comp_descs
    lo_list = result.get("lo_handoff", [])
    for lo in lo_list:
        if not lo.get("competency_text"):
            lo["competency_text"] = comp_descs.get(lo.get("c_code", ""), "")
    return lo_list


def _normalise_assessment_sections(result: dict, comp_descs: dict = None) -> list:
    """
    Return assessment_sections[] in the shape lpa_page.html renderAssessment() expects.

    New format:  result["assessment_items"] — flat list of question objects.
                 Each item must have: c_code, question_type, question_text,
                 options[], annotation, period_ref, weight_label,
                 competency_text, chapter_section.
    Legacy:      result["assessment_sections"] already populated — return as-is.
    Falls back to [] when neither is present.
    """
    if result.get("assessment_sections"):
        return result["assessment_sections"]

    items = result.get("assessment_items", [])
    if not items:
        return []

    # ── Mathematics format detection (v2.1 shape) ──────────────────────────
    # Maths assessment ships as a list of section-objects, each with its own
    # nested `items[]` array — distinct from the flat per-item list used by
    # Science and Social Sciences. Detect on the presence of `section_code`
    # ("A" / "B" / "C") at the top level of the first element with a nested
    # `items` array.
    _is_maths_assessment = (
        isinstance(items, list)
        and len(items) > 0
        and isinstance(items[0], dict)
        and "section_code" in items[0]
        and isinstance(items[0].get("items"), list)
    )
    if _is_maths_assessment:
        # Middle-stage fallback descriptions (A/B/C fixed).
        # Prep-stage section_title comes directly from the JSON (Explore/Reason/
        # Practise/Solve) so this dict is only consulted when note is absent AND
        # section_title is missing — treat unknown codes as title-cased intent.
        _MATHS_SECTION_DESC = {
            "A": "Recall and Apply — short answers, definitions, and procedural fluency.",
            "B": "Reason and Explain — proofs, justifications, and constructions.",
            "C": "Apply in Context — case-based and multi-concept problems.",
        }
        _maths_sections = []
        for sec in items:
            if not isinstance(sec, dict):
                continue
            _code  = sec.get("section_code", "")
            _title = sec.get("section_title", "")
            _note  = sec.get("note", "")
            _qs    = []
            for it in (sec.get("items") or []):
                if not isinstance(it, dict):
                    continue
                _qtype = it.get("question_type", "")
                _prompt = it.get("prompt", "")
                # ── Strip pipe-table from prompt when visual_stimulus already
                # carries the same table (avoids double-printing in HTML + PDF).
                _vs_raw = it.get("visual_stimulus", "") or ""
                if _vs_raw.strip() and "|" in _vs_raw:
                    _vs_lines = set(ln.strip() for ln in _vs_raw.strip().splitlines() if ln.strip())
                    _prompt_lines = _prompt.splitlines()
                    _cleaned_lines = []
                    for _ln in _prompt_lines:
                        if _ln.strip() in _vs_lines:
                            continue  # this line is the table — skip it
                        _cleaned_lines.append(_ln)
                    # Collapse multiple consecutive blank lines and trim trailing
                    import re as _re
                    _prompt = _re.sub(r'\n{3,}', '\n\n', "\n".join(_cleaned_lines)).rstrip()
                # Exercise companion (Constitution v3.2 Rule 9) — pointer to
                # textbook item that anchors this goal. Both fields empty when
                # the LP gamut walk found no anchor.
                _ex            = it.get("exercise") or {}
                _ex_book_ref   = _ex.get("book_ref", "") or ""
                _ex_description = _ex.get("description", "") or ""
                # ── Parse structured teacher_guide (v3.2) ────────────────────
                # teacher_guide is a JSON object:
                #   { expected_answer, method_one_line,
                #     what_each_option_reveals, inclusivity }
                # Tolerate legacy piped string for any saved plans pre-v3.2.
                _tg = it.get("teacher_guide", {}) or {}
                if isinstance(_tg, str):
                    _parts = [p.strip() for p in _tg.split(" | ")]
                    _tg_legacy_expected = _parts[1] if len(_parts) > 1 else ""
                    if _tg_legacy_expected.lower().startswith("expected answer:"):
                        _tg_legacy_expected = _tg_legacy_expected[len("expected answer:"):].strip()
                    _tg = {
                        "expected_answer":          _tg_legacy_expected,
                        "method_one_line":          "",
                        "what_each_option_reveals": {},
                        "inclusivity":              _parts[2] if len(_parts) > 2 else "",
                    }
                _tg_expected      = _tg.get("expected_answer", "") or ""
                _tg_method        = _tg.get("method_one_line", "") or ""
                _tg_what_reveals  = _tg.get("what_each_option_reveals", {}) or {}
                _tg_inclusivity   = _tg.get("inclusivity", "") or ""
                _maths_guide = {
                    "expected_answer":          _tg_expected,
                    "method_one_line":          _tg_method,
                    "what_each_option_reveals": _tg_what_reveals,
                    "inclusivity":              _tg_inclusivity,
                }
                _qs.append({
                    "type":               _qtype,
                    "question":           _prompt,
                    # OPEN_TASK is not a Mathematics question type per Constitution
                    # v3.2 Rule 10; these fields stay empty for math items.
                    "task":               "",
                    "scaffold":           "",
                    "format_of_output":   [],
                    "task_instructions":  "",
                    "options":            it.get("options", []) or [],
                    "annotation":         _ex_book_ref,
                    "period_ref":         _ex_book_ref,
                    "title":              (
                        (_qtype + ": " + (_prompt[:56] + "…" if len(_prompt) > 56 else _prompt))
                        if _prompt else _qtype
                    ),
                    "expected": (
                        next(
                            (o.get("text", "") for o in (it.get("options") or [])
                             if isinstance(o, dict) and o.get("is_correct")),
                            ""
                        )
                        if _qtype == "MCQ" else _tg_expected
                    ),
                    "cognitive_demand":         "",
                    "guide":                    _maths_guide,
                    "what_each_option_reveals": _tg_what_reveals,
                    "inclusivity":              _tg_inclusivity,
                    "visual_stimulus":          it.get("visual_stimulus", None),
                    "correct_answer":           "",
                    "implied_lo":               "",
                    # ── Maths-specific question fields surfaced to renderer ──
                    "is_mathematics":           True,
                    "section_ref":              it.get("section_ref", ""),
                    "section_title":            it.get("section_title", ""),
                    "goal":                     it.get("goal", ""),
                    "expected_answer":          _tg_expected,
                    # Exercise companion (Constitution v3.2 Rule 9)
                    "exercise":                 {
                        "book_ref":    _ex_book_ref,
                        "description": _ex_description,
                    },
                })
            _types_in_order = []
            for q in _qs:
                t = q["type"]
                if t and t not in _types_in_order:
                    _types_in_order.append(t)
            _maths_sections.append({
                # Maps onto SS-shape fields the HTML already knows how to
                # render. The HTML's SS branch renders these as:
                #   c_code (badge) | weight_label (right-side label)
                #   competency_short (description below)
                "c_code":           ("Section " + _code) if _code else "",
                "weight_label":     _title,
                "competency_short": _note or _title or _MATHS_SECTION_DESC.get(_code, ""),
                "drawing_on":       _title,
                "question_types":   " · ".join(_types_in_order),
                "questions":        _qs,
                "is_science":       False,
                "is_mathematics":   True,
                "section_code":     _code,
                "section_title":    _title,
                "stage_label":      None,
            })
        return _maths_sections

    # ── English format detection (spine-grouped schema) ─────────────────────
    # English assessment ships as a list of spine-objects, each with its own
    # nested `items[]` array. Spines: reading_for_comprehension, listening,
    # speaking, writing, vocabulary_grammar, beyond_text. Detect on
    # `spine_code` at the top level of the first element with a nested
    # `items[]` array.
    _is_english_assessment = (
        isinstance(items, list)
        and len(items) > 0
        and isinstance(items[0], dict)
        and "spine_code" in items[0]
        and isinstance(items[0].get("items"), list)
    )
    if _is_english_assessment:
        _ENGLISH_SPINE_DESC = {
            # Middle-stage spines
            "reading_for_comprehension":    "Encountering text and demonstrating comprehension — recall, inference, reflection.",
            "listening":                    "Active listening — meaning, attitude, summarisation.",
            "speaking":                     "Structured talk — conversation, discussion, debate.",
            "writing":                      "Drafting and editing — formal and creative composition.",
            "vocabulary_grammar":           "Word-building and grammar embedded in context.",
            "beyond_text":                  "Library work, projects, and interdisciplinary extensions.",
            # Preparatory-stage spines (listening folded into oracy)
            "reading":                      "Encountering text and demonstrating comprehension at the preparatory stage.",
            "oracy":                        "Merged listening and speaking — recitation, conversation, sound discrimination.",
            "word_work": "Phonics, sight words, vocabulary, word games, and grammar-in-context.",
        }
        # Closed types render `teacher_guide.suggested_answer`; open types
        # render `teacher_guide.expected_elements` as bullets.
        _ENG_CLOSED_TYPES = {"MCQ", "FILL_IN", "MATCH", "TRUE_FALSE"}

        def _eng_resolve_answer(qtype: str, tg: dict, options: list = None) -> tuple[str, list]:
            qtype_u = (qtype or "").strip().upper()
            tg = tg if isinstance(tg, dict) else {}
            # For MCQ: the correct option is already highlighted visually.
            # teacher_guide.note holds distractor analysis; render that instead.
            if qtype_u == "MCQ":
                sug = tg.get("note", "") or ""
            else:
                sug = tg.get("suggested_answer", "") or ""
            exp = tg.get("expected_elements") or []
            if not isinstance(exp, list):
                exp = []
            is_closed = qtype_u in _ENG_CLOSED_TYPES
            expected = sug if is_closed else "\n".join(str(e) for e in exp)
            return expected, exp

        _eng_sections = []
        for sec in items:
            if not isinstance(sec, dict):
                continue
            _spine_code  = (sec.get("spine_code") or "").strip().lower()
            _spine_title = sec.get("spine_title") or _spine_code.replace("_", " ").title()
            _qs = []
            for it in (sec.get("items") or []):
                if not isinstance(it, dict):
                    continue
                # New shape carries task_prompt + sub_items[]; legacy carries
                # prompt + question_type + teacher_guide. Read task_prompt
                # first, fall back to legacy prompt.
                _task_prompt = it.get("task_prompt") or it.get("item_stem") or it.get("prompt") or ""
                _outer_qtype = (it.get("question_type") or "").strip().upper()
                _outer_tg    = it.get("teacher_guide") or {}
                _sub_items_raw = it.get("sub_items")
                if not isinstance(_sub_items_raw, list):
                    _sub_items_raw = []

                # Build a normalised sub_items list for the renderer. When the
                # composite has no sub_items (open task with no textbook
                # sub-items, or a generated item), we synthesise ONE pseudo
                # sub-item from the outer task_prompt + outer teacher_guide so
                # the downstream HTML can always iterate sub_items.
                _sub_items_render = []
                if _sub_items_raw:
                    for si in _sub_items_raw:
                        if not isinstance(si, dict):
                            continue
                        _si_qtype = (si.get("question_type") or "").strip().upper()
                        _si_tg    = si.get("teacher_guide") or {}
                        _si_opts = si.get("options") or []
                        _si_expected, _si_exp_elems = _eng_resolve_answer(_si_qtype, _si_tg, _si_opts)
                        _sub_items_render.append({
                            "stem":              si.get("stem", "") or "",
                            "type":              _si_qtype,
                            "options":           _si_opts,
                            "visual_stimulus":   si.get("visual_stimulus", "") or "",
                            "expected":          _si_expected,
                            "expected_elements": _si_exp_elems,
                            "suggested_answer":  (_si_tg or {}).get("suggested_answer", "") or "",
                            "verified":          bool(si.get("verified", False)),
                        })
                else:
                    # Outer task itself owns the answer layer.
                    _outer_expected, _outer_exp_elems = _eng_resolve_answer(_outer_qtype, _outer_tg, it.get("options") or [])
                    _sub_items_render.append({
                        "stem":              _task_prompt,
                        "type":              _outer_qtype,
                        "options":           it.get("options") or [],
                        "visual_stimulus":   it.get("visual_stimulus", "") or "",
                        "expected":          _outer_expected,
                        "expected_elements": _outer_exp_elems,
                        "suggested_answer":  (_outer_tg if isinstance(_outer_tg, dict) else {}).get("suggested_answer", "") or "",
                        "verified":          bool(it.get("verified", False)),
                    })

                # Card-level type chip: when the composite has explicit
                # sub-items, leave the outer type empty (the card shows the
                # task framing only); otherwise reflect the outer task's type.
                _card_type = "" if _sub_items_raw else _outer_qtype
                _card_title = (
                    (_card_type + ": " + (_task_prompt[:56] + "…" if len(_task_prompt) > 56 else _task_prompt))
                    if (_card_type and _task_prompt) else (_card_type or "Task")
                )
                # Card-level expected: when the outer task owns the answer
                # layer (no sub_items), surface it; otherwise leave empty —
                # the per-sub-item rows carry their own answers.
                # Exception: EXTRACT_ANALYSIS always builds a synthetic sub_item
                # that carries expected_elements — never duplicate them on the card.
                _card_expected = ""
                _card_exp_elems = []
                _card_suggested = ""
                if not _sub_items_raw and _outer_qtype != "EXTRACT_ANALYSIS":
                    _card_expected, _card_exp_elems = _eng_resolve_answer(_outer_qtype, _outer_tg, it.get("options") or [])
                    _card_suggested = (_outer_tg if isinstance(_outer_tg, dict) else {}).get("suggested_answer", "") or ""

                _qs.append({
                    "type":               _card_type,
                    "question":           _task_prompt,
                    "task":               "",
                    "scaffold":           "",
                    "format_of_output":   [],
                    "task_instructions":  "",
                    "options":            it.get("options") or [],
                    "annotation":         "",
                    "period_ref":         "",
                    "title":              _card_title,
                    "expected":           _card_expected,
                    "cognitive_demand":   "",
                    "guide":              {},
                    "expected_elements":  _card_exp_elems,
                    "look_for":           [],
                    "what_each_option_reveals": {},
                    "inclusivity":        "",
                    "visual_stimulus":    it.get("visual_stimulus", "") or "",
                    "correct_answer":     "",
                    "implied_lo":         it.get("source_lo", "") or it.get("implied_lo", "") or "",
                    # ── English-specific question fields surfaced to renderer ──
                    "is_english":           True,
                    "task_prompt":          _task_prompt,
                    "sub_items":            _sub_items_render,
                    "has_sub_items":        bool(_sub_items_raw),
                    "suggested_answer":     _card_suggested,
                    "source_section_id":    it.get("source_section_id", "") or "",
                    "source_section_title": it.get("source_section_title", "") or "",
                    "source_section_type":  it.get("source_section_type", "") or "",
                    "source_spine_section": it.get("source_spine_section", "") or "",
                    "source":               it.get("source", "") or "",
                    "source_task_index":    it.get("source_task_index", -1),
                    "transcript_ref":       it.get("transcript_ref", "") or "",
                    "verified":             bool(it.get("verified", False)),
                })
            _types_in_order = []
            for q in _qs:
                # When the composite carries sub_items, the spine-level type
                # chip strip should reflect the sub-items' types; otherwise
                # use the outer card type.
                _src_types = (
                    [si.get("type") for si in (q.get("sub_items") or [])]
                    if q.get("has_sub_items")
                    else [q.get("type")]
                )
                for t in _src_types:
                    if t and t not in _types_in_order:
                        _types_in_order.append(t)
            _eng_sections.append({
                "c_code":           _spine_title,
                "weight_label":     "",
                "competency_short": sec.get("note") or _ENGLISH_SPINE_DESC.get(_spine_code, ""),
                "drawing_on":       "",
                "question_types":   " · ".join(_types_in_order),
                "questions":        _qs,
                "is_science":       False,
                "is_mathematics":   False,
                "is_english":       True,
                "spine_code":       _spine_code,
                "stage_label":      None,
            })
        return _eng_sections

    # ── Fix 1 helper: short title ≤ 60 chars from type + first words of text ──
    def _build_title(qtype: str, qtext: str) -> str:
        prefix = (qtype or "Q").strip()
        budget = 58 - len(prefix)          # leaves 2 chars for ": "
        if budget <= 0:
            return prefix[:60]
        snippet = (qtext or "").strip()
        if len(snippet) > budget:
            snippet = snippet[:budget].rsplit(" ", 1)[0]
        return (prefix + ": " + snippet) if snippet else prefix

    # ── Fix 2 helper: derive expected-answer text by question type ─────────────
    def _build_expected(item: dict) -> str:
        qtype = (item.get("question_type") or "").strip().upper()
        if qtype == "MCQ":
            opts = item.get("options") or []
            if isinstance(opts, dict):
                # Science format: {"A": "text", ...} + separate "correct_answer" key
                correct_key = item.get("correct_answer", "")
                text = opts.get(correct_key, "")
                return (correct_key + ": " + text).strip(": ") if correct_key else ""
            for opt in opts:
                if not isinstance(opt, dict):
                    continue
                if opt.get("is_correct"):
                    label = opt.get("label", opt.get("key", ""))
                    text  = opt.get("text",  opt.get("value", ""))
                    return (label + ": " + text).strip(": ") if label else text
            return ""
        if qtype == "SCR":
            elems = item.get("expected_elements") or []
            return "\n".join(str(e) for e in elems)
        if qtype == "ECR":
            elems = item.get("look_for") or []
            return "\n".join(str(e) for e in elems)
        if qtype == "OPEN_TASK":
            # New schema: format_of_output is a list; join for display
            fof = item.get("format_of_output") or []
            if isinstance(fof, list):
                return "\n".join(str(f) for f in fof)
            return str(fof)
        return ""

    # Weight integer → label, mirroring WEIGHT_LABEL in lpa_page.html
    _WLBL = {3: "Central", 2: "Substantive", 1: "Present"}

    from collections import OrderedDict
    sections: dict = OrderedDict()
    for item in items:
        # ── Science format detection ────────────────────────────────────────
        # Science items carry at least one of these fields; SS items never do.
        _is_science = (
            item.get("stage_label") is not None or
            item.get("implied_lo_assessed") is not None or
            bool(item.get("marking_guidance")) or          # Science-only field
            bool(item.get("what_each_option_reveals")) or  # Science MCQ top-level
            bool(item.get("correct_answer"))               # Science MCQ correct key
        )

        if _is_science:
            _comp      = {}
            c_code     = ""
            _group_key = item.get("stage_label") or item.get("implied_lo_assessed") or f"_sci_{len(sections)}"
        else:
            # c_code may be a top-level field OR nested under item["competency"]["c_code"],
            # exactly as in lesson-plan periods (see _normalise_lo_handoff).
            _comp = item.get("competency") or {}
            if not isinstance(_comp, dict):
                _comp = {}
            c_code     = item.get("c_code") or _comp.get("c_code", "")
            _group_key = c_code

        if _group_key not in sections:
            if _is_science:
                _wlabel     = item.get("stage_label", "")
                _ctext      = item.get("implied_lo_assessed", "")
                _drawing_on = item.get("stage_label", "")
            else:
                # weight_label: prefer explicit string; fall back to integer from competency
                _wlabel = item.get("weight_label") or ""
                if not _wlabel:
                    _w = _comp.get("weight")
                    try:
                        _wlabel = _WLBL.get(int(_w), "") if _w is not None else ""
                    except (TypeError, ValueError):
                        _wlabel = ""

                # competency_text: prefer canonical lookup from comp_descs (authoritative
                # framework descriptions); fall back to AI-generated text only if the
                # lookup misses (e.g. comp_descs not loaded).
                _ctext = (
                    (comp_descs.get(c_code, "") if comp_descs and c_code else "") or
                    item.get("competency_text") or
                    _comp.get("competency_text", "") or
                    _comp.get("text", "")
                )
                _drawing_on = item.get("chapter_section", "")

            sections[_group_key] = {
                "c_code":           c_code,
                "weight_label":     _wlabel,
                "competency_short": _ctext,
                "drawing_on":       _drawing_on,
                "question_types":   "",
                "questions":        [],
                # ── Science-detection fields for lpa_page.html renderAssessment() ──
                # is_science is the canonical flag; stage_label carried for display.
                "is_science":  _is_science,
                "stage_label": item.get("stage_label", "") if _is_science else None,
            }
        qtype = item.get("question_type", "")
        sections[_group_key]["questions"].append({
            "type":               qtype,
            "question":           item.get("question_text", ""),
            "task":               item.get("task", ""),
            "scaffold":           item.get("scaffold", ""),
            "format_of_output":   item.get("format_of_output", []),
            "task_instructions":  item.get("task_instructions", ""),
            "options":            item.get("options", []),
            "annotation":         item.get("marking_guidance", "") if _is_science else item.get("annotation", ""),
            "period_ref":         item.get("period_ref", ""),
            "title":              _build_title(qtype, item.get("task", "") or item.get("question_text", "")),
            "expected":           _build_expected(item),
            "cognitive_demand":   item.get("cognitive_demand", ""),
            # TWAU: OPEN_TASK behavioural subtype → renderers show a "Performance Task" label.
            "performance_task":   bool(item.get("performance_task", False)),
            "guide":                    item.get("guide", {}),
            "expected_elements":        item.get("expected_elements", []),
            "look_for":                 item.get("look_for", []),
            # Science-specific fields for HTML rendering.
            # Science MCQ stores distractor notes at item["guide"]["MCQ"][...];
            # try top-level first (flat schema), fall back to the nested path.
            "what_each_option_reveals": (
                item.get("what_each_option_reveals")
                or (item.get("guide") or {}).get(qtype.upper() if qtype else "MCQ", {}).get("what_each_option_reveals", {})
                or {}
            ),
            "inclusivity": (
                item.get("inclusivity")
                or (item.get("guide") or {}).get(qtype.upper() if qtype else "MCQ", {}).get("inclusivity", "")
                or ""
            ),
            "visual_stimulus":          item.get("visual_stimulus", None),
            "correct_answer":           item.get("correct_answer", ""),
            # Learning Outcome for Assessment Question column display.
            # Science: sourced from implied_lo_assessed on the item itself.
            # Social Science: sourced from implied_lo on the item itself (not competency_text).
            "implied_lo": (
                item.get("implied_lo_assessed", "")
                if _is_science else
                item.get("implied_lo", "")
            ),
        })

    # Fix 3: populate question_types — unique types in order of first appearance
    for sec in sections.values():
        seen: list = []
        for q in sec["questions"]:
            t = q["type"]
            if t and t not in seen:
                seen.append(t)
        sec["question_types"] = " · ".join(seen)

    return list(sections.values())


# ── LRM allocation helpers ────────────────────────────────────────────────────

def _ch_w3_codes(ch: dict) -> list:
    return [item["c_code"] for item in ch.get("primary", []) if item.get("weight") == 3]

def _ch_w2_codes(ch: dict) -> list:
    return [item["c_code"] for item in ch.get("primary", []) if item.get("weight") == 2]

def _ch_w1_codes(ch: dict) -> list:
    return [item["c_code"] for item in ch.get("primary", []) if item.get("weight") == 1]

def _alloc_chapter_weight(ch: dict, subject: str = "") -> int:
    """Return the allocation weight for a chapter.

    For TWAU, Science, and Mathematics the effort_index is the weight.
    For Social Sciences and other subjects the stored chapter_weight
    (competency load) is used.
    """
    _effort_subjects = {"Science", "Mathematics", "The World Around Us"}
    if subject in _effort_subjects:
        effort = ch.get("effort_index")
        if isinstance(effort, (int, float)) and effort > 0:
            return int(round(effort))
        return 0
    # Social Sciences / other: competency-load weight
    stored = ch.get("chapter_weight")
    if isinstance(stored, (int, float)) and stored > 0:
        return int(stored)
    return sum(item.get("weight", 0) for item in ch.get("primary", []))

def _lrm(raw_floats: list, total: int) -> list:
    """Largest Remainder Method: distribute `total` integer slots proportionally."""
    floors = [math.floor(f) for f in raw_floats]
    remainders = sorted(enumerate(raw_floats), key=lambda x: -(x[1] - math.floor(x[1])))
    deficit = total - sum(floors)
    result = floors[:]
    for k in range(deficit):
        result[remainders[k][0]] += 1
    return result

def _compute_allocation(chs: list, period_types: list, subject: str = "") -> list:
    """
    Returns one allocation dict per chapter.
    Each dict has {mins: count, ..., 'total': int}.
    """
    if not chs or not period_types:
        return []
    if len(chs) == 1:
        alloc = {pt["mins"]: pt["count"] for pt in period_types}
        alloc["total"] = sum(pt["count"] for pt in period_types)
        return [alloc]

    weights  = [_alloc_chapter_weight(ch, subject) for ch in chs]
    sum_w    = sum(weights) or 1
    sorted_types  = sorted(period_types, key=lambda pt: -pt["mins"])
    total_periods = sum(pt["count"] for pt in sorted_types)

    pass1     = _lrm([w / sum_w * total_periods for w in weights], total_periods)
    remaining = pass1[:]
    allocs    = [{} for _ in chs]

    for pt in sorted_types[:-1]:
        raw    = [min(w / sum_w * pt["count"], remaining[i]) for i, w in enumerate(weights)]
        result = _lrm(raw, pt["count"])
        for i, v in enumerate(result):
            allocs[i][pt["mins"]]  = v
            remaining[i]          -= v

    shortest = sorted_types[-1]
    for i, v in enumerate(remaining):
        allocs[i][shortest["mins"]] = max(0, v)

    for i in range(len(chs)):
        allocs[i]["total"] = pass1[i]

    return allocs


def _generate_pdf_bytes_alloc(
    chs: list,
    allocs: list,
    sorted_pts: list,
    grade: str,
    subject: str,
) -> bytes:
    """PDF export for the period allocation report (landscape)."""

    class _AllocPDF(FPDF):
        """Allocation report PDF with a consistent top margin on every page.

        The title block is drawn once in the body (page 1 only). Without an
        explicit header, FPDF's auto page break starts continuation pages
        (page 2+) flush against the top edge with no breathing room. This
        header() reserves the same top margin on every page so the table
        never sits tight against the top edge — applies to all subjects and
        stages, since the layout below is shared across them.
        """

        def header(self):
            self.set_y(self.t_margin)

    pdf = _AllocPDF(orientation="L", unit="mm", format="A4")
    pdf.set_top_margin(16)
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(44, 62, 80)
    pdf.cell(0, 9, "Aruvi - Period Allocation Report", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 6, f"Grade: {grade}   Subject: {subject}", ln=True)
    pdf.ln(3)

    # Detect subject group for column layout and footnote wording
    is_twau      = subject == "The World Around Us"
    is_math_prep = subject == "Mathematics" and get_stage(grade) == "preparatory"
    is_science   = subject in ("Science", "Mathematics") or is_twau
    is_english   = subject == "English"
    uses_effort_index = is_science or is_english

    # Column layout — switches on subject group
    # Mirrors the HTML: # | Chapter | Total | [pt-type cols...] | Effort Idx or Weight
    pt_headers = [f"{pt['mins']}-min Periods" for pt in sorted_pts]
    if uses_effort_index:
        # Science / Mathematics / English / TWAU: Total 3rd, EI last
        all_headers = ["#", "Chapter", "Total Periods"] + pt_headers + ["Effort Index"]
        fixed_w     = [8, 88, 20]
    else:
        # Social Sciences / other languages: Total 3rd, Weight last
        all_headers = ["#", "Chapter", "Total Periods"] + pt_headers + ["Weight"]
        fixed_w     = [8, 88, 20]
    pt_w       = [22] * len(sorted_pts)
    tail_w     = [24]
    col_widths = fixed_w + pt_w + tail_w

    # Header row
    pdf.set_font("Helvetica", "B", 7)
    pdf.set_fill_color(44, 62, 80)
    pdf.set_text_color(255, 255, 255)
    for h, w in zip(all_headers, col_widths):
        pdf.cell(w, 7, h, border=1, fill=True, align="C")
    pdf.ln()

    # Data rows
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(30, 30, 30)
    grand_p = 0
    grand_m = 0
    col_sums = {pt["mins"]: 0 for pt in sorted_pts}

    # Chapter title is always col index 1; its column width drives wrapping
    TITLE_COL_IDX = 1
    title_col_w = col_widths[TITLE_COL_IDX]
    LINE_H = 6  # standard cell line height (mm)

    def _safe_str(s: str, w_col: int) -> str:
        """Encode to latin-1 safely for fpdf.

        FPDF's built-in Helvetica is latin-1 only, so common Unicode
        punctuation that appears in chapter titles (em/en dashes, curly
        quotes, ellipsis) would otherwise be rendered as `?`. Map those
        to ASCII equivalents first, then fall back to replacement for
        anything else.
        """
        _PUNCT_MAP = {
            "‒": "-", "–": "-", "—": "-", "―": "-",
            "‘": "'", "’": "'", "‚": "'", "‛": "'",
            "“": '"', "”": '"', "„": '"', "‟": '"',
            "…": "...",
            "•": "*",
            " ": " ",
        }
        s = s.translate({ord(k): v for k, v in _PUNCT_MAP.items()})
        try:
            s.encode("latin-1")
            return s
        except Exception:
            return s.encode("latin-1", "replace").decode("latin-1")

    def _count_title_lines(title: str, col_w: float) -> int:
        """Estimate how many wrapped lines the title will occupy."""
        # fpdf wraps at ~(col_w / char_width_approx) chars; at font size 7,
        # Helvetica average char width ≈ 1.7 mm.
        chars_per_line = max(1, int(col_w / 1.7))
        words = title.split()
        line_len = 0
        lines = 1
        for word in words:
            token_len = len(word) + (1 if line_len > 0 else 0)
            if line_len + token_len > chars_per_line:
                lines += 1
                line_len = len(word)
            else:
                line_len += token_len
        return lines

    for idx, (ch, alloc) in enumerate(zip(chs, allocs)):
        if uses_effort_index:
            ei = ch.get("effort_index", 0)
            tail_val = str(ei) if (isinstance(ei, (int, float)) and ei > 0) else "-"
        else:
            wt = ch.get("chapter_weight", 0)
            tail_val = str(wt) if wt else "-"
        tp = alloc.get("total", 0)
        grand_p += tp
        grand_m += sum(alloc.get(pt["mins"], 0) * pt["mins"] for pt in sorted_pts)
        fill = (idx % 2 == 0)
        if fill:
            pdf.set_fill_color(248, 247, 245)
        else:
            pdf.set_fill_color(255, 255, 255)

        chapter_title = ch.get("chapter_title", "")

        # # | Chapter | Total | [pt-type cols...] | Effort Idx or Weight
        row_vals = [f"{ch['chapter_number']:02d}", chapter_title, str(tp)]
        for pt in sorted_pts:
            v = alloc.get(pt["mins"], 0)
            col_sums[pt["mins"]] += v
            row_vals.append(str(v))
        row_vals.append(tail_val)

        # Determine row height — title may wrap
        n_title_lines = _count_title_lines(chapter_title, title_col_w)
        row_h = max(LINE_H, n_title_lines * LINE_H)

        row_start_x = pdf.get_x()
        row_start_y = pdf.get_y()

        for col_i, (val, w) in enumerate(zip(row_vals, col_widths)):
            x_cur = row_start_x + sum(col_widths[:col_i])
            if col_i == TITLE_COL_IDX:
                # Multi-cell for the title column so long titles wrap left-aligned
                safe_val = _safe_str(val, w)
                # Fill background then draw border rectangle
                if fill:
                    pdf.set_fill_color(248, 247, 245)
                    pdf.set_xy(x_cur, row_start_y)
                    pdf.cell(w, row_h, "", border=1, fill=True, align="L")
                else:
                    pdf.set_xy(x_cur, row_start_y)
                    pdf.cell(w, row_h, "", border=1, fill=False, align="L")
                # Overlay the wrapped text inside the bordered cell
                pdf.set_xy(x_cur + 1, row_start_y + 0.5)
                pdf.multi_cell(w - 2, LINE_H, safe_val, border=0, align="L", fill=False)
            else:
                pdf.set_xy(x_cur, row_start_y)
                safe_val = _safe_str(val, w)
                pdf.cell(w, row_h, safe_val, border=1, fill=fill, align="C")

        pdf.set_xy(row_start_x, row_start_y + row_h)

    # Footer row
    pdf.set_font("Helvetica", "B", 7)
    pdf.set_fill_color(214, 228, 248)
    # Number of blank cells before the period-type sums mirrors the fixed column count
    # Footer: blank | Total | grand_p | [pt-col sums...] | blank (last col = EI or Weight)
    foot_vals = ["", "Total", str(grand_p)]
    for pt in sorted_pts:
        foot_vals.append(str(col_sums[pt["mins"]]))
    foot_vals.append("")
    for val, w in zip(foot_vals, col_widths):
        pdf.cell(w, 6, val, border=1, fill=True, align="C")
    pdf.ln()

    # Footnote — wording depends on subject group
    pdf.ln(3)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(100, 100, 100)
    h_g, m_g = divmod(grand_m, 60)
    time_str = f"{h_g}h {m_g}min" if h_g else f"{m_g} min"
    if is_english:
        footnote = (
            f"Total: {grand_p} periods · {time_str}   |   "
            "Periods allocated using the Largest Remainder Method (LRM) weighted by chapter effort index."
        )
    elif is_science:
        footnote = (
            f"Total: {grand_p} periods · {time_str}   |   "
            "Periods allocated using the Largest Remainder Method (LRM) weighted by chapter effort index."
        )
    else:
        footnote = (
            "Periods allocated using the Largest Remainder Method (LRM) "
            f"weighted by chapter competency load.   "
            f"Total: {grand_p} periods · {time_str}"
        )
    pdf.cell(0, 5, footnote, ln=True)

    # ── "About the Effort Index" block — English and Science/Mathematics ──────
    if is_english:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(26, 68, 128)
        pdf.cell(0, 5, "About the Effort Index", ln=True)
        pdf.set_draw_color(147, 188, 232)
        pdf.set_line_width(0.3)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(1)

        _ei_rows_en = [
            ("What it measures:",
             "The effort index tells you how much classroom time a chapter typically needs compared to other"
             " chapters in the subject. Chapters with a higher effort index get more periods; chapters with"
             " a lower one get fewer. It is calculated from four signals, each scored on a simple scale."),
            ("Spine load (x2):",
             "How many types of classroom work (reading for comprehension, listening, speaking, writing,"
             " vocabulary, beyond-text) appear on average per section. More types = higher score."),
            ("Task density (x1.5):",
             "How many tasks appear on average within each block of work. More tasks per block = higher score."),
            ("Writing demand (x1.5):",
             "Total exercise items under Writing and Beyond-the-Text across the chapter. These take longer"
             " to complete and assess, so a heavier count raises the score."),
            ("Project load (x1):",
             "How many Beyond-the-Text sections the chapter has. Each one adds to the score as these"
             " activities need extra planning time."),
            ("Note:",
             "The four scores are combined with fixed weights to give the effort index. Only relative values"
             " matter - it is used to share your available periods across chapters in proportion to their load."),
        ]
        _lbl_w = 44
        _body_w = 180 - _lbl_w
        for lbl, body in _ei_rows_en:
            y0 = pdf.get_y()
            pdf.set_font("Helvetica", "B", 6.5)
            pdf.set_text_color(26, 68, 128)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(_lbl_w, 4, lbl, ln=False)
            y1 = pdf.get_y()
            pdf.set_xy(pdf.l_margin + _lbl_w, y0)
            pdf.set_font("Helvetica", "", 6.5)
            pdf.set_text_color(75, 75, 75)
            pdf.multi_cell(_body_w, 4, body)
            y2 = pdf.get_y()
            pdf.set_y(max(y1, y2))
            pdf.ln(1)

    elif is_twau:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(26, 68, 128)
        pdf.cell(0, 5, "About the Effort Index", ln=True)
        pdf.set_draw_color(147, 188, 232)
        pdf.set_line_width(0.3)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(1)

        _ei_rows_twau = [
            ("Formula:",
             "effort_index = (Conceptual demand x3) + (Task load x2) + (Project load x1.5) + Map work"),
            ("Conceptual demand (x3):",
             "How abstract the chapter's reasoning is. 1 = concrete/tangible (e.g. family, simple"
             " observations); 2 = slight abstraction — comparisons, simple cause-and-effect, or"
             " categorisation of familiar objects; 3 = classification or material properties;"
             " 4 = multi-step reasoning or concepts requiring inference beyond direct observation"
             " (e.g. seasonal cycles, ecosystem interdependence); 5 = geological, astronomical,"
             " or cultural-history abstraction. Judged from the chapter content, not the grade alone."),
            ("Task load (x2):",
             "Discrete score (0-3) based on the total count of student tasks (Activities, Discuss,"
             " Write, Find out, Draw, Let us reflect items). 0 = fewer than 10; 1 = 10-20;"
             " 2 = 21-30; 3 = more than 30."),
            ("Project load (x1.5):",
             "0 = none; 1 = light (multi-day observation, e.g. watch a plant grow over a week);"
             " 2 = substantial (artefact construction or sustained build project)."),
            ("Map work (x1):",
             "0 = no maps; 1 = map reading; 2 = map drawing or regional comparison."),
            ("Note:",
             "Only relative values matter - the effort index is used to share your available"
             " periods across chapters in proportion to their load."),
        ]
        _lbl_w = 54
        _body_w = 180 - _lbl_w
        for lbl, body in _ei_rows_twau:
            y0 = pdf.get_y()
            pdf.set_font("Helvetica", "B", 6.5)
            pdf.set_text_color(26, 68, 128)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(_lbl_w, 4, lbl, ln=False)
            y1 = pdf.get_y()
            pdf.set_xy(pdf.l_margin + _lbl_w, y0)
            pdf.set_font("Helvetica", "", 6.5)
            pdf.set_text_color(75, 75, 75)
            pdf.multi_cell(_body_w, 4, body)
            y2 = pdf.get_y()
            pdf.set_y(max(y1, y2))
            pdf.ln(1)

    elif is_math_prep:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(26, 68, 128)
        pdf.cell(0, 5, "About the Effort Index", ln=True)
        pdf.set_draw_color(147, 188, 232)
        pdf.set_line_width(0.3)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(1)

        _ei_rows_mp = [
            ("Formula:",
             "effort_index = (Conceptual demand x2) + (Task load x2)"
             " + (Exploration load x1.5) + (Procedural load x1.5)"),
            ("Conceptual demand (x2):",
             "How abstract the chapter's reasoning is on the concrete-to-symbolic path."
             " 1 = fully concrete (counting, matching, sorting tangible objects);"
             " 2 = slight abstraction (place value, simple patterns, measurement with standard units);"
             " 3 = symbolic or multi-step reasoning (multi-digit operations, fraction concepts,"
             " area/perimeter reasoning)."),
            ("Task load (x2):",
             "Discrete tier from total task count."
             " 0 = fewer than 8 tasks; 1 = 8-15; 2 = 16-25; 3 = more than 25."),
            ("Exploration load (x1.5):",
             "Share of hands-on, manipulative, or game-based tasks."
             " 0 = none; 1 = a few; 2 = prominently exploratory."),
            ("Procedural load (x1.5):",
             "Share of compute, convert, or drill tasks."
             " 0 = none; 1 = moderate; 2 = heavily procedural."),
            ("Note:",
             "Only relative values matter - the effort index is used to share your available"
             " periods across chapters in proportion to their load."),
        ]
        _lbl_w = 54
        _body_w = 180 - _lbl_w
        for lbl, body in _ei_rows_mp:
            y0 = pdf.get_y()
            pdf.set_font("Helvetica", "B", 6.5)
            pdf.set_text_color(26, 68, 128)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(_lbl_w, 4, lbl, ln=False)
            y1 = pdf.get_y()
            pdf.set_xy(pdf.l_margin + _lbl_w, y0)
            pdf.set_font("Helvetica", "", 6.5)
            pdf.set_text_color(75, 75, 75)
            pdf.multi_cell(_body_w, 4, body)
            y2 = pdf.get_y()
            pdf.set_y(max(y1, y2))
            pdf.ln(1)

    elif is_science:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(26, 68, 128)
        pdf.cell(0, 5, "About the Effort Index", ln=True)
        pdf.set_draw_color(147, 188, 232)
        pdf.set_line_width(0.3)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(1)

        _ei_rows_sci = [
            ("What it measures:",
             "The effort index tells you how much classroom time a chapter typically needs compared to other"
             " chapters in the subject. Chapters with a higher effort index get more periods; chapters with"
             " a lower one get fewer. It is calculated from four signals read from the chapter content."),
            ("Formula:",
             "effort_index = (Conceptual demand x2) + (Activity load x2)"
             " + (Demo load x1.5) + (Exercise execution load x2)"),
            ("Conceptual demand (x2):",
             "The cognitive complexity of exercises and questions in the chapter (1-3). High-order thinking or"
             " multi-step reasoning raises the score."),
            ("Activity load (x2):",
             "A discrete 0-3 tier from the number of hands-on activities students perform themselves:"
             " 0 = none; 1 = 1-3 (light); 2 = 4-7 (standard); 3 = 8 or more (activity-heavy)."),
            ("Demo load (x1.5):",
             "A discrete 0-2 tier from the number of teacher demonstrations: 0 = none; 1 = 1-2;"
             " 2 = 3 or more. These need preparation and focused class attention."),
            ("Exercise execution load (x2):",
             "The weight of multi-step calculation or diagram production in the exercises (0-2)."
             " A heavier execution load means more time for guided practice and assessment."),
            ("Note:",
             "All four signals are bounded discrete tiers combined with fixed weights, so no single signal"
             " can dominate. Only relative values matter - the index shares your available periods across"
             " chapters in proportion to their load."),
        ]
        _lbl_w = 50
        _body_w = 180 - _lbl_w
        for lbl, body in _ei_rows_sci:
            y0 = pdf.get_y()
            pdf.set_font("Helvetica", "B", 6.5)
            pdf.set_text_color(26, 68, 128)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(_lbl_w, 4, lbl, ln=False)
            y1 = pdf.get_y()
            pdf.set_xy(pdf.l_margin + _lbl_w, y0)
            pdf.set_font("Helvetica", "", 6.5)
            pdf.set_text_color(75, 75, 75)
            pdf.multi_cell(_body_w, 4, body)
            y2 = pdf.get_y()
            pdf.set_y(max(y1, y2))
            pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ── Constants ─────────────────────────────────────────────────────────────────

LOGO_PATH     = MISC_DIR / "aruvi_logo-transparent.png"

DURATION_OPTIONS = [30, 35, 40, 45, 50, 60]
WEIGHT_LABEL     = {3: "Central", 2: "Substantive", 1: "Present"}

GRADES = [
    "Grade III", "Grade IV", "Grade V", "Grade VI",
    "Grade VII", "Grade VIII", "Grade IX",
]

SUBJECTS = [
    "English", "Mathematics", "Science",
    "Social Science", "The World Around Us",
]

# Grades available per subject.
# "The World Around Us" → preparatory only (III–V)
# Science / Social Science → middle + secondary only (VI–X)
# All other subjects → all grades (III–X)
_PREPARATORY_GRADES = ["Grade III", "Grade IV", "Grade V"]
_MIDDLE_SEC_GRADES  = ["Grade VI", "Grade VII", "Grade VIII", "Grade IX"]

def grades_for_subject(subject: str) -> list:
    if subject == "The World Around Us":
        return _PREPARATORY_GRADES
    if subject in ("Science", "Social Science"):
        return _MIDDLE_SEC_GRADES
    return GRADES

# ── Page config ───────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Aruvi",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Restore state from URL query params (survives pill-click reloads) ─────────

query = st.query_params

if "role"    in query and query["role"]    in ("Allocate", "Generate", "My Plans"):
    _prev_role = st.session_state.get("role", "")
    st.session_state.role    = query["role"]
    if query["role"] == "My Plans" and _prev_role != "My Plans":
        st.session_state.myplans_should_collapse = True
if "grade" not in query and st.session_state.get("grade"):
    pass  # keep existing session state grade
if "subject" not in query and st.session_state.get("subject"):
    pass  # keep existing session state subject
if "grade"   in query and query["grade"]   in GRADES:
    st.session_state.grade   = query["grade"]
if "subject" in query and query["subject"] in SUBJECTS:
    st.session_state.subject = query["subject"]
if "ch"      in query:
    try: st.session_state.teacher_ch_idx = int(query["ch"])
    except ValueError: pass

# Persist alloc_chs and alloc_pts from URL into session state
# so they survive across reruns and are available wherever the tab renders
if "alloc_chs" in query and query["alloc_chs"]:
    st.session_state["alloc_chs"] = query["alloc_chs"]
if "alloc_pts" in query and query["alloc_pts"]:
    st.session_state["alloc_pts"] = query["alloc_pts"]
# Clear them if a fresh load (no alloc params in URL)
if "alloc_chs" not in query and "alloc_pts" not in query:
    if "alloc_chs" in st.session_state: del st.session_state["alloc_chs"]
    if "alloc_pts" in st.session_state: del st.session_state["alloc_pts"]

# Defaults on first load
if "role"    not in st.session_state: st.session_state.role    = "Allocate"
if "grade"   not in st.session_state: st.session_state.grade   = None
if "subject" not in st.session_state: st.session_state.subject = None

# ── Image helpers ─────────────────────────────────────────────────────────────

def _img_src(path: Path) -> str:
    """Load a PNG file as a base64 data URI. Returns '' if the file is missing."""
    try:
        b64 = base64.b64encode(path.read_bytes()).decode()
        return f"data:image/png;base64,{b64}"
    except Exception:
        return ""


LOGO_SRC    = _img_src(LOGO_PATH)
GRADE_SRC   = _img_src(MISC_DIR / "grade.png")
SUBJECT_SRC = _img_src(MISC_DIR / "subject.png")
CHAPTER_SRC = _img_src(MISC_DIR / "chapter.png")
SAVED_SRC   = _img_src(MISC_DIR / "saved.png")            # My Plans "Saved" filter icon
PERIOD_SRC      = _img_src(MISC_DIR / "period.png")       # row header add-icon
TIME_SRC        = _img_src(MISC_DIR / "time.png")         # "Available time" label icon
FULL_PERIOD_SRC = _img_src(MISC_DIR / "full_period.png")  # Principal "Period Budget" label icon
WATERMARK_SRC   = _img_src(MISC_DIR / "aruvi_logo-transparent.png")  # Main body watermark
_rotate_logo_src = _img_src(
    PROJECT_ROOT / "aruvi_streamlit" / "static" / "aruvi_logo_rotate.png"
)                                                                        # Progress popup spinning logo


# ── CSS + JS ───────────────────────────────────────────────────────────────────

st.markdown("""
<style>

/* ═══════════════════════════════════════════════════
   FIXED TOP NAV BAR
   Change 1: width 100vw, left 0 — spans full viewport
   including over the sidebar. overflow:visible ensures
   logo/brand are never clipped regardless of sidebar state.
   ═══════════════════════════════════════════════════ */
.aruvi-topnav {
    position: fixed !important;
    top: 0 !important;
    left: 0 !important;
    width: 100vw !important;
    z-index: 99999 !important;
    background: #f5f3ef;
    border-bottom: 1px solid #d9d6d0;
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0.6rem 1.5rem;
    box-sizing: border-box;
    min-height: 72px;
    overflow: visible;
}

/* Left: logo + brand — never clip or hide */
.topnav-left {
    flex: 0 0 auto;
    min-width: 180px;
    display: flex;
    flex-direction: row;
    align-items: center;
    gap: 0.75rem;
    overflow: visible;
}
.topnav-left img {
    width: 56px;
    height: 56px;
    object-fit: contain;
    display: block;
    background: transparent;
    flex-shrink: 0;
}
/* Brand: wordmark above slogan */
.topnav-brand {
    display: flex;
    flex-direction: column;
    gap: 0.18rem;
    overflow: visible;
    white-space: nowrap;
}
.topnav-wordmark {
    font-size: 0.65rem;
    font-weight: 500;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    color: #5a5754;
    white-space: nowrap;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
    line-height: 1;
}
.topnav-slogan {
    font-size: 0.55rem;
    font-weight: 400;
    letter-spacing: 0.01em;
    color: #5a5754;
    white-space: nowrap;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
    line-height: 1;
}

/* Centre: pill toggle */
.topnav-center {
    flex: 1;
    display: flex;
    align-items: center;
    justify-content: center;
}

/* Right: empty balancer */
.topnav-right {
    flex: 0 0 auto;
    min-width: 180px;
}

/* Pill container */
.aruvi-topnav-inner {
    display: inline-flex;
    align-items: center;
    background: #e8e5e0;
    border-radius: 999px;
    padding: 3px 4px;
    gap: 2px;
}

/* Individual pills */
.aruvi-pill {
    display: inline-block;
    padding: 0.3rem 1.45rem;
    border-radius: 999px;
    font-size: 0.82rem;
    font-weight: 500;
    letter-spacing: 0.01em;
    cursor: pointer;
    transition: background 0.15s, color 0.15s;
    color: #6b6866;          /* warm grey matching logo palette */
    background: transparent;
    border: none;
    text-decoration: none !important;
    user-select: none;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
}
.aruvi-pill:link    { color: #6b6866 !important; text-decoration: none !important; }
.aruvi-pill:visited { color: #6b6866 !important; text-decoration: none !important; }
.aruvi-pill:active  { color: #6b6866 !important; text-decoration: none !important; }
.aruvi-pill:focus   { color: #6b6866 !important; text-decoration: none !important; }
.aruvi-pill:hover   { color: #2c2a27 !important; text-decoration: none !important; background: rgba(0,0,0,0.04); }
.aruvi-pill.active {
    background: #ffffff;
    color: #2c2a27;          /* same dark tone as topnav wordmark */
    font-weight: 600;
    box-shadow: 0 1px 3px rgba(0,0,0,0.12);
}



/* ═══════════════════════════════════════════════════
   PUSH CONTENT DOWN below the fixed top nav
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] > div:first-child {
    padding-top: 5.8rem !important;
    display: flex !important;
    flex-direction: column !important;
    min-height: 100vh !important;
    box-sizing: border-box !important;
}
.main .block-container {
    background-color: #ffffff !important;
    padding: 5.8rem 3rem 2rem 2.5rem !important;
    max-width: none;
}
header[data-testid="stHeader"] {
    background: rgba(0,0,0,0) !important;
    top: 72px !important;
}

/* ═══════════════════════════════════════════════════
   GLOBAL
   ═══════════════════════════════════════════════════ */
html, body {
    background-color: #f5f3ef;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
}
.stApp {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
}
[data-testid="stAppViewContainer"] {
    background-color: #ffffff !important;
}

/* ═══════════════════════════════════════════════════
   MAIN BODY WATERMARK
   Aruvi logo rendered as a very-light-grey centred
   watermark behind all content. Uses ::before so the
   opacity does not bleed through to child elements.
   ═══════════════════════════════════════════════════ */
[data-testid="stMain"] {
    background-color: #ffffff !important;
    position: relative;
}
/* watermark rule injected below via f-string */

/* 24 × 24 px grid — very faint lines */
[data-testid="stMain"]::after {
    content: "";
    position: fixed;
    inset: 0;
    background-image:
        linear-gradient(rgba(180, 174, 165, 0.07) 1px, transparent 1px),
        linear-gradient(90deg, rgba(180, 174, 165, 0.07) 1px, transparent 1px);
    background-size: 24px 24px;
    pointer-events: none;
    z-index: 0;
}

/* ═══════════════════════════════════════════════════
   SIDEBAR — must render below our topnav
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] {
    background-color: #eeece8;
    border-right: 1px solid #d9d6d0;
    z-index: 100 !important;
}

/* ═══════════════════════════════════════════════════
   PREVENT STACKING-CONTEXT BREAKS
   Streamlit may apply CSS transforms to app containers
   for animations. Any ancestor with transform:non-none
   makes position:fixed children act like position:absolute,
   breaking left:0/width:100vw on the topnav.
   Force transforms off on every Streamlit wrapper.
   ═══════════════════════════════════════════════════ */
.stApp,
[data-testid="stAppViewContainer"],
[data-testid="stMain"],
[data-testid="block-container"],
.main {
    transform: none !important;
    will-change: auto !important;
}

/* ═══════════════════════════════════════════════════
   SIDEBAR FIELD LABEL ROW  (icon + uppercase name)
   Rendered above each selectbox via st.markdown.
   ═══════════════════════════════════════════════════ */
.sidebar-field-label {
    display: flex;
    align-items: center;
    gap: 0.38rem;
    margin-top: 0.85rem;
    margin-bottom: 0.75rem;
}
.field-icon {
    width: 23px;
    height: 23px;
    object-fit: contain;
    opacity: 0.72;
    flex-shrink: 0;
}
.field-icon-grade {
    width: 27px;
    height: 27px;
    object-fit: contain;
    opacity: 1.0;
    flex-shrink: 0;
}
.field-label-text {
    font-size: 0.70rem;
    font-weight: 500;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: #5a5754;
    line-height: 1;
}

/* ═══════════════════════════════════════════════════
   SIDEBAR SELECTBOX: flat / no-box style
   The box border and background are stripped away.
   Value sits flush-left directly below the label row.
   A › chevron (via ::after) signals the dropdown.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [data-testid="stSelectbox"] {
    position: relative !important;
    margin-top: 0 !important;
    margin-bottom: 0.5rem !important;
}
/* Strip box chrome from the BaseUI select control */
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] > div:first-child {
    border: none !important;
    background: transparent !important;
    box-shadow: none !important;
    padding: 0 1.2rem 0 0 !important;
    min-height: 28px !important;
}
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"],
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="base-input"] {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    padding: 0 !important;
}
/* Value text: flush left, medium-dark grey */
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] [data-baseweb="input"],
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] > div > div > div {
    color: #3d3b38 !important;
    font-size: 0.84rem !important;
    padding: 0 !important;
    line-height: 1.4 !important;
}

/* ═══════════════════════════════════════════════════
   GRADE / SUBJECT / CHAPTER SELECTBOX
   Grey rounded box (filled with sidebar bg) + smaller
   dark-grey value font.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] div[class*="st-key-grade_select"] [data-baseweb="select"] > div:first-child,
section[data-testid="stSidebar"] div[class*="st-key-subject_select"] [data-baseweb="select"] > div:first-child,
section[data-testid="stSidebar"] div[class*="st-key-teacher_ch_select"] [data-baseweb="select"] > div:first-child,
section[data-testid="stSidebar"] div[class*="st-key-mp_grade_select"] [data-baseweb="select"] > div:first-child,
section[data-testid="stSidebar"] div[class*="st-key-mp_subject_select"] [data-baseweb="select"] > div:first-child,
section[data-testid="stSidebar"] div[class*="st-key-mp_saved_select"] [data-baseweb="select"] > div:first-child {
    border: 1px solid #d0cdc9 !important;
    border-radius: 8px !important;
    background: #ffffff !important;
    padding: 6px 8px 6px 10px !important;
    box-shadow: none !important;
    min-height: 34px !important;
    display: flex !important;
    align-items: center !important;
}
/* Value text + placeholder: legible size, high-contrast dark */
section[data-testid="stSidebar"] div[class*="st-key-grade_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-grade_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-grade_select"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] div[class*="st-key-subject_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-subject_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-subject_select"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] div[class*="st-key-teacher_ch_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-teacher_ch_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-teacher_ch_select"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] div[class*="st-key-mp_grade_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-mp_grade_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-mp_grade_select"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] div[class*="st-key-mp_subject_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-mp_subject_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-mp_subject_select"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] div[class*="st-key-mp_saved_select"] [data-baseweb="select"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-mp_saved_select"] [data-baseweb="select"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-mp_saved_select"] [data-baseweb="select"] span {
    font-size: 0.76rem !important;
    color: #2c2a27 !important;
}

/* ═══════════════════════════════════════════════════
   DURATION SELECTBOX  — centre value text
   Targets dur_sel_0 (Teacher) and dur_sel_p0 (Principal).
   Strategy:
     • Remove the right-offset padding on the control container
       (the arrow div is a flex-sibling, so it stays right naturally).
     • Give the VALUE sub-container flex:1 + justify-content:center
       so it fills remaining space and centres its content.
     • Force every text node inside it to centre.
   ═══════════════════════════════════════════════════ */

/* 1. Control container: flex row, no extra right padding */
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="select"] > div:first-child {
    padding: 0 !important;
    display: flex !important;
    align-items: center !important;
}
/* 2. Value sub-container (first child of control): fill width, centre */
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="select"] > div:first-child > div:first-child {
    flex: 1 1 0% !important;
    display: flex !important;
    justify-content: center !important;
    align-items: center !important;
    min-width: 0 !important;
}
/* 3. Value / placeholder text nodes — block + centred */
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="value"],
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="placeholder"],
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="select"] > div:first-child span {
    display: block !important;
    text-align: center !important;
    width: 100% !important;
}
/* 4. White rounded box — overrides the flat/no-border style */
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="select"] > div:first-child {
    border: 1px solid #d0cdc9 !important;
    border-radius: 8px !important;
    background: #ffffff !important;
    padding: 4px 6px !important;
}
section[data-testid="stSidebar"] div[class*="st-key-dur_sel_"] [data-baseweb="select"] > div:first-child > div {
    background: #ffffff !important;
}

/* ═══════════════════════════════════════════════════
   COUNT NUMBER INPUT  — centre value, hide built-in steps
   Applies to all cnt_ keys (Teacher cnt_0/1/… and
   Principal cnt_p0/1/…). Custom +/− buttons are used.
   ═══════════════════════════════════════════════════ */
/* Hide Streamlit's built-in step buttons — custom ±
   st.buttons are used on both Teacher and Principal. */
section[data-testid="stSidebar"] [class*="st-key-cnt_"] [data-testid="stNumberInput"] button {
    display: none !important;
}

/* ═══════════════════════════════════════════════════
   SIDEBAR SECTION LABEL
   ═══════════════════════════════════════════════════ */
.sect-label {
    font-size: 0.68rem;
    font-weight: 500;
    letter-spacing: 0.10em;
    text-transform: uppercase;
    color: #5a5754 !important;
    margin: 0.5rem 0 0.5rem 0;
    display: flex;
    align-items: center;
    gap: 0.35rem;
}

/* ═══════════════════════════════════════════════════
   TIGHTER SIDEBAR VERTICAL RHYTHM
   Reduces Streamlit's default block-container gaps
   so Grade / Subject / Chapter / Generate sit closer.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] > div[data-testid="element-container"] {
    margin-bottom: 0 !important;
    padding-bottom: 0 !important;
}
/* Collapse space between a sect-label row and the row immediately below it */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] > div[data-testid="element-container"]:has(.sect-label) + div[data-testid="element-container"] {
    margin-top: 0 !important;
    padding-top: 0 !important;
}
section[data-testid="stSidebar"] .sidebar-field-label {
    margin-top: 0.55rem !important;
    margin-bottom: 0.45rem !important;
}
section[data-testid="stSidebar"] [data-testid="stSelectbox"] {
    margin-bottom: 0.3rem !important;
}
section[data-testid="stSidebar"] hr,
section[data-testid="stSidebar"] [data-testid="stDivider"] {
    margin-top: 0.5rem !important;
    margin-bottom: 0.5rem !important;
}

/* ═══════════════════════════════════════════════════
   PERIOD ⊕ ICON  (clickable, adds a period block)
   ═══════════════════════════════════════════════════ */
.period-icon {
    width: 20px;
    height: 20px;
    object-fit: contain;
    cursor: pointer;
    opacity: 1.0;
    transition: opacity 0.15s;
    flex-shrink: 0;
    vertical-align: middle;
}
.period-icon:hover { opacity: 0.75; }
/* Text fallback when PNG is missing */
.period-icon-text {
    font-size: 1.0rem;
    line-height: 1;
    cursor: pointer;
    color: #9c9693;
    user-select: none;
    transition: color 0.15s, opacity 0.15s;
    opacity: 0.75;
}
.period-icon-text:hover { color: #c96442; opacity: 1.0; }

/* ═══════════════════════════════════════════════════
   REMOVE-ROW ✕ BUTTON  — dark, vertically centred
   ═══════════════════════════════════════════════════ */
/* Align the whole rm column contents to centre vertically */
section[data-testid="stSidebar"] div[class*="st-key-rm_"] {
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    height: 100% !important;
}
section[data-testid="stSidebar"] div[class*="st-key-rm_"] button {
    color: #3d3b38 !important;
    font-size: 0.85rem !important;
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
    min-height: unset !important;
    line-height: 1 !important;
}
section[data-testid="stSidebar"] div[class*="st-key-rm_"] button:hover {
    color: #c0392b !important;
    background: transparent !important;
}

/* ═══════════════════════════════════════════════════
   SIDEBAR BUTTONS — ensure text always visible
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] button,
section[data-testid="stSidebar"] button p,
section[data-testid="stSidebar"] button span,
section[data-testid="stSidebar"] button div {
    color: #3d3b38 !important;
    visibility: visible !important;
    opacity: 1 !important;
}

/* ═══════════════════════════════════════════════════
   PERIOD COUNT STEPPER  (+/− buttons)
   ::after pseudo-elements supply the visible symbol.
   Streamlit's theme cannot override ::after content
   color, so this always renders on the native background.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [class*="st-key-plus_"] button,
section[data-testid="stSidebar"] [class*="st-key-minus_"] button {
    background: transparent !important;
    border: 1px solid #c8c4be !important;
    border-radius: 4px !important;
    min-height: 28px !important;
    padding: 0 !important;
    position: relative !important;
}
/* Hide Streamlit's own (theme-coloured) label */
section[data-testid="stSidebar"] [class*="st-key-plus_"] button *,
section[data-testid="stSidebar"] [class*="st-key-minus_"] button * {
    visibility: hidden !important;
}
/* Overlay our own symbol via ::after — CSS-owned, theme-proof */
section[data-testid="stSidebar"] [class*="st-key-plus_"] button::after {
    content: "+";
    color: #3d3b38;
    font-size: 1.05rem;
    font-weight: 500;
    position: absolute;
    top: 50%;
    left: 50%;
    transform: translate(-50%, -50%);
    visibility: visible !important;
    pointer-events: none;
}
section[data-testid="stSidebar"] [class*="st-key-minus_"] button::after {
    content: "−";
    color: #3d3b38;
    font-size: 1.05rem;
    font-weight: 500;
    position: absolute;
    top: 50%;
    left: 50%;
    transform: translate(-50%, -50%);
    visibility: visible !important;
    pointer-events: none;
}
section[data-testid="stSidebar"] [class*="st-key-plus_"] button:hover,
section[data-testid="stSidebar"] [class*="st-key-minus_"] button:hover {
    background: #e8e5e0 !important;
    border-color: #9c9693 !important;
}

/* ═══════════════════════════════════════════════════
   DURATION NUMBER INPUT  (replaces selectbox)
   In the time-per-period column — keep native arrows
   so user can step through values or type directly.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [class*="st-key-dur_"] input,
section[data-testid="stSidebar"] [class*="st-key-dur_p"] input {
    font-size: 0.84rem !important;
    color: #3d3b38 !important;
    padding: 0.25rem 0.4rem !important;
}

/* ═══════════════════════════════════════════════════
   COUNT NUMBER INPUT  (editable, flanked by +/−)
   Hide native spin arrows — custom buttons are used.
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [class*="st-key-cnt_"] input {
    font-size: 0.84rem !important;
    color: #3d3b38 !important;
    text-align: center !important;
    padding: 0.25rem 0.1rem !important;
    -moz-appearance: textfield !important;
}
section[data-testid="stSidebar"] [class*="st-key-cnt_"] input::-webkit-outer-spin-button,
section[data-testid="stSidebar"] [class*="st-key-cnt_"] input::-webkit-inner-spin-button {
    -webkit-appearance: none !important;
    margin: 0 !important;
}
/* White rounded box on all count inputs (Teacher + Principal) */
section[data-testid="stSidebar"] [class*="st-key-cnt_"] [data-baseweb="input"] {
    background: #ffffff !important;
    border: 1px solid #d0cdc9 !important;
    border-radius: 8px !important;
    box-shadow: none !important;
}
/* BaseUI renders a nested inner-wrapper div inside [data-baseweb="input"]
   that carries its own background — override it to match white. */
section[data-testid="stSidebar"] [class*="st-key-cnt_"] [data-baseweb="input"] > div {
    background: #ffffff !important;
}
/* Keep the inner <input> element transparent so the
   container background shows through cleanly */
section[data-testid="stSidebar"] [class*="st-key-cnt_"] input {
    background: transparent !important;
}

/* ═══════════════════════════════════════════════════
   PERIOD BLOCK COLUMN HEADERS  (Change 4)
   Match .sect-label style but zero top-margin for first row
   ═══════════════════════════════════════════════════ */
.block-col-label {
    font-size: 0.68rem;
    font-weight: 500;
    letter-spacing: 0.10em;
    text-transform: uppercase;
    color: #9c9693;
    margin: 0.5rem 0 0.15rem 0;
    line-height: 1;
    display: block;
}

/* ═══════════════════════════════════════════════════
   SELECTBOX DROPDOWN OPTION LIST
   BaseUI portals the menu outside the sidebar so these
   must be global. Match chosen-value font / size / colour.
   ═══════════════════════════════════════════════════ */
[data-baseweb="popover"] [data-baseweb="menu"],
[data-baseweb="popover"] [role="listbox"] {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif !important;
    background: #f5f3ef !important;
    border: 1px solid #d9d6d0 !important;
    border-radius: 6px !important;
    box-shadow: 0 4px 12px rgba(0,0,0,0.10) !important;
}
[data-baseweb="popover"] [role="option"] {
    font-size: 0.76rem !important;
    color: #2c2a27 !important;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif !important;
    background: transparent !important;
    padding: 0.42rem 0.75rem !important;
}
[data-baseweb="popover"] [role="option"]:hover {
    background: #e8e5e0 !important;
    color: #1a1a1a !important;
}
[data-baseweb="popover"] [aria-selected="true"] {
    background: #e0ddd8 !important;
    color: #1a1a1a !important;
}
/* Placeholder text — same size & colour as selected value */
section[data-testid="stSidebar"] [data-testid="stSelectbox"] [data-baseweb="select"] [data-baseweb="placeholder"] {
    font-size: 0.84rem !important;
    color: #3d3b38 !important;
}

/* ═══════════════════════════════════════════════════
   DROPDOWN ARROW — keep native SVG, just size + colour it
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] [data-baseweb="select"] svg {
    display: block !important;
    width: 14px !important;
    height: 14px !important;
    color: #9c9693 !important;
    opacity: 0.8;
}

/* ═══════════════════════════════════════════════════
   WORKSPACE TABS  (inner tab strip)
   ═══════════════════════════════════════════════════ */
.stTabs [data-baseweb="tab-list"] {
    background: transparent !important;
    border-bottom: 1px solid #d9d6d0 !important;
    gap: 0 !important;
    padding: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    background: transparent !important;
    color: #9c9693 !important;
    border: none !important;
    border-bottom: 2px solid transparent !important;
    border-radius: 0 !important;
    padding: 0.5rem 1.4rem 0.45rem !important;
    font-size: 0.82rem !important;
    font-weight: 400 !important;
    letter-spacing: 0.01em;
    margin-bottom: -1px !important;
    transition: color 0.12s;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #5a5754 !important;
    background: transparent !important;
}
.stTabs [aria-selected="true"] {
    color: #1a1a1a !important;
    border-bottom: 2px solid #2c3e50 !important;
}
.stTabs [data-baseweb="tab-highlight"],
.stTabs [data-baseweb="tab-border"] {
    display: none !important;
}
.stTabs [data-baseweb="tab-panel"] {
    padding: 1.5rem 0 0 0 !important;
}

/* ═══════════════════════════════════════════════════
   NAV PILL BUTTONS — lift into fixed top nav bar
   ═══════════════════════════════════════════════════ */
div[class*="st-key-nav_allocate"],
div[class*="st-key-nav_generate"],
div[class*="st-key-nav_myplans"] {
    position: fixed !important;
    top: 18px !important;
    z-index: 100000 !important;
    margin: 0 !important;
    padding: 0 !important;
}
div[class*="st-key-nav_allocate"] { left: calc(50% - 148px) !important; }
div[class*="st-key-nav_generate"] { left: calc(50% - 44px)  !important; }
div[class*="st-key-nav_myplans"]  { left: calc(50% + 60px)  !important; }

/* Style all three as pill-shaped */
div[class*="st-key-nav_allocate"] button,
div[class*="st-key-nav_generate"] button,
div[class*="st-key-nav_myplans"]  button {
    background: transparent !important;
    border: none !important;
    border-radius: 999px !important;
    color: #6b6866 !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    padding: 0.3rem 1.45rem !important;
    box-shadow: none !important;
    transition: background 0.15s, color 0.15s !important;
}
div[class*="st-key-nav_allocate"] button:hover,
div[class*="st-key-nav_generate"] button:hover,
div[class*="st-key-nav_myplans"]  button:hover {
    background: rgba(0,0,0,0.04) !important;
    color: #2c2a27 !important;
    border: none !important;
}
/* Active pill — white background with shadow, matching original design */
div[class*="st-key-nav_allocate"] button[kind="primary"],
div[class*="st-key-nav_generate"] button[kind="primary"],
div[class*="st-key-nav_myplans"]  button[kind="primary"] {
    background: #ffffff !important;
    color: #2c2a27 !important;
    font-weight: 600 !important;
    box-shadow: 0 1px 3px rgba(0,0,0,0.12) !important;
    border: none !important;
}

/* ═══════════════════════════════════════════════════
   BUTTONS
   ═══════════════════════════════════════════════════ */
div.stButton > button {
    background: transparent;
    border: 1px solid #d9d6d0;
    color: #5a5754;
    border-radius: 5px;
    font-size: 0.82rem;
    transition: border-color 0.15s, color 0.15s, background 0.15s;
}
div.stButton > button:hover {
    border-color: #c96442;
    color: #1a1a1a;
    background: rgba(201, 100, 66, 0.05);
}
div.stButton > button[kind="primary"] {
    background: #c96442 !important;
    border: none !important;
    color: #fff !important;
    font-weight: 500 !important;
}
div.stButton > button[kind="primary"]:hover {
    background: #d97050 !important;
}

/* ═══════════════════════════════════════════════════
   GENERATE BUTTONS  — Teacher + Principal, tall pill style
   Dark slate background · white text · bold · centred
   ═══════════════════════════════════════════════════ */
section[data-testid="stSidebar"] div[class*="st-key-teacher_gen"] button,
section[data-testid="stSidebar"] div[class*="st-key-principal_gen"] button {
    height: 56px !important;
    min-height: 56px !important;
    border-radius: 12px !important;
    background: #2c3e50 !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.18) !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    color: #ffffff !important;
    letter-spacing: 0.02em !important;
    justify-content: center !important;
}
section[data-testid="stSidebar"] div[class*="st-key-teacher_gen"] button:hover,
section[data-testid="stSidebar"] div[class*="st-key-principal_gen"] button:hover {
    background: #3d5166 !important;
    box-shadow: 0 4px 14px rgba(0,0,0,0.24) !important;
}
section[data-testid="stSidebar"] div[class*="st-key-teacher_gen"] button *,
section[data-testid="stSidebar"] div[class*="st-key-principal_gen"] button * {
    color: #ffffff !important;
    visibility: visible !important;
}
/* ✦ icon via ::before — theme-proof on both buttons */
section[data-testid="stSidebar"] div[class*="st-key-teacher_gen"] button::before,
section[data-testid="stSidebar"] div[class*="st-key-principal_gen"] button::before {
    content: "✦";
    font-size: 0.85rem;
    color: #ffffff;
    flex-shrink: 0;
    visibility: visible !important;
    margin-right: 0.35rem;
}
div.stButton > button[disabled],
div.stButton > button:disabled {
    background: #eeece8 !important;
    border: 1px solid #d9d6d0 !important;
    color: #c8c4be !important;
}
div[class*="st-key-lpa_confirm"] button {
    background: #2c3e50 !important;
    border: none !important;
    color: #ffffff !important;
    font-weight: 600 !important;
    border-radius: 9px !important;
}
div[class*="st-key-lpa_confirm"] button:hover {
    background: #3d5166 !important;
}

/* ═══════════════════════════════════════════════════
   TOTAL / ALLOCATION LINE
   ═══════════════════════════════════════════════════ */
.total-line {
    font-size: 0.79rem;
    color: #c96442;
    margin: 0.4rem 0 0.25rem 0;
}
.over-line {
    font-size: 0.76rem;
    color: #c04040;
    margin: 0.1rem 0;
}

/* ═══════════════════════════════════════════════════
   WORKSPACE: CHAPTER HEADER
   ═══════════════════════════════════════════════════ */
.ch-title {
    font-size: 1.35rem;
    font-weight: 400;
    color: #1a1a1a;
    margin-bottom: 0.2rem;
    line-height: 1.3;
}
.ch-meta {
    font-size: 0.74rem;
    color: #9c9693;
    margin-bottom: 1.75rem;
    letter-spacing: 0.015em;
}

/* ═══════════════════════════════════════════════════
   COMPETENCY ROW
   ═══════════════════════════════════════════════════ */
.comp-row {
    display: flex;
    align-items: baseline;
    gap: 0.7rem;
    margin-bottom: 0.15rem;
}
.comp-code   { font-size: 0.92rem; color: #1a1a1a; }
.comp-cg     { font-size: 0.74rem; color: #9c9693; }
.comp-weight {
    font-size: 0.72rem; color: #5a5754;
    background: #e8e5e0; padding: 1px 7px; border-radius: 3px;
}

/* ═══════════════════════════════════════════════════
   INCIDENTAL FOOTNOTE
   ═══════════════════════════════════════════════════ */
.incidental-line {
    font-size: 0.73rem;
    color: #9c9693;
    margin-top: 1.75rem;
    line-height: 1.7;
}

/* ═══════════════════════════════════════════════════
   WORKSPACE PLACEHOLDER
   ═══════════════════════════════════════════════════ */
.ws-placeholder {
    color: #9c9693;
    font-size: 0.88rem;
    padding: 5rem 0 3rem 0;
    text-align: center;
    letter-spacing: 0.01em;
}

/* ═══════════════════════════════════════════════════
   NO-DATA SIDEBAR NOTICE
   ═══════════════════════════════════════════════════ */
.no-data-notice {
    font-size: 0.78rem;
    color: #9c9693;
    margin-top: 1.25rem;
    line-height: 1.6;
}

/* ═══════════════════════════════════════════════════
   EXPANDER
   ═══════════════════════════════════════════════════ */
details > summary {
    font-size: 0.76rem !important;
    color: #5a5754 !important;
    padding: 0.2rem 0 !important;
}
details[open] > summary { color: #1a1a1a !important; }
details > div {
    font-size: 0.82rem !important;
    color: #5a5754 !important;
    line-height: 1.7 !important;
    padding: 0.4rem 0 0.2rem 0 !important;
}

/* ═══════════════════════════════════════════════════
   INFO / WARNING
   ═══════════════════════════════════════════════════ */
div[data-testid="stInfo"] {
    background: #fef8f5 !important;
    border: 1px solid #e8d0c0 !important;
    color: #8b5e4a !important;
    border-radius: 6px !important;
}
div[data-testid="stWarning"] {
    background: #fdf8ec !important;
    border: 1px solid #e8d898 !important;
    color: #7a6520 !important;
    border-radius: 6px !important;
}

/* ═══════════════════════════════════════════════════
   CHECKBOX — label text, box size, spacing, tick colour
   ═══════════════════════════════════════════════════ */

/* Label text — identical to Select All / field-label-text
   Target every element Streamlit may use: span, p, or bare div  */
section[data-testid="stSidebar"] .stCheckbox label span,
section[data-testid="stSidebar"] .stCheckbox label p,
section[data-testid="stSidebar"] .stCheckbox label > div,
section[data-testid="stSidebar"] [data-baseweb="checkbox"] > div,
section[data-testid="stSidebar"] [data-baseweb="checkbox"] > div p,
section[data-testid="stSidebar"] [data-baseweb="checkbox"] > div span {
    font-size: 0.70rem !important;
    color: #5a5754 !important;
    font-weight: 500 !important;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif !important;
}

/* Shrink the tick-box itself */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"] {
    width: 13px !important;
    height: 13px !important;
    min-width: 13px !important;
    min-height: 13px !important;
    border-radius: 3px !important;
    flex-shrink: 0 !important;
}

/* Checked state — dark grey fill, white tick.
   BaseUI injects background as an inline style attribute which
   defeats stylesheet !important rules.  The ::before pseudo-element
   is a child layer that paints ON TOP of the parent's own background,
   so it covers the orange regardless of how it was applied.          */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"] {
    border-color: #c8c4be !important;   /* unchecked border warm grey */
}
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"][aria-checked="true"] {
    border-color: #2c3e50 !important;
    background: #2c3e50 !important;
    background-color: #2c3e50 !important;
    position: relative !important;
    overflow: hidden !important;
}
/* ::before covers any BaseUI inline-style orange injection */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"][aria-checked="true"]::before {
    content: "" !important;
    position: absolute !important;
    inset: 0 !important;
    background: #2c3e50 !important;
    background-color: #2c3e50 !important;
    z-index: 0 !important;
    pointer-events: none !important;
}
/* Also target the inner div BaseUI may use as the colour layer */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"][aria-checked="true"] > div {
    background: #2c3e50 !important;
    background-color: #2c3e50 !important;
}
/* SVG sits above all layers so the tick stays white */
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"][aria-checked="true"] svg {
    position: relative !important;
    z-index: 1 !important;
    width: 10px !important;
    height: 10px !important;
}
section[data-testid="stSidebar"] [data-baseweb="checkbox"] [role="checkbox"][aria-checked="true"] svg path {
    fill: #ffffff !important;
    stroke: #ffffff !important;
}

/* Minimum row spacing between chapters */
section[data-testid="stSidebar"] .stCheckbox {
    margin: 0 !important;
    padding: 0 !important;
    line-height: 1 !important;
}
section[data-testid="stSidebar"] .stCheckbox > label {
    padding-top: 0.03rem !important;
    padding-bottom: 0.03rem !important;
    min-height: 0 !important;
    line-height: 1.2 !important;
    gap: 0.35rem !important;
}
/* Collapse the flex gap on every vertical block that contains checkboxes.
   This is the real source of the large inter-row spacing in Streamlit —
   the parent block's gap property, not the checkbox's own margins.     */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"]:has(.stCheckbox) {
    gap: 0 !important;
}
/* Also zero the element-container wrapper Streamlit puts around each widget */
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"]:has(.stCheckbox)
    > div[data-testid="element-container"] {
    margin: 0 !important;
    padding: 0 !important;
}

/* ═══════════════════════════════════════════════════
   SELECT ALL / DESELECT ALL  — workspace body (Plan workspace)
   Dark slate, matches Generate button
   ═══════════════════════════════════════════════════ */
div[class*="st-key-sel_all"] button,
div[class*="st-key-desel_all"] button {
    background: #2c3e50 !important;
    border: none !important;
    border-radius: 6px !important;
    color: #ffffff !important;
    font-size: 0.72rem !important;
    font-weight: 500 !important;
    height: 28px !important;
    min-height: 28px !important;
    padding: 0 10px !important;
}
div[class*="st-key-sel_all"] button:hover,
div[class*="st-key-desel_all"] button:hover {
    background: #3d5166 !important;
}
div[class*="st-key-sel_all"] button *,
div[class*="st-key-desel_all"] button * {
    color: #ffffff !important;
    visibility: visible !important;
}

/* ═══════════════════════════════════════════════════
   CHAPTER TILE CHECKBOXES — styled as clean tiles
   ═══════════════════════════════════════════════════ */
[data-testid="stMain"] div[class*="st-key-chk_"] label {
    background: #f5f3ef !important;
    border: 1px solid #d0cdc9 !important;
    border-radius: 8px !important;
    padding: 10px 12px !important;
    width: 100% !important;
    margin-bottom: 8px !important;
    cursor: pointer !important;
    display: flex !important;
    align-items: flex-start !important;
    gap: 8px !important;
    min-height: 52px !important;
}
[data-testid="stMain"] div[class*="st-key-chk_"] label:hover {
    border-color: #2c3e50 !important;
    background: #f0f3f6 !important;
}
[data-testid="stMain"] div[class*="st-key-chk_"] input:checked + label,
[data-testid="stMain"] div[class*="st-key-chk_"] [aria-checked="true"] ~ div {
    border-color: #2c3e50 !important;
    background: #f0f3f6 !important;
}
[data-testid="stMain"] div[class*="st-key-chk_"] label span,
[data-testid="stMain"] div[class*="st-key-chk_"] label p {
    font-size: 0.76rem !important;
    color: #2c2a27 !important;
    line-height: 1.35 !important;
}

/* ═══════════════════════════════════════════════════
   DIVIDERS
   ═══════════════════════════════════════════════════ */
hr { border-color: #d9d6d0 !important; }

/* ═══════════════════════════════════════════════════
   SIDEBAR USER FOOTER
   Sticky at bottom of sidebar via flex-column parent
   ═══════════════════════════════════════════════════ */
.sidebar-spacer {
    flex: 1 1 auto;
    min-height: 1.5rem;
}
.sidebar-user-footer {
    flex-shrink: 0;
    position: sticky;
    bottom: 0;
    padding-top: 0;
    padding-bottom: 1rem;
    background: #eeece8;
}
.user-footer-inner {
    display: flex;
    align-items: center;
    gap: 0.65rem;
    padding-top: 0.65rem;
}
.user-avatar {
    width: 28px;
    height: 28px;
    border-radius: 50%;
    background: #d9d6d0;
    color: #5a5754;
    font-size: 0.72rem;
    font-weight: 600;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
    letter-spacing: 0.02em;
    user-select: none;
}
.user-info { display: flex; flex-direction: column; gap: 0.1rem; }
.user-name {
    font-size: 0.82rem;
    font-weight: 500;
    color: #1a1a1a;
    line-height: 1;
}
.user-plan {
    font-size: 0.72rem;
    color: #9c9693;
    line-height: 1;
}

/* ═══════════════════════════════════════════════════
   ASK ARUVI FAB  — fixed bottom-right floating button
   ═══════════════════════════════════════════════════ */
div[class*="st-key-ask_aruvi_fab"] button {
    position: fixed !important;
    bottom: 28px !important;
    right: 28px !important;
    width: 52px !important;
    height: 52px !important;
    border-radius: 50% !important;
    background: #1B2A3B !important;
    color: #ffffff !important;
    font-size: 1.3rem !important;
    border: none !important;
    z-index: 99999 !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.18) !important;
    min-height: unset !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_fab"] button:hover {
    background: #2C7A7B !important;
}
/* Panel — slides in from right */
.aruvi-chat-panel {
    position: fixed;
    top: 72px;
    right: 0;
    width: 260px;
    height: calc(100vh - 72px);
    background: #f5f3ef;
    border-left: 1px solid #d9d6d0;
    z-index: 99998;
    display: flex;
    flex-direction: column;
    box-shadow: -4px 0 16px rgba(0,0,0,0.08);
}
.aruvi-chat-panel-header {
    padding: 14px 16px 10px;
    border-bottom: 1px solid #d9d6d0;
    font-size: 0.68rem;
    font-weight: 500;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: #5a5754;
}
.aruvi-chat-panel-body {
    flex: 1;
    display: flex;
    align-items: center;
    justify-content: center;
    padding: 20px;
    font-size: 0.72rem;
    color: #9c9693;
    text-align: center;
    line-height: 1.6;
}

/* ═══════════════════════════════════════════════════
   EXPORT BUTTONS
   ═══════════════════════════════════════════════════ */
div[class*="st-key-export_docx"] button,
div[class*="st-key-export_pdf"] button {
    background: transparent !important;
    border: 1px solid #2c3e50 !important;
    border-radius: 6px !important;
    color: #2c3e50 !important;
    font-size: 0.76rem !important;
    font-weight: 500 !important;
    height: 32px !important;
    min-height: 32px !important;
}
div[class*="st-key-export_docx"] button:hover,
div[class*="st-key-export_pdf"] button:hover {
    background: #2c3e50 !important;
    color: #ffffff !important;
}
div[class*="st-key-export_docx"] button *,
div[class*="st-key-export_pdf"] button * {
    color: inherit !important;
    visibility: visible !important;
}

/* ═══════════════════════════════════════════════════
   MY PLANS — VIEW / PDF BUTTONS
   ═══════════════════════════════════════════════════ */
div[class*="st-key-view_"] button,
div[class*="st-key-pdf_"] button,
div[class*="st-key-mp_gen_assess_"] button {
    background: #2c3e50 !important;
    border: none !important;
    color: #ffffff !important;
    font-weight: 600 !important;
    border-radius: 8px !important;
}
div[class*="st-key-view_"] button,
div[class*="st-key-pdf_"] button {
    font-size: 0.78rem !important;
}
div[class*="st-key-view_"] button:hover,
div[class*="st-key-pdf_"] button:hover,
div[class*="st-key-mp_gen_assess_"] button:hover {
    background: #3d5166 !important;
}
/* "Generate Assessment" sits in the same narrow column as View — drop the
   font a notch and let the label wrap so "Generate" / "Assessment" stack
   on two lines instead of clipping. */
div[class*="st-key-mp_gen_assess_"] button {
    font-size: 0.78rem !important;
    line-height: 1.1 !important;
    padding: 0.25rem 0.4rem !important;
    white-space: normal !important;
}
div[class*="st-key-mp_gen_assess_"] button *,
div[class*="st-key-mp_gen_assess_"] button p,
div[class*="st-key-mp_gen_assess_"] button span {
    color: #ffffff !important;
    white-space: normal !important;
    word-break: normal !important;
    font-size: 0.78rem !important;
    margin: 0 !important;
}

/* MY PLANS — BACK BUTTONS (match primary / Generate button colours) */
div[class*="st-key-mp_back_"] button {
    background: #c96442 !important;
    border: none !important;
    color: #ffffff !important;
}
div[class*="st-key-mp_back_"] button:hover {
    background: #d97050 !important;
}


/* ═══════════════════════════════════════════════════
   HIDE STREAMLIT CHROME
   ═══════════════════════════════════════════════════ */
#MainMenu, footer { visibility: hidden; }

/* Hide Streamlit's built-in running / status indicator (cyclist / runner animation) */
[data-testid="stStatusWidget"],
[data-testid="stDecoration"],
div[class*="StatusWidget"],
.stSpinner > div > div { display: none !important; }

</style>

""", unsafe_allow_html=True)

# ── MutationObserver: override BaseUI inline-style orange on checkboxes ───────
st.markdown("""<script>
(function() {
    var TARGET_COLOR = '#2c3e50';
    var SIDEBAR_SELECTOR = 'section[data-testid="stSidebar"]';

    function fixCheckbox(el) {
        if (el && el.getAttribute('role') === 'checkbox' &&
                el.getAttribute('aria-checked') === 'true') {
            el.style.setProperty('background', TARGET_COLOR, 'important');
            el.style.setProperty('background-color', TARGET_COLOR, 'important');
            el.style.setProperty('border-color', TARGET_COLOR, 'important');
            var inner = el.querySelector('div');
            if (inner) {
                inner.style.setProperty('background', TARGET_COLOR, 'important');
                inner.style.setProperty('background-color', TARGET_COLOR, 'important');
            }
        }
    }

    function fixAll() {
        var sidebar = document.querySelector(SIDEBAR_SELECTOR);
        if (!sidebar) return;
        sidebar.querySelectorAll('[role="checkbox"][aria-checked="true"]')
               .forEach(fixCheckbox);
    }

    var observer = new MutationObserver(function(mutations) {
        mutations.forEach(function(m) {
            if (m.type === 'attributes') {
                fixCheckbox(m.target);
            } else {
                m.addedNodes.forEach(function(n) {
                    if (n.nodeType === 1) {
                        if (n.getAttribute && n.getAttribute('role') === 'checkbox') {
                            fixCheckbox(n);
                        }
                        n.querySelectorAll && n.querySelectorAll('[role="checkbox"]')
                                               .forEach(fixCheckbox);
                    }
                });
            }
        });
    });

    function attach() {
        var sidebar = document.querySelector(SIDEBAR_SELECTOR);
        if (sidebar) {
            fixAll();
            observer.observe(sidebar, {
                childList: true, subtree: true,
                attributes: true, attributeFilter: ['aria-checked', 'style']
            });
        } else {
            setTimeout(attach, 300);
        }
    }

    if (document.readyState === 'loading') {
        document.addEventListener('DOMContentLoaded', attach);
    } else {
        attach();
    }
})();
</script>""", unsafe_allow_html=True)

# ── Watermark: inject separately so we can embed the base64 data URI ──────────
if WATERMARK_SRC:
    st.markdown(f"""<style>
[data-testid="stMain"]::before {{
    content: "";
    position: fixed;
    inset: 0;
    background-image: url('{WATERMARK_SRC}');
    background-repeat: no-repeat;
    background-position: center center;
    background-size: 480px auto;
    opacity: 0.025;
    filter: grayscale(100%);
    pointer-events: none;
    z-index: 0;
}}
</style>""", unsafe_allow_html=True)

# ── Data loading ──────────────────────────────────────────────────────────────

def _mappings_cache_key(grade: str, subject: str) -> str:
    """Return a string that changes whenever the mappings directory contents change."""
    subj_f  = subject_to_folder(subject)
    grade_f = grade_to_folder(grade)
    mappings_dir = PROJECT_ROOT / f"mirror/chapters/{subj_f}/{grade_f}/mappings"
    if not mappings_dir.exists():
        return "empty"
    files = sorted(mappings_dir.glob("ch_*_mapping.json"))
    return f"{len(files)}:{':'.join(f.name for f in files)}"


@st.cache_data
def load_all_chapters(grade: str, subject: str, _cache_key: str = "") -> list[dict]:
    """Load chapter mapping JSONs for the given grade and subject."""
    subj_f  = subject_to_folder(subject)
    grade_f = grade_to_folder(grade)
    mappings_dir = PROJECT_ROOT / f"mirror/chapters/{subj_f}/{grade_f}/mappings"
    chapters = []
    if not mappings_dir.exists():
        return []
    for path in sorted(mappings_dir.glob("ch_*_mapping.json")):
        try:
            with open(path, encoding="utf-8") as fh:
                data = json.load(fh)
            chapters.append(data)
        except Exception:
            continue
    chapters.sort(key=lambda c: c["chapter_number"])
    return chapters


if st.session_state.grade and st.session_state.subject:
    chapters = load_all_chapters(
        st.session_state.grade,
        st.session_state.subject,
        _cache_key=_mappings_cache_key(st.session_state.grade, st.session_state.subject),
    )
else:
    chapters = []


def ch_label(ch: dict) -> str:
    return f"Ch {ch['chapter_number']:02d} — {ch['chapter_title']}"


def ch_short(ch: dict) -> str:
    t = ch["chapter_title"]
    # Drop any subtitle after " - " or ":" — whichever comes first
    for sep in [" - ", ":"]:
        if sep in t:
            t = t.split(sep)[0].strip()
            break
    return f"Ch {ch['chapter_number']:02d} · {t}"


# ── Period-row callbacks (run before script on each interaction) ───────────────

def _cb_add_row():
    if "_next_row_id" not in st.session_state:
        st.session_state["_next_row_id"] = len(st.session_state.get("period_rows", [0])) + 1
    _new_id = st.session_state["_next_row_id"]
    st.session_state["_next_row_id"] = _new_id + 1
    st.session_state["period_rows"] = st.session_state.get("period_rows", []) + [_new_id]

def _cb_del_row(rid):
    st.session_state["period_rows"] = [
        r for r in st.session_state.get("period_rows", []) if r != rid
    ]

def _cb_inc_cnt(rid, delta):
    st.session_state[f"cnt_{rid}"] = max(
        1, min(999, st.session_state.get(f"cnt_{rid}", 1) + delta)
    )

def _cb_add_row_p():
    _new_id = st.session_state["_next_row_id_p"]
    st.session_state["_next_row_id_p"] = _new_id + 1
    st.session_state["period_rows_p"] = st.session_state["period_rows_p"] + [_new_id]

def _cb_del_row_p(rid):
    st.session_state["period_rows_p"] = [
        r for r in st.session_state["period_rows_p"] if r != rid
    ]


# ── Session state ─────────────────────────────────────────────────────────────

if "role"              not in st.session_state: st.session_state.role              = "Allocate"
if "grade"             not in st.session_state: st.session_state.grade             = None
if "subject"           not in st.session_state: st.session_state.subject           = None

# Teacher
if "period_blocks"     not in st.session_state: st.session_state.period_blocks     = [{"id": 0, "duration": None, "count": None}]
if "next_block_id"     not in st.session_state: st.session_state.next_block_id     = 1
if "teacher_generated" not in st.session_state: st.session_state.teacher_generated = False
if "teacher_ch_idx"    not in st.session_state: st.session_state.teacher_ch_idx    = None

# Principal
if "principal_period_blocks"  not in st.session_state: st.session_state.principal_period_blocks  = [{"id": 0, "duration": None, "count": None}]
if "principal_next_block_id"  not in st.session_state: st.session_state.principal_next_block_id  = 1
# Always sync ch_selected / ch_periods with the live chapter list so that
# newly-added chapters (e.g. English after mappings are generated) appear
# immediately without requiring a full session restart.
if "ch_selected" not in st.session_state:
    st.session_state.ch_selected = {ch["chapter_number"]: False for ch in chapters}
else:
    for ch in chapters:
        st.session_state.ch_selected.setdefault(ch["chapter_number"], False)

if "ch_periods" not in st.session_state:
    st.session_state.ch_periods = {ch["chapter_number"]: 6 for ch in chapters}
else:
    for ch in chapters:
        st.session_state.ch_periods.setdefault(ch["chapter_number"], 6)
if "principal_generated"      not in st.session_state: st.session_state.principal_generated      = False
if "ask_aruvi_open"           not in st.session_state: st.session_state.ask_aruvi_open           = False
st.session_state.setdefault("ask_aruvi_session_id",   str(uuid.uuid4()))
st.session_state.setdefault("ask_aruvi_category",     None)
st.session_state.setdefault("ask_aruvi_response",     "")
st.session_state.setdefault("ask_aruvi_last_query",   "")
st.session_state.setdefault("ask_aruvi_show_thumbs",  False)
st.session_state.setdefault("ask_aruvi_thumb_done",   False)
st.session_state.setdefault("ask_aruvi_show_followup", False)
st.session_state.setdefault("ask_aruvi_detail_cat", None)
st.session_state.setdefault("ask_aruvi_fb_sent",    False)
st.session_state.setdefault("ask_aruvi_fb_reset",   0)
# Managed-agent secondary panel state
st.session_state.setdefault("ask_aruvi_agent_open",         False)
st.session_state.setdefault("ask_aruvi_agent_response",     "")
st.session_state.setdefault("ask_aruvi_agent_last_query",   "")
st.session_state.setdefault("ask_aruvi_agent_show_thumbs",  False)
st.session_state.setdefault("ask_aruvi_agent_thumb_done",   False)
st.session_state.setdefault("ask_aruvi_agent_show_followup",False)
st.session_state.setdefault("ask_aruvi_agent_fb_sent",      False)
st.session_state.setdefault("ask_aruvi_agent_fb_reset",     0)
st.session_state.setdefault("ask_aruvi_agent_mode",         "question")  # "question" | "feedback"
st.session_state.setdefault("ask_aruvi_agent_fu_done",      False)  # thumbs-down follow-up submitted
st.session_state.setdefault("ask_aruvi_agent_fu_skipped",   False)  # True when Skip pressed (not Submit)
if "lpa_result"               not in st.session_state: st.session_state.lpa_result               = None
if "lpa_generating"           not in st.session_state: st.session_state.lpa_generating           = False
if "lpa_start_ts"             not in st.session_state: st.session_state.lpa_start_ts             = None
if "lpa_stop_event"           not in st.session_state: st.session_state.lpa_stop_event           = None
if "lpa_thread"               not in st.session_state: st.session_state.lpa_thread               = None
if "lpa_result_queue"         not in st.session_state: st.session_state.lpa_result_queue         = None
if "no_chapter_warning"       not in st.session_state: st.session_state.no_chapter_warning       = False
if "plan_just_saved"          not in st.session_state: st.session_state.plan_just_saved          = False
# LP/A split — Generate-tab confirmation dialog state
if "gen_include_assessment"        not in st.session_state: st.session_state.gen_include_assessment        = False
if "show_gen_confirm"              not in st.session_state: st.session_state.show_gen_confirm              = False
# LP/A split — deferred assessment (My Plans) state
if "mp_deferred_assess_generating" not in st.session_state: st.session_state.mp_deferred_assess_generating = False
if "mp_deferred_assess_plan"       not in st.session_state: st.session_state.mp_deferred_assess_plan       = None
if "mp_da_thread"                  not in st.session_state: st.session_state.mp_da_thread                  = None
if "mp_da_stop_event"              not in st.session_state: st.session_state.mp_da_stop_event              = None
if "mp_da_result_queue"            not in st.session_state: st.session_state.mp_da_result_queue            = None

@st.dialog(" ")
def _no_chapter_dialog():
    st.markdown(
        '<div style="text-align:center;padding:4px 0 8px;">'
        '<div style="font-size:2.2rem;margin-bottom:10px;">📖</div>'
        '<div style="font-size:1rem;font-weight:600;color:#3d3b38;margin-bottom:8px;">'
        'No chapter selected</div>'
        '<div style="font-size:0.85rem;color:#6b6965;margin-bottom:4px;">'
        'Please pick a chapter from the sidebar<br>before generating.'
        '</div></div>',
        unsafe_allow_html=True,
    )
    col = st.columns([1, 2, 1])[1]
    with col:
        if st.button("OK", key="no_chapter_ok_dlg", type="primary", use_container_width=True):
            st.session_state.no_chapter_warning = False
            st.rerun()
if "mp_viewing_plan"          not in st.session_state: st.session_state.mp_viewing_plan          = None
if "period_rows"              not in st.session_state: st.session_state["period_rows"]            = []
if "myplans_should_collapse"  not in st.session_state: st.session_state.myplans_should_collapse  = False
if "show_save_prompt"         not in st.session_state: st.session_state.show_save_prompt         = False
if "plan_already_saved"       not in st.session_state: st.session_state.plan_already_saved       = False

has_chapter_data = len(chapters) > 0

# ── Fixed top nav bar ─────────────────────────────────────────────────────────
# Logo/brand rendered as HTML; pill buttons rendered as CSS-positioned st.buttons.

logo_img_tag = (
    f'<img src="{LOGO_SRC}" alt="Aruvi logo">'
    if LOGO_SRC else '<div style="width:56px;height:56px;"></div>'
)

st.markdown(f"""
<div class="aruvi-topnav">
  <div class="topnav-left">
    {logo_img_tag}
    <div class="topnav-brand">
      <span class="topnav-wordmark">Aruvi</span>
      <span class="topnav-slogan">AI powered teaching assistant</span>
    </div>
  </div>
  <div class="topnav-center" id="aruvi-pill-anchor"></div>
  <div class="topnav-right"></div>
</div>
""", unsafe_allow_html=True)

_nc1, _nc2, _nc3, _nc4, _nc5 = st.columns([2, 1, 1, 1, 2])
with _nc2:
    if st.button("Allocate", key="nav_allocate", type="primary" if st.session_state.role == "Allocate" else "secondary"):
        st.session_state.role = "Allocate"
        st.query_params["role"] = "Allocate"
        st.rerun()
with _nc3:
    if st.button("Generate", key="nav_generate", type="primary" if st.session_state.role == "Generate" else "secondary"):
        st.session_state.role = "Generate"
        st.query_params["role"] = "Generate"
        st.rerun()
with _nc4:
    if st.button("My Plans", key="nav_myplans", type="primary" if st.session_state.role == "My Plans" else "secondary"):
        st.session_state.role = "My Plans"
        st.session_state.myplans_should_collapse = True
        st.query_params["role"] = "My Plans"
        st.rerun()


# ── Sidebar ───────────────────────────────────────────────────────────────────
# Change 3: Grade / Subject / Chapter selectboxes use label_visibility="visible".
#           CSS floats each label inside the selectbox border at top-left.
#           No separate icon-label-row div above each selectbox.

with st.sidebar:
    if st.session_state.role == "My Plans":
        # On the My Plans tab, the Grade and Subject selectboxes act as filters
        # for the saved-plans table. Unlike Allocate/Generate they expose an
        # explicit "All" option (the default), and they write through to the
        # shared session-state keys so any specific selection persists when
        # the user switches tabs. "All" maps to None in shared state.

        _s_icon = f'<img src="{SUBJECT_SRC}" class="field-icon" alt="">' if SUBJECT_SRC else ""
        st.markdown(
            f'<div class="sidebar-field-label">{_s_icon}'
            f'<span class="field-label-text">Subject</span></div>',
            unsafe_allow_html=True,
        )
        _mp_subj_opts = ["All"] + SUBJECTS
        _mp_subj_cur  = st.session_state.subject if st.session_state.subject in SUBJECTS else "All"
        subject = st.selectbox(
            "Subject",
            _mp_subj_opts,
            index=_mp_subj_opts.index(_mp_subj_cur),
            label_visibility="collapsed",
            key="mp_subject_select",
        )
        _subject_val = None if subject == "All" else subject
        if _subject_val != st.session_state.subject:
            st.session_state.subject             = _subject_val
            st.session_state.teacher_ch_idx      = None
            # Clear grade if it is not valid for the new subject
            if st.session_state.grade and _subject_val and \
                    st.session_state.grade not in grades_for_subject(_subject_val):
                st.session_state.grade = None
            st.session_state.teacher_generated   = False
            st.session_state.principal_generated = False
            if _subject_val:
                st.query_params["subject"] = _subject_val
            else:
                st.query_params.pop("subject", None)
            st.rerun()

        # Grade options depend on selected subject (All = no subject filter → show all grades)
        _g_icon = f'<img src="{GRADE_SRC}" class="field-icon-grade" alt="">' if GRADE_SRC else ""
        st.markdown(
            f'<div class="sidebar-field-label">{_g_icon}'
            f'<span class="field-label-text">Grade</span></div>',
            unsafe_allow_html=True,
        )
        _mp_base_grades = grades_for_subject(_subject_val) if _subject_val else GRADES
        _mp_grade_opts  = ["All"] + _mp_base_grades
        _mp_grade_cur   = st.session_state.grade if st.session_state.grade in _mp_base_grades else "All"
        grade = st.selectbox(
            "Grade",
            _mp_grade_opts,
            index=_mp_grade_opts.index(_mp_grade_cur),
            label_visibility="collapsed",
            key="mp_grade_select",
        )
        _grade_val = None if grade == "All" else grade
        if _grade_val != st.session_state.grade:
            st.session_state.grade               = _grade_val
            st.session_state.teacher_ch_idx      = None
            st.session_state.teacher_generated   = False
            st.session_state.principal_generated = False
            if _grade_val:
                st.query_params["grade"] = _grade_val
            else:
                st.query_params.pop("grade", None)
            st.rerun()

        # ── Saved date filter ─────────────────────────────────────────────────
        _sv_icon = f'<img src="{SAVED_SRC}" class="field-icon" alt="">' if SAVED_SRC else ""
        st.markdown(
            f'<div class="sidebar-field-label">{_sv_icon}'
            f'<span class="field-label-text">Saved</span></div>',
            unsafe_allow_html=True,
        )
        _mp_saved_opts = ["Today", "Yesterday", "This week", "This month", "All"]
        if "mp_saved_filter" not in st.session_state:
            st.session_state.mp_saved_filter = "All"
        _saved_choice = st.selectbox(
            "Saved",
            _mp_saved_opts,
            index=_mp_saved_opts.index(st.session_state.mp_saved_filter),
            label_visibility="collapsed",
            key="mp_saved_select",
        )
        if _saved_choice != st.session_state.mp_saved_filter:
            st.session_state.mp_saved_filter = _saved_choice
            st.rerun()

        st.markdown('<div class="sidebar-spacer"></div>', unsafe_allow_html=True)
        st.markdown("""
    <div class="sidebar-user-footer">
      <hr style="border:none;border-top:1px solid #d9d6d0;margin:0;" />
      <div class="user-footer-inner">
        <div class="user-avatar">RT</div>
        <div class="user-info">
          <span class="user-name">Ramesh Tripathi</span>
          <span class="user-plan">Free plan</span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)
    else:

        # ── Subject selector — label above, value below (left-aligned) ───────────
        _s_icon = f'<img src="{SUBJECT_SRC}" class="field-icon" alt="">' if SUBJECT_SRC else ""
        st.markdown(
            f'<div class="sidebar-field-label">{_s_icon}'
            f'<span class="field-label-text">Subject</span></div>',
            unsafe_allow_html=True,
        )
        subject = st.selectbox(
            "Subject",
            SUBJECTS,
            index=None if st.session_state.subject is None
                  else SUBJECTS.index(st.session_state.subject),
            placeholder="Choose a subject",
            label_visibility="collapsed",
            key="subject_select",
        )
        if subject != st.session_state.subject:
            st.session_state.subject             = subject
            st.session_state.teacher_ch_idx      = None
            # If current grade is not valid for the new subject, clear it
            if st.session_state.grade and subject and \
                    st.session_state.grade not in grades_for_subject(subject):
                st.session_state.grade = None
            st.session_state.teacher_generated   = False
            st.session_state.principal_generated = False
            if subject:
                st.query_params["subject"] = subject
            st.rerun()

        # ── Grade selector — label above, value below (left-aligned) ─────────────
        # Available grades depend on the selected subject.
        _g_icon = f'<img src="{GRADE_SRC}" class="field-icon-grade" alt="">' if GRADE_SRC else ""
        st.markdown(
            f'<div class="sidebar-field-label">{_g_icon}'
            f'<span class="field-label-text">Grade</span></div>',
            unsafe_allow_html=True,
        )
        _avail_grades = grades_for_subject(st.session_state.subject) if st.session_state.subject else GRADES
        _cur_grade    = st.session_state.grade if st.session_state.grade in _avail_grades else None
        grade = st.selectbox(
            "Grade",
            _avail_grades,
            index=None if _cur_grade is None
                  else _avail_grades.index(_cur_grade),
            placeholder="Choose a grade",
            label_visibility="collapsed",
            key="grade_select",
        )
        if grade != st.session_state.grade:
            st.session_state.grade               = grade
            st.session_state.teacher_ch_idx      = None
            st.session_state.teacher_generated   = False
            st.session_state.principal_generated = False
            if grade:
                st.query_params["grade"] = grade
            st.rerun()

        # ── No data for this combination ──────────────────────────────────────────
        if not has_chapter_data:
            st.markdown(
                '<div class="no-data-notice">'
                f'Chapter data for {st.session_state.subject}, '
                f'{st.session_state.grade} is not available yet.'
                '</div>',
                unsafe_allow_html=True,
            )

        # ── Teacher inputs ────────────────────────────────────────────────────────
        elif st.session_state.role == "Generate":

            st.divider()

            # Chapter selector — label above, value below (left-aligned)
            _c_icon = f'<img src="{CHAPTER_SRC}" class="field-icon" alt="">' if CHAPTER_SRC else ""
            st.markdown(
                f'<div class="sidebar-field-label">{_c_icon}'
                f'<span class="field-label-text">Chapter</span></div>',
                unsafe_allow_html=True,
            )
            ch_labels = [ch_label(ch) for ch in chapters]
            # Guard: reset index if it is out of range for the current chapter list
            if st.session_state.teacher_ch_idx is not None and (
                st.session_state.teacher_ch_idx < 0
                or st.session_state.teacher_ch_idx >= len(ch_labels)
            ):
                st.session_state.teacher_ch_idx = None
            sel_label = st.selectbox(
                "Chapter",
                ch_labels,
                index=st.session_state.teacher_ch_idx,
                placeholder="Choose a chapter",
                label_visibility="collapsed",
                key="teacher_ch_select",
            )
            if sel_label is not None:
                new_idx = ch_labels.index(sel_label)
                if new_idx != st.session_state.teacher_ch_idx:
                    st.session_state.teacher_ch_idx    = new_idx
                    st.session_state.teacher_generated = False
                    st.query_params["ch"] = str(new_idx)

            st.divider()

            # ── Available Time section ─────────────────────────────────────────────
            # "Available time" label — uses time.png icon
            if TIME_SRC:
                time_icon_html = f'<img src="{TIME_SRC}" class="period-icon" alt="⏱">'
            else:
                time_icon_html = ''
            st.markdown(
                f'<div class="sect-label">{time_icon_html}'
                f'<span>Available time</span></div>',
                unsafe_allow_html=True,
            )

            # ── Multi-row period state bootstrap ──────────────────────────────────
            if not st.session_state.get("period_rows"):
                st.session_state["period_rows"] = [0]
                st.session_state["_next_row_id"] = 1

            # Ensure cnt is initialised for every active row (new rows only)
            for _rid in st.session_state.get("period_rows", []):
                if f"cnt_{_rid}" not in st.session_state:
                    st.session_state[f"cnt_{_rid}"] = 1

            # Inject dynamic CSS: add-row button shows period.png icon
            if PERIOD_SRC:
                st.markdown(
                    f"""<style>
section[data-testid="stSidebar"] div[class*="st-key-add_period_row"] button {{
    background: url('{PERIOD_SRC}') center / 16px 16px no-repeat transparent !important;
    color: transparent !important;
    border: none !important;
    box-shadow: none !important;
    padding: 0 !important;
    min-height: 22px !important;
}}
section[data-testid="stSidebar"] div[class*="st-key-add_period_row"] button:hover {{
    opacity: 0.60;
    background: url('{PERIOD_SRC}') center / 16px 16px no-repeat transparent !important;
}}
</style>""",
                unsafe_allow_html=True,
            )

            # ── Column header labels — rendered once above the first row ──────────
            _hc_dur, _hc_cnt, _hc_add = st.columns([4, 4, 1])
            with _hc_dur:
                st.markdown('<div style="font-size:0.68rem;font-weight:500;letter-spacing:0.08em;text-transform:uppercase;color:#5a5754;margin-bottom:0.15rem;">Mins / Period</div>', unsafe_allow_html=True)
            with _hc_cnt:
                st.markdown('<div style="font-size:0.68rem;font-weight:500;letter-spacing:0.08em;text-transform:uppercase;color:#5a5754;margin-bottom:0.15rem;">No. of Periods</div>', unsafe_allow_html=True)
            with _hc_add:
                st.button(
                    "⊕",
                    key="add_period_row",
                    use_container_width=True,
                    help="Add another period type",
                    on_click=_cb_add_row,
                )

            # ── Period rows — all [4, 4, 1]; first row's delete slot stays empty ──
            for _rid in st.session_state.get("period_rows", []):
                _is_first = (_rid == st.session_state.get("period_rows", [0])[0])
                c_dur, c_cnt, c_del = st.columns([4, 4, 1])

                with c_dur:
                    st.selectbox("Time per period", options=DURATION_OPTIONS, index=DURATION_OPTIONS.index(40),
                                 label_visibility="collapsed", key=f"dur_sel_{_rid}")

                with c_cnt:
                    cm, cv, cp = st.columns([1, 3, 1])
                    with cm:
                        st.button("−", key=f"minus_{_rid}",
                                  use_container_width=True,
                                  on_click=_cb_inc_cnt, args=(_rid, -1))
                    with cv:
                        st.number_input(
                            "count",
                            min_value=1,
                            max_value=999,
                            step=1,
                            label_visibility="collapsed",
                            key=f"cnt_{_rid}",
                        )
                    with cp:
                        st.button("+", key=f"plus_{_rid}",
                                  use_container_width=True,
                                  on_click=_cb_inc_cnt, args=(_rid, 1))

                with c_del:
                    if not _is_first:
                        st.button("×", key=f"del_{_rid}",
                                  use_container_width=True,
                                  on_click=_cb_del_row, args=(_rid,))

            # Total across all rows
            total_m = sum(
                (st.session_state.get(f"dur_sel_{r}") or 0) * (st.session_state.get(f"cnt_{r}") or 0)
                for r in st.session_state.get("period_rows", [])
            )
            total_p = sum(
                (st.session_state.get(f"cnt_{r}") or 0)
                for r in st.session_state.get("period_rows", [])
            )
            if total_m > 0:
                _h, _min = divmod(total_m, 60)
                if _h == 0:
                    _time_str = f"{_min} minute{'s' if _min != 1 else ''}"
                elif _min == 0:
                    _time_str = f"{_h} hour{'s' if _h != 1 else ''}"
                else:
                    _time_str = f"{_h} hour{'s' if _h != 1 else ''} and {_min} minute{'s' if _min != 1 else ''}"
                _p_label = f"{total_p} period{'s' if total_p != 1 else ''}"
                st.markdown(
                    f'<div style="font-size:0.79rem;color:#3d3b38;margin:0.4rem 0 0.25rem 0;">'
                    f'Total · {_time_str}, {_p_label}'
                    f'</div>',
                    unsafe_allow_html=True,
                )

            st.divider()

            can_gen = (
                len(st.session_state.get("period_rows", [])) > 0
                and all(
                    (st.session_state.get(f"dur_sel_{r}") or 0) >= 1
                    and (st.session_state.get(f"cnt_{r}") or 0) >= 1
                    for r in st.session_state.get("period_rows", [])
                )
            )
            if st.button(
                "Generate Lesson Plan",
                disabled=not can_gen,
                type="primary",
                use_container_width=True,
                key="teacher_gen",
            ):
                if st.session_state.teacher_ch_idx is None:
                    st.session_state.no_chapter_warning = True
                else:
                    st.session_state.no_chapter_warning = False
                    # LP/A split — open the confirmation dialog before kicking
                    # off generation. The checkbox in the dialog decides whether
                    # assessment is included in the same run or deferred.
                    st.session_state.show_gen_confirm   = True
                    st.session_state.plan_already_saved = False
                    st.rerun()

            # ── No-chapter warning popup ──────────────────────────────────────────
            # The Include-Assessment confirmation card lives in the main
            # workspace (right of the sidebar) — see the Generate role block.
            if st.session_state.get("no_chapter_warning"):
                _no_chapter_dialog()

        # ── Principal inputs ──────────────────────────────────────────────────────
        else:

            # ── Period Budget section — same block architecture as Teacher ─────────
            if FULL_PERIOD_SRC:
                fp_icon_html = f'<img src="{FULL_PERIOD_SRC}" class="period-icon" alt="">'
            else:
                fp_icon_html = ''
            st.markdown(
                f'<div class="sect-label" style="margin-bottom:0.45rem;">{fp_icon_html}'
                f'<span>Period Budget</span></div>',
                unsafe_allow_html=True,
            )

            # ── Multi-row period state bootstrap (Principal) ──────────────────────
            if "period_rows_p" not in st.session_state:
                st.session_state["period_rows_p"] = [0]
                st.session_state["_next_row_id_p"] = 1

            # Ensure cnt is initialised (default 1) for every active row (new rows only)
            for _rid_p in st.session_state["period_rows_p"]:
                if f"cnt_p{_rid_p}" not in st.session_state:
                    st.session_state[f"cnt_p{_rid_p}"] = 1

            # Dynamic CSS: ⊕ add-row button shows period.png icon (Plan)
            if PERIOD_SRC:
                st.markdown(
                    f"""<style>
    section[data-testid="stSidebar"] div[class*="st-key-add_period_row_p"] button {{
        background: url('{PERIOD_SRC}') center / 16px 16px no-repeat transparent !important;
        color: transparent !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
        min-height: 22px !important;
    }}
    section[data-testid="stSidebar"] div[class*="st-key-add_period_row_p"] button:hover {{
        opacity: 0.60;
        background: url('{PERIOD_SRC}') center / 16px 16px no-repeat transparent !important;
    }}
    </style>""",
                    unsafe_allow_html=True,
                )

            # ── Column header labels — rendered once above the first row ──────────
            _hc_dur_p, _hc_cnt_p, _hc_add_p = st.columns([4, 4, 1])
            with _hc_dur_p:
                st.markdown('<div style="font-size:0.68rem;font-weight:500;letter-spacing:0.08em;text-transform:uppercase;color:#5a5754;margin-bottom:0.15rem;">Mins / Period</div>', unsafe_allow_html=True)
            with _hc_cnt_p:
                st.markdown('<div style="font-size:0.68rem;font-weight:500;letter-spacing:0.08em;text-transform:uppercase;color:#5a5754;margin-bottom:0.15rem;">No. of Periods</div>', unsafe_allow_html=True)
            with _hc_add_p:
                st.button(
                    "⊕",
                    key="add_period_row_p",
                    use_container_width=True,
                    help="Add another period type",
                    on_click=_cb_add_row_p,
                )

            # ── Period rows — all [4, 4, 1]; first row's delete slot stays empty ──
            for _rid_p in st.session_state["period_rows_p"]:
                _is_first_p = (_rid_p == st.session_state["period_rows_p"][0])
                pc_dur, pc_cnt, pc_del = st.columns([4, 4, 1])

                with pc_dur:
                    st.selectbox("Time per period", options=DURATION_OPTIONS, index=DURATION_OPTIONS.index(40),
                                 label_visibility="collapsed", key=f"dur_sel_p{_rid_p}")

                with pc_cnt:
                    st.number_input(
                        "count",
                        min_value=1,
                        max_value=999,
                        step=1,
                        label_visibility="collapsed",
                        key=f"cnt_p{_rid_p}",
                    )

                with pc_del:
                    if not _is_first_p:
                        st.button("×", key=f"del_p{_rid_p}",
                                  use_container_width=True,
                                  on_click=_cb_del_row_p, args=(_rid_p,))

            # Total across all rows
            p_total_m = sum(
                (st.session_state.get(f"dur_sel_p{r}") or 0) * (st.session_state.get(f"cnt_p{r}") or 0)
                for r in st.session_state["period_rows_p"]
            )
            p_total_p = sum(
                (st.session_state.get(f"cnt_p{r}") or 0)
                for r in st.session_state["period_rows_p"]
            )
            if p_total_m > 0:
                _ph, _pmin = divmod(p_total_m, 60)
                if _ph == 0:
                    _p_time_str = f"{_pmin} minute{'s' if _pmin != 1 else ''}"
                elif _pmin == 0:
                    _p_time_str = f"{_ph} hour{'s' if _ph != 1 else ''}"
                else:
                    _p_time_str = f"{_ph} hour{'s' if _ph != 1 else ''} and {_pmin} minute{'s' if _pmin != 1 else ''}"
                _pp_label = f"{p_total_p} period{'s' if p_total_p != 1 else ''}"
                st.markdown(
                    f'<div style="font-size:0.79rem;color:#3d3b38;margin:0.4rem 0 0.25rem 0;">'
                    f'Total · {_p_time_str}, {_pp_label}'
                    f'</div>',
                    unsafe_allow_html=True,
                )

        # ── Sidebar spacer + user footer (sticky at bottom) ───────────────────────
        st.markdown('<div class="sidebar-spacer"></div>', unsafe_allow_html=True)
        st.markdown("""
    <div class="sidebar-user-footer">
      <hr style="border:none;border-top:1px solid #d9d6d0;margin:0;" />
      <div class="user-footer-inner">
        <div class="user-avatar">RT</div>
        <div class="user-info">
          <span class="user-name">Ramesh Tripathi</span>
          <span class="user-plan">Free plan</span>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)


# ── Workspace ─────────────────────────────────────────────────────────────────

if (not has_chapter_data and st.session_state.role != "My Plans"
        and st.session_state.lpa_result is None
        and not st.session_state.get("mp_deferred_assess_generating")):
    if st.session_state.grade is None or st.session_state.subject is None:
        _msg = "Choose a grade and subject to get started."
    else:
        _msg = f"No content available yet for {st.session_state.subject}, {st.session_state.grade}."
    st.markdown(
        f'<div class="ws-placeholder">{_msg}</div>',
        unsafe_allow_html=True,
    )

# ═════════════════════════════════════════════════
#  GENERATE WORKSPACE
#  Change 2: tabs = Competencies · Lesson Plan · Assessment
# ═════════════════════════════════════════════════
elif st.session_state.role == "Generate":

    # ── Deferred assessment generation ────────────────────────────────────────
    # Triggered when "Generate Assessment" is clicked on an lp_only row in
    # My Plans. The role flip there lands the teacher here so the popup
    # appears in the regular Generate-tab spot. On completion, role flips
    # back to My Plans so the now-PDF row is visible.
    if (st.session_state.get("mp_deferred_assess_generating")
        and st.session_state.get("mp_deferred_assess_plan") is not None):

        # Render a blank full-height placeholder so no stale DOM from the
        # previous My Plans render bleeds through as a shadow behind the popup.
        st.markdown(
            '<div style="min-height:80vh;"></div>',
            unsafe_allow_html=True,
        )

        _dap = st.session_state.mp_deferred_assess_plan

        _existing_da_thread = st.session_state.get("mp_da_thread")
        if (_existing_da_thread is not None and _existing_da_thread.is_alive()
                and st.session_state.mp_da_result_queue is not None):
            _da_stop_ev = st.session_state.mp_da_stop_event
            _da_rq      = st.session_state.mp_da_result_queue
        else:
            _da_stop_ev = threading.Event()
            _da_rq      = queue.Queue()
            st.session_state.mp_da_stop_event   = _da_stop_ev
            st.session_state.mp_da_result_queue = _da_rq
            _da_t = threading.Thread(
                target=generate_assessment_only,
                kwargs=dict(
                    saved_plan   = _dap,
                    result_queue = _da_rq,
                    stop_event   = _da_stop_ev,
                ),
                daemon=True,
            )
            add_script_run_ctx(_da_t)
            _da_t.start()
            st.session_state.mp_da_thread = _da_t

        # Hidden Stop button — the visible pill inside the popup clicks this
        # via DOM querySelector (see _hdr_working in generate_assessment_only).
        st.markdown(
            '<style>'
            'div[class*="st-key-btn_stop_da_generation"]{display:none!important;}'
            '</style>',
            unsafe_allow_html=True,
        )
        if st.button("stop", key="btn_stop_da_generation"):
            if st.session_state.mp_da_stop_event is not None:
                st.session_state.mp_da_stop_event.set()

        _da_heartbeat = st.empty()
        _da_result = None
        while _da_result is None:
            try:
                _da_result = _da_rq.get(timeout=0.25)
            except queue.Empty:
                _da_heartbeat.markdown("")
        _da_heartbeat.empty()

        st.session_state.mp_da_thread       = None
        st.session_state.mp_da_stop_event   = None
        st.session_state.mp_da_result_queue = None

        if (not _da_result.get("stopped")
            and not _da_result.get("error")
            and _da_result.get("assessment_items")):
            update_saved_plan_with_assessment(
                filename      = _dap.get("filename", ""),
                grade         = _dap.get("grade", ""),
                subject       = _dap.get("subject", ""),
                assess_result = _da_result,
            )
        elif _da_result.get("error"):
            st.error(f"Assessment generation failed: {_da_result['error']}")

        st.session_state.mp_deferred_assess_generating = False
        st.session_state.mp_deferred_assess_plan       = None
        # Hand control back to My Plans so the row re-renders with PDF ⬇.
        st.session_state.role = "My Plans"
        st.query_params["role"] = "My Plans"
        st.rerun()

    # ── Confirmation card: Include Assessment? ──────────────────────────────
    # Rendered in the main workspace (not the sidebar) so the card sits over
    # the body content. Layout: stacked rows (Grade / Subject / Chapter title)
    # then the Include-Assessment toggle, then the action buttons.
    if (st.session_state.get("show_gen_confirm")
        and not st.session_state.lpa_generating
        and st.session_state.lpa_result is None):
        _confirm_ch = None
        if st.session_state.teacher_ch_idx is not None and st.session_state.teacher_ch_idx < len(chapters):
            _confirm_ch = chapters[st.session_state.teacher_ch_idx]
        st.markdown("""<style>
div[class*="st-key-gen_confirm_box"] {
    background:#fff;
    border:1px solid #d9d6d0;
    border-radius:12px;
    padding:1.25rem 1.5rem;
    margin:0.5rem 0 1rem 0;
    box-shadow:0 2px 8px rgba(0,0,0,0.08);
    max-width:540px;
}
/* Generate button — match the sidebar Generate button (dark slate pill with
   the ✦ AI twinkle prepended via ::before). */
div[class*="st-key-gen_confirm_go"] button {
    height: 56px !important;
    min-height: 56px !important;
    border-radius: 12px !important;
    background: #2c3e50 !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.18) !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    color: #ffffff !important;
    letter-spacing: 0.02em !important;
    border: none !important;
    justify-content: center !important;
}
div[class*="st-key-gen_confirm_go"] button:hover {
    background: #3d5166 !important;
    box-shadow: 0 4px 14px rgba(0,0,0,0.24) !important;
}
div[class*="st-key-gen_confirm_go"] button * {
    color: #ffffff !important;
    visibility: visible !important;
}
div[class*="st-key-gen_confirm_go"] button::before {
    content: "✦";
    font-size: 0.85rem;
    color: #ffffff;
    flex-shrink: 0;
    visibility: visible !important;
    margin-right: 0.35rem;
}
/* Cancel button — white text on dark grey, font size matches Generate. */
div[class*="st-key-gen_confirm_cancel"] button {
    height: 56px !important;
    min-height: 56px !important;
    border-radius: 12px !important;
    background: #5a5754 !important;
    color: #ffffff !important;
    border: none !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.18) !important;
}
div[class*="st-key-gen_confirm_cancel"] button:hover {
    background: #6e6b67 !important;
}
div[class*="st-key-gen_confirm_cancel"] button * {
    color: #ffffff !important;
}
</style>""", unsafe_allow_html=True)
        with st.container(key="gen_confirm_box"):
            st.markdown(
                '<div style="font-size:1.05rem;font-weight:600;color:#3d3b38;'
                'margin-bottom:0.85rem;">Ready to generate</div>',
                unsafe_allow_html=True,
            )
            _meta_row_style = (
                'font-size:0.88rem;color:#3d3b38;margin:0.15rem 0;'
                'display:flex;gap:0.5rem;'
            )
            _meta_label_style = (
                'font-size:0.72rem;text-transform:uppercase;letter-spacing:0.06em;'
                'color:#7a776f;width:5.5rem;flex-shrink:0;align-self:center;'
            )
            _grade_val   = st.session_state.grade   or ""
            _subject_val = st.session_state.subject or ""
            _chapter_val = _confirm_ch.get("chapter_title", "") if _confirm_ch else ""
            st.markdown(
                f'<div style="{_meta_row_style}">'
                f'<span style="{_meta_label_style}">Grade</span>'
                f'<span>{_grade_val}</span></div>'
                f'<div style="{_meta_row_style}">'
                f'<span style="{_meta_label_style}">Subject</span>'
                f'<span>{_subject_val}</span></div>'
                f'<div style="{_meta_row_style}">'
                f'<span style="{_meta_label_style}">Chapter</span>'
                f'<span>{_chapter_val}</span></div>',
                unsafe_allow_html=True,
            )
            st.markdown('<div style="height:0.75rem;"></div>', unsafe_allow_html=True)
            st.checkbox(
                "Include Assessment",
                key="gen_confirm_include_assess",
                value=False,
                help="Assessment can also be generated later from My Plans.",
            )
            st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)
            _gc1, _gc2, _gc_rest = st.columns([1.3, 1.3, 2])
            with _gc1:
                if st.button("Generate", key="gen_confirm_go",
                             type="primary", use_container_width=True):
                    st.session_state.gen_include_assessment = bool(
                        st.session_state.get("gen_confirm_include_assess", False)
                    )
                    st.session_state.lpa_generating  = True
                    st.session_state.lpa_result      = None
                    st.session_state.show_gen_confirm = False
                    st.rerun()
            with _gc2:
                if st.button("Cancel", key="gen_confirm_cancel",
                             use_container_width=True):
                    st.session_state.show_gen_confirm = False
                    st.rerun()

    # ── Generation (needs chapter selected) ──────────────────────────────────
    if st.session_state.lpa_generating and st.session_state.teacher_ch_idx is not None and st.session_state.teacher_ch_idx < len(chapters):
        if st.session_state.teacher_ch_idx >= len(chapters):
            st.session_state.lpa_generating = False
            st.session_state.teacher_ch_idx = None
            st.rerun()
        selected_ch = chapters[st.session_state.teacher_ch_idx]
        if st.session_state.lpa_generating:

            # ── Launch background thread (only once per generate run) ────────
            # generate_lpa calls st.* inside the thread for progress rendering.
            # add_script_run_ctx propagates the current Streamlit script context
            # to the thread so those calls are valid and non-silent.
            #
            # IMPORTANT: when the user clicks Stop, Streamlit triggers a rerun
            # of THIS script block. We must NOT spawn a second worker thread on
            # the rerun — instead, we re-bind to the existing thread/queue/
            # event already stored in session state.
            _existing_thread = st.session_state.get("lpa_thread")
            if (_existing_thread is not None and _existing_thread.is_alive()
                and st.session_state.lpa_result_queue is not None):
                # Mid-run rerun — reuse the already-running worker.
                _stop_ev = st.session_state.lpa_stop_event
                _rq      = st.session_state.lpa_result_queue
            else:
                # Fresh generation — spawn the worker.
                _stop_ev = threading.Event()
                _rq      = queue.Queue()
                st.session_state.lpa_stop_event   = _stop_ev
                st.session_state.lpa_result_queue = _rq
                _t = threading.Thread(
                    target=generate_lp_only,
                    kwargs=dict(
                        grade              = st.session_state.grade,
                        subject            = st.session_state.subject,
                        chapter            = selected_ch,
                        period_rows        = st.session_state.get("period_rows", [0]),
                        session            = st.session_state,
                        include_assessment = st.session_state.get("gen_include_assessment", False),
                        result_queue       = _rq,
                        stop_event         = _stop_ev,
                    ),
                    daemon=True,
                )
                add_script_run_ctx(_t)   # allow st.* calls from the thread
                _t.start()
                st.session_state.lpa_thread = _t

            # ── Hidden Streamlit stop button ──────────────────────────────────
            # The visible stop pill in the progress popup clicks this button
            # directly via DOM querySelector (see the pill HTML above), which
            # triggers a Streamlit rerun where the click is registered here.
            st.markdown(
                '<style>'
                'div[class*="st-key-btn_stop_generation"]{display:none!important;}'
                '</style>',
                unsafe_allow_html=True,
            )
            if st.button("stop", key="btn_stop_generation"):
                if st.session_state.lpa_stop_event is not None:
                    st.session_state.lpa_stop_event.set()

            # ── Poll the result queue with an st.* heartbeat ─────────────────
            # A naked _rq.get() blocks the script-runner thread, so Streamlit
            # can never inject a rerun in response to the Stop click. The
            # heartbeat below is a no-op st.* call that gives Streamlit a
            # checkpoint to raise RerunException when a rerun is queued.
            _heartbeat = st.empty()
            result = None
            while result is None:
                try:
                    result = _rq.get(timeout=0.25)
                except queue.Empty:
                    _heartbeat.markdown("")
            _heartbeat.empty()

            st.session_state.lpa_thread       = None
            st.session_state.lpa_stop_event   = None
            st.session_state.lpa_result_queue = None

            if result.get("stopped"):
                # User stopped — silently reset to state before Generate was pressed
                st.session_state.lpa_generating = False
                st.session_state.lpa_result     = None
                st.rerun()
            else:
                st.session_state.lpa_result        = result
                st.session_state.lpa_generating    = False
                st.session_state.teacher_generated = True
                st.session_state.show_save_prompt  = True
                st.session_state.grade             = None
                st.session_state.subject           = None
                st.session_state.period_rows       = []
                st.rerun()

    # ── Result block ─────────────────────────────────────────────────────────
    result = st.session_state.lpa_result
    if result is None and st.session_state.teacher_ch_idx is None:
        st.markdown(
            '<div class="ws-placeholder">Choose a chapter to get started, '
            'or view a saved plan from My Plans.</div>',
            unsafe_allow_html=True,
        )
    elif result is None:
        # Skip the placeholder while the confirmation card is on screen —
        # the card already occupies the workspace.
        if not st.session_state.get("show_gen_confirm"):
            st.markdown(
                '<div class="ws-placeholder">'
                'Set your period budget and click Generate Lesson Plan.'
                '</div>',
                unsafe_allow_html=True,
            )
    elif result.get("error"):
        st.error(f"Generation failed: {result['error']}")
    else:
        # Get chapter data — from index if available, else use result metadata
        if st.session_state.teacher_ch_idx is not None and st.session_state.teacher_ch_idx < len(chapters):
            _chapter_export = chapters[st.session_state.teacher_ch_idx]
        else:
            # result now carries grade/subject/chapter_title/chapter_number from generate_lpa
            _ch_num_from_result = result.get("chapter_number")
            _chapter_export = next(
                (c for c in chapters if c["chapter_number"] == _ch_num_from_result),
                None
            )
            if _chapter_export is None:
                # Fallback: also try lo_handoff for backwards compat with old saved results
                _lo_list = result.get("lo_handoff", [])
                _ch_num_from_lo = _lo_list[0].get("chapter_number") if _lo_list else None
                _chapter_export = next(
                    (c for c in chapters if c["chapter_number"] == _ch_num_from_lo),
                    None
                )
            if _chapter_export is None:
                # Reconstruct minimal chapter dict from result metadata
                _chapter_export = {
                    "chapter_title":   result.get("chapter_title", "Chapter"),
                    "chapter_number":  result.get("chapter_number") or 0,
                    "chapter_weight":  "",
                    "primary":         [],
                }
        _safe_title = re.sub(r"[^\w\s-]", "", _chapter_export.get("chapter_title", "chapter")).strip().replace(" ", "_")[:40]
        _filename_stem = f"Aruvi_{_safe_title}"

        # ── Resolve grade / subject for PDF and save operations ───────────────
        # After generation, session grade/subject are cleared; use result's own copy.
        _res_grade   = st.session_state.grade   or result.get("grade",   "Grade VII")
        _res_subject = st.session_state.subject or result.get("subject", "Social Science")

        # ── "Do you want to save the plan?" popup — shown once after generation ─
        if st.session_state.get("show_save_prompt"):
            st.markdown("""<style>
div[class*="st-key-save_prompt_box"] {
    background:#fff;
    border:1px solid #d9d6d0;
    border-radius:10px;
    padding:1rem 1.25rem;
    margin-bottom:1rem;
    box-shadow:0 2px 8px rgba(0,0,0,0.08);
}
div[class*="st-key-save_prompt_yes"] button {
    background-color:#2c3e50 !important;
    color:#fff !important;
    border:none !important;
}
div[class*="st-key-save_prompt_no"] button {
    background-color:#f2f0ec !important;
    color:#3d3b38 !important;
    border:1px solid #d9d6d0 !important;
}
</style>""", unsafe_allow_html=True)
            with st.container(key="save_prompt_box"):
                st.markdown(
                    '<div style="font-size:0.95rem;font-weight:500;color:#3d3b38;'
                    'margin-bottom:0.75rem;">Do you want to save the plan?</div>',
                    unsafe_allow_html=True,
                )
                _sp_c1, _sp_c2, _sp_rest = st.columns([1, 1, 3])
                with _sp_c1:
                    if st.button("Yes", key="save_prompt_yes", type="primary",
                                 use_container_width=True):
                        save_plan(
                            grade       = _res_grade,
                            subject     = _res_subject,
                            chapter     = _chapter_export,
                            period_rows = st.session_state.get("period_rows_snapshot",
                                          st.session_state.get("period_rows", [])),
                            session     = st.session_state,
                            result      = result,
                        )
                        st.session_state.show_save_prompt  = False
                        st.session_state.plan_just_saved   = True
                        st.session_state.plan_already_saved = True
                        st.rerun()
                with _sp_c2:
                    if st.button("No", key="save_prompt_no",
                                 use_container_width=True):
                        st.session_state.show_save_prompt = False
                        st.rerun()

        # ── Primary-style LP / Assessment / Save / Clear buttons ─────────────
        # CSS: match Generate button colour scheme; orange for Save button;
        #      Clear uses Streamlit primary style (same as LP / Assessment).
        st.markdown("""<style>
div[data-testid="stDownloadButton"] button[kind="primary"] {
    font-size: 0.82rem !important;
}
div[class*="st-key-gen_save_top"] button,
div[class*="st-key-gen-save-top"] button {
    background-color: #E87722 !important;
    color: #ffffff !important;
    border: none !important;
    font-size: 0.82rem !important;
}
div[class*="st-key-gen_clear_top"] button,
div[class*="st-key-gen-clear-top"] button {
    background-color: #1e2a38 !important;
    color: #ffffff !important;
    border: none !important;
    font-size: 0.82rem !important;
}
div[class*="st-key-gen_clear_top"] button:hover,
div[class*="st-key-gen-clear-top"] button:hover {
    background-color: #2c3e52 !important;
    color: #ffffff !important;
    border: none !important;
}
div[class*="st-key-gen_clear_top"] button p,
div[class*="st-key-gen-clear-top"] button p {
    color: #ffffff !important;
}
</style>""", unsafe_allow_html=True)
        _pdl_c1, _pdl_c2, _pdl_c3, _pdl_c4, _pdl_spc = st.columns([1, 1, 1, 1, 1])
        with _pdl_c1:
            try:
                from lp_pdf_generator import build_lp_pdf_bytes as _blpb_gen
                _gen_lp_payload = {
                    "saved_at":       datetime.now().isoformat(timespec="seconds"),
                    "grade":          _res_grade,
                    "subject":        _res_subject,
                    "chapter_number": _chapter_export.get("chapter_number", 0),
                    "chapter_title":  _chapter_export.get("chapter_title",  ""),
                    "result":         {"lesson_plan": result.get("lesson_plan", {})},
                }
                _gen_lp_bytes = _blpb_gen(_gen_lp_payload)
                st.download_button(
                    label="Lesson plan  ⬇",
                    data=_gen_lp_bytes,
                    file_name=f"{_filename_stem}_LP.pdf",
                    mime="application/pdf",
                    key="gen_lp_primary_dl",
                    type="primary",
                    use_container_width=True,
                )
            except Exception as _gen_lp_err:
                st.caption(f"LP PDF error: {_gen_lp_err}")
        with _pdl_c2:
            # LP/A split — only show the Assessment download when the result
            # actually carries assessment_items. LP-only runs leave this slot
            # empty; the teacher generates assessment later from My Plans.
            _gen_has_assessment = (
                result.get("plan_status", "full_lpa") == "full_lpa"
                and bool(result.get("assessment_items"))
            )
            if _gen_has_assessment:
                try:
                    from assessment_pdf_generator import build_assessment_pdf_bytes as _bapb_gen
                    _gen_assess_payload = {
                        "saved_at":       datetime.now().isoformat(timespec="seconds"),
                        "grade":          _res_grade,
                        "subject":        _res_subject,
                        "chapter_number": _chapter_export.get("chapter_number", 0),
                        "chapter_title":  _chapter_export.get("chapter_title",  ""),
                        "result": {
                            "lesson_plan":      result.get("lesson_plan", {}),
                            "assessment_items": result.get("assessment_items", []),
                        },
                    }
                    _gen_assess_bytes = _bapb_gen(_gen_assess_payload)
                except Exception as _gen_assess_err:
                    _gen_assess_bytes = b""
                st.download_button(
                    label="Assessment  ⬇",
                    data=_gen_assess_bytes if _gen_assess_bytes else b"",
                    file_name=f"{_filename_stem}_Assessment.pdf",
                    mime="application/pdf",
                    key="gen_assess_primary_dl",
                    type="primary",
                    use_container_width=True,
                )
        with _pdl_c3:
            _already_saved = st.session_state.get("plan_already_saved", False)
            if st.button(
                "Saved ✓" if _already_saved else "Save to my plans",
                key="gen_save_top",
                use_container_width=True,
                disabled=_already_saved,
            ):
                save_plan(
                    grade       = _res_grade,
                    subject     = _res_subject,
                    chapter     = _chapter_export,
                    period_rows = st.session_state.get("period_rows", [0]),
                    session     = st.session_state,
                    result      = result,
                )
                st.session_state.plan_just_saved    = True
                st.session_state.plan_already_saved = True
                st.rerun()
        with _pdl_c4:
            if st.button(
                "Clear",
                key="gen_clear_top",
                use_container_width=True,
            ):
                st.session_state.lpa_result         = None
                st.session_state.show_save_prompt   = False
                st.session_state.plan_already_saved = False
                st.session_state.plan_just_saved    = False
                st.rerun()
        if st.session_state.get("plan_just_saved"):
            st.components.v1.html(
                """
                <script>
                (function() {
                    // Build modal overlay
                    var overlay = document.createElement('div');
                    overlay.style.cssText = [
                        'position:fixed','top:0','left:0','width:100%','height:100%',
                        'background:rgba(0,0,0,0.45)','z-index:99999',
                        'display:flex','align-items:center','justify-content:center'
                    ].join(';');

                    var box = document.createElement('div');
                    box.style.cssText = [
                        'background:#ffffff','border-radius:12px',
                        'padding:40px 48px','text-align:center',
                        'box-shadow:0 8px 32px rgba(0,0,0,0.18)',
                        'min-width:280px','max-width:380px'
                    ].join(';');

                    var icon = document.createElement('div');
                    icon.textContent = '✓';
                    icon.style.cssText = [
                        'font-size:2.4rem','color:#2e7d32',
                        'font-weight:700','margin-bottom:12px'
                    ].join(';');

                    var msg = document.createElement('p');
                    msg.textContent = 'Saved — view it in My Plans.';
                    msg.style.cssText = [
                        'font-size:1.05rem','color:#1a1a1a',
                        'margin:0 0 24px 0','font-family:sans-serif'
                    ].join(';');

                    var btn = document.createElement('button');
                    btn.textContent = 'OK';
                    btn.style.cssText = [
                        'background:#2e7d32','color:#fff','border:none',
                        'border-radius:6px','padding:10px 36px',
                        'font-size:1rem','cursor:pointer','font-family:sans-serif'
                    ].join(';');
                    btn.onmouseover = function() { btn.style.background = '#1b5e20'; };
                    btn.onmouseout  = function() { btn.style.background = '#2e7d32'; };
                    btn.onclick = function() { overlay.remove(); };

                    box.appendChild(icon);
                    box.appendChild(msg);
                    box.appendChild(btn);
                    overlay.appendChild(box);

                    // Attach to the top-level parent document (outside the iframe)
                    var target = window.parent ? window.parent.document.body : document.body;
                    target.appendChild(overlay);
                })();
                </script>
                """,
                height=0,
            )
            st.session_state.plan_just_saved = False
        st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)

        # ── LPA HTML page ─────────────────────────────────────────────────
        _lpa_html_path = PROJECT_ROOT / "lpa_page.html"
        try:
            _lpa_tpl = _lpa_html_path.read_text(encoding="utf-8")
        except Exception:
            _lpa_tpl = "<p>lpa_page.html not found.</p>"

        if st.session_state.teacher_ch_idx is not None and st.session_state.teacher_ch_idx < len(chapters):
            _ch_data = chapters[st.session_state.teacher_ch_idx]
        else:
            _ch_data = _chapter_export
        _period_schedule = " · ".join(
            f'{st.session_state.get(f"cnt_{r}", 1)} × {st.session_state.get(f"dur_sel_{r}", 40)}-min'
            for r in st.session_state.get("period_rows", [0])
        )

        # ── Normalise to lpa_page.html-compatible shape (handles old + new JSON) ─
        _grade_ctx   = _res_grade
        _subject_ctx = _res_subject
        _stage  = get_stage(_grade_ctx)
        _subj_f = subject_to_folder(_subject_ctx)
        try:
            if _subj_f == "the_world_around_us":
                _comp_descs_raw = json.loads(
                    (PROJECT_ROOT / f"mirror/framework/{_subj_f}"
                     / "preparatory/competency_descriptions_twau.json")
                    .read_text(encoding="utf-8")
                )
                # TWAU file has curricular_goals list — flatten to {c_code: description}
                _comp_descs = {}
                for _cg in _comp_descs_raw.get("curricular_goals", []):
                    for _c in _cg.get("competencies", []):
                        _comp_descs[_c["code"]] = _c["description"]
            else:
                _comp_descs = json.loads(
                    (PROJECT_ROOT / f"mirror/framework/{_subj_f}/{_stage}"
                     / f"competency_descriptions_{_stage}.json")
                    .read_text(encoding="utf-8")
                )
        except Exception:
            _comp_descs = {}

        _lo_handoff          = _normalise_lo_handoff(result, _comp_descs)
        _assessment_sections = _normalise_assessment_sections(result, _comp_descs)

        _lpa_data = {
            "chapter_title":       _ch_data.get("chapter_title", ""),
            "chapter_number":      _ch_data.get("chapter_number", ""),
            "grade":               _res_grade,
            "subject":             _res_subject,
            "period_schedule":     _period_schedule,
            "lo_handoff":          _lo_handoff,
            "assessment_sections": _assessment_sections,
        }
        _lpa_inject = "window.LPA_DATA = " + json.dumps(_lpa_data, ensure_ascii=False) + ";\n"
        _lpa_html = _lpa_tpl.replace("/* __LPA_DATA__ */", _lpa_inject)
        _lpa_height_script = """
<script>
(function() {
  /* Measure the actual .lpa content element — avoids the scrollHeight==viewport
     problem (scrollHeight equals iframe height when content is shorter). */
  function fitIframe() {
    var lpa = document.querySelector('.lpa');
    if (!lpa) return;
    var h = Math.ceil(lpa.getBoundingClientRect().height) + 24;
    if (h < 100) return;

    /* Method A: window.frameElement (works when Streamlit serves component
       via a same-origin URL, i.e. localhost:8501/component/...) */
    try {
      var fe = window.frameElement;
      if (fe) {
        fe.style.height = h + 'px';
        var p1 = fe.parentElement;
        if (p1) { p1.style.height = h + 'px'; p1.style.minHeight = '0'; }
        var p2 = p1 && p1.parentElement;
        if (p2) { p2.style.height = h + 'px'; p2.style.minHeight = '0'; }
      }
    } catch(e) {}

    /* Method B: access the Streamlit page DOM directly from the parent window
       (same-origin: Streamlit app and component URL share localhost:8501).
       Find the component iframe by its Streamlit-assigned height attribute and
       collapse it + its wrapper containers to the measured content height. */
    try {
      var pDoc = window.parent.document;
      /* Streamlit renders: <iframe height="2200" ...> inside a wrapper div */
      var targets = pDoc.querySelectorAll('iframe[height="2200"]');
      for (var i = 0; i < targets.length; i++) {
        var fr = targets[i];
        fr.setAttribute('height', String(h));
        fr.style.height = h + 'px';
        var w1 = fr.parentElement;
        if (w1) { w1.style.height = h + 'px'; w1.style.minHeight = '0'; }
        var w2 = w1 && w1.parentElement;
        if (w2) { w2.style.height = h + 'px'; w2.style.minHeight = '0'; }
      }
    } catch(e) {}

    /* Method C: Streamlit postMessage protocol (both recognised formats) */
    try {
      window.parent.postMessage({ type: 'streamlit:setFrameHeight', height: h }, '*');
      window.parent.postMessage(
        { isStreamlitMessage: true, type: 'streamlit:setFrameHeight', height: h }, '*'
      );
    } catch(e) {}
  }

  /* lpa_page.html fires its own (broken) reportHeight at 100 ms, 400 ms, 800 ms.
     Those calls measure scrollHeight which equals the 2200 px viewport height and
     would override a correct value. We fire JUST AFTER each one to win the race,
     plus an early shot and a late cleanup pass. */
  setTimeout(fitIframe,   50);
  setTimeout(fitIframe,  150);  /* override broken 100 ms call  */
  setTimeout(fitIframe,  450);  /* override broken 400 ms call  */
  setTimeout(fitIframe,  850);  /* override broken 800 ms call  */
  setTimeout(fitIframe, 1200);  /* final cleanup                */

  /* Re-fit on every collapsible expand / collapse */
  var t = null;
  function dFit() { clearTimeout(t); t = setTimeout(fitIframe, 150); }
  if (document.body) {
    new MutationObserver(dFit).observe(
      document.body, { childList: true, subtree: true, attributes: true }
    );
  }
})();
</script>
"""
        _lpa_html = _lpa_html + _lpa_height_script
        components.html(_lpa_html, height=2200, scrolling=False)



# ═════════════════════════════════════════════════
#  ALLOCATE WORKSPACE
# ═════════════════════════════════════════════════
elif st.session_state.role == "Allocate":

    # ── Gather period types from sidebar state ─────────────────────────────────
    _pt_rows = st.session_state.get("period_rows_p", [0])
    _period_types = [
        {
            "mins":  int(st.session_state.get(f"dur_sel_p{r}") or 40),
            "count": int(st.session_state.get(f"cnt_p{r}")     or 1),
        }
        for r in _pt_rows
    ]
    _period_types = [pt for pt in _period_types if pt["mins"] > 0 and pt["count"] > 0]
    _sorted_pts   = sorted(_period_types, key=lambda pt: -pt["mins"])

    # ── Load full mapping JSONs for all chapters ───────────────────────────────
    def _load_chapter_mappings(grade, subject):
        _stage  = get_stage(grade)
        _subj_f = subject_to_folder(subject)

        # Load competency descriptions lookup (keyed by c_code → description text)
        if _subj_f == "the_world_around_us":
            _comp_desc_path = (
                PROJECT_ROOT
                / f"mirror/framework/{_subj_f}"
                / "preparatory/competency_descriptions_twau.json"
            )
        else:
            _comp_desc_path = (
                PROJECT_ROOT
                / f"mirror/framework/{_subj_f}/{_stage}"
                / f"competency_descriptions_{_stage}.json"
            )
        try:
            _raw_descs = json.loads(_comp_desc_path.read_text(encoding="utf-8"))
            # Three formats exist:
            #   Science:     {curricular_goals: [{cg_code, competencies: [{code, description}]}]}  ← list of objects
            #   Mathematics: {curricular_goals: {"CG-1": {competency_codes: {"C-1.1": "text"}}}}   ← dict of dicts
            #   SS/Lang:     {c_code: description_string, ...}                                      ← flat dict
            if "curricular_goals" in _raw_descs:
                _cg_val = _raw_descs["curricular_goals"]
                _comp_descs = {}
                if isinstance(_cg_val, list):
                    # Science format: list of CG objects
                    for _cg in _cg_val:
                        for _comp in _cg.get("competencies", []):
                            _comp_descs[_comp.get("code", "")] = _comp.get("description", "")
                elif isinstance(_cg_val, dict):
                    # Mathematics format: dict keyed by CG code
                    for _cg_code, _cg_body in _cg_val.items():
                        _ccodes = _cg_body.get("competency_codes", {})
                        for _c_code, _desc in _ccodes.items():
                            _comp_descs[_c_code] = _desc
            else:
                _comp_descs = _raw_descs
        except Exception:
            _comp_descs = {}

        _result = []
        for _ch in chapters:
            _paths = resolve_paths(grade, subject, _ch["chapter_number"])
            try:
                _mapping = json.loads(
                    _paths["chapter_mapping"].read_text(encoding="utf-8")
                )
            except Exception:
                _mapping = {}

            # English ch_02 schema nests effort signals under "effort_signals" and
            # competencies under "competencies.primary" (a list of strings).
            # Flatten both into the top-level mapping dict for uniform downstream access.
            if "effort_signals" in _mapping:
                _mapping.update(_mapping.pop("effort_signals"))
            if "competencies" in _mapping and isinstance(_mapping["competencies"], dict):
                _comp_block = _mapping.pop("competencies")
                # Only promote if "primary" not already at top level
                if "primary" not in _mapping:
                    _mapping["primary"] = _comp_block.get("primary", [])
                if "incidental" not in _mapping:
                    _mapping["incidental"] = _comp_block.get("incidental", [])

            # Enrich primary competencies with full description text.
            # Key order: "primary" (Science VII / English), "competencies" (VI schema),
            # then Mathematics which uses "core_competencies" + "adjunct_competencies".
            if "primary" in _mapping or "competencies" in _mapping:
                _primary_entries = _mapping.get("primary", _mapping.get("competencies", []))
            else:
                # Mathematics schema: merge core + adjunct
                _primary_entries = (
                    _mapping.get("core_competencies", []) +
                    _mapping.get("adjunct_competencies", [])
                )
            _enriched_primary = []
            for _entry in _primary_entries:
                # _entry may be a dict {"c_code": ..., "weight": ...} or a bare string "C-1.1"
                if isinstance(_entry, str):
                    _e = {"c_code": _entry, "weight": 1}
                else:
                    _e = dict(_entry)
                # Prefer the shared descriptions lookup; fall back to inline
                # competency_text (TWAU stores it directly in the mapping entry).
                _e["description"] = (
                    _comp_descs.get(_e.get("c_code", ""), "")
                    or _e.get("competency_text", "")
                )
                _enriched_primary.append(_e)

            # TWAU: effort signals live in the summary JSON (not the mapping).
            # Read summary to pull conceptual_demand, task_load, project_load, map_work.
            _twau_signals = {}
            if subject == "The World Around Us":
                try:
                    _summary_raw = json.loads(
                        _paths["chapter_summary"].read_text(encoding="utf-8")
                    )
                    _twau_signals = {
                        "conceptual_demand": _summary_raw.get("conceptual_demand", 0),
                        "task_load":         _summary_raw.get("task_load", 0),
                        "project_load":      _summary_raw.get("project_load", 0),
                        "map_work":          _summary_raw.get("map_work", 0),
                    }
                except Exception:
                    _twau_signals = {
                        "conceptual_demand": 0, "task_load": 0,
                        "project_load": 0, "map_work": 0,
                    }

            _result.append({
                "chapter_number":    _ch["chapter_number"],
                "chapter_title":     _ch.get("chapter_title", ""),
                "chapter_weight":    _mapping.get("chapter_weight", 0),
                "effort_index":      _mapping.get("effort_index", 0),
                # Science signals
                "conceptual_demand": _twau_signals.get("conceptual_demand",
                                         _mapping.get("conceptual_demand", 0)),
                "activity_count":    _mapping.get("activity_count", 0),
                "activity_load":     _mapping.get("activity_load", 0),
                "demo_count":        _mapping.get("demo_count", 0),
                "demo_load":         _mapping.get("demo_load", 0),
                "exec_load":         _mapping.get("exec_load", 0),
                # English signals
                "spine_load":        _mapping.get("spine_load", 0),
                "task_density":      _mapping.get("task_density", 0),
                "writing_demand":    _mapping.get("writing_demand", 0),
                "project_load":      _twau_signals.get("project_load",
                                         _mapping.get("project_load", 0)),
                # TWAU-specific signals
                "task_load":         _twau_signals.get("task_load",
                                         _mapping.get("task_load", 0)),
                "map_work":          _twau_signals.get("map_work", 0),
                # Mathematics preparatory signals
                "exploration_load":  _mapping.get("exploration_load", 0),
                "procedural_load":   _mapping.get("procedural_load", 0),
                "primary":           _enriched_primary,
                "incidental":        _mapping.get("incidental", []),
            })
        return _result

    _chapters_data = _load_chapter_mappings(
        st.session_state.grade, st.session_state.subject
    )

    # ── Load HTML template and inject data ────────────────────────────────────
    _html_path = PROJECT_ROOT / "allocate_page.html"
    _html_tpl  = _html_path.read_text(encoding="utf-8")

    import base64 as _b64
    _logo_path = PROJECT_ROOT / "miscellaneous/aruvi_logo-transparent.png"
    try:
        _logo_b64 = _b64.b64encode(_logo_path.read_bytes()).decode()
    except Exception:
        _logo_b64 = ""

    # Load English spine data for the PDF report (textbook_section_names + competency codes)
    _english_spine_data = {}
    if st.session_state.subject == "English":
        _stage = get_stage(st.session_state.grade)
        _spine_path = PROJECT_ROOT / f"mirror/framework/english/{_stage}/spine_to_cg.json"
        try:
            _spine_raw = json.loads(_spine_path.read_text(encoding="utf-8"))
            _english_spine_data = _spine_raw.get("spines", {})
        except Exception:
            _english_spine_data = {}

    _inject = (
        f"const CHAPTERS_DATA  = {json.dumps(_chapters_data, ensure_ascii=False)};\n"
        f"const PERIOD_TYPES   = {json.dumps(_sorted_pts)};\n"
        f"const GRADE_LABEL    = {json.dumps(st.session_state.grade    or '')};\n"
        f"const SUBJECT_LABEL  = {json.dumps(st.session_state.subject  or '')};\n"
        f"const IS_SCIENCE     = {json.dumps(st.session_state.subject in ('Science', 'Mathematics'))};\n"
        f"const IS_MATH_PREP   = {json.dumps(st.session_state.subject == 'Mathematics' and get_stage(st.session_state.grade or '') == 'preparatory')};\n"
        f"const IS_ENGLISH     = {json.dumps(st.session_state.subject == 'English')};\n"
        f"const IS_TWAU        = {json.dumps(st.session_state.subject == 'The World Around Us')};\n"
        f"const ARUVI_LOGO_B64 = {json.dumps(_logo_b64)};\n"
        f"const ENGLISH_SPINE_DATA = {json.dumps(_english_spine_data, ensure_ascii=False)};\n"
    )
    _html = _html_tpl.replace("/* __CHAPTERS_DATA__ */", _inject)

    # Inject the correct footnote text directly into the static HTML so it is
    # correct on first render, before any JS runs.
    _subject = st.session_state.subject or ""
    if _subject == "English":
        _fn1_text = (
            '<div class="about-ei">'
            '<h4>About the Effort Index</h4>'
            '<p>The effort index is a number that tells you how much classroom '
            'time a chapter typically needs compared to other chapters in the '
            'subject. Chapters with a higher effort index get more periods; '
            'chapters with a lower one get fewer. It is calculated from four '
            'signals, each scored on a simple scale.</p>'
            '<ul>'
            '<li><b>Spine load</b> — How many types of classroom work '
            '(reading for comprehension, listening, speaking, writing, '
            'vocabulary, beyond-text) appear on average per section. '
            'More types = higher score.</li>'
            '<li><b>Task density</b> — How many tasks appear on average '
            'within each block of work. More tasks per block = higher '
            'score.</li>'
            '<li><b>Writing demand</b> — Total exercise items under '
            'Writing and Beyond-the-Text across the chapter. These take '
            'longer to complete and assess, so a heavier count raises the '
            'score.</li>'
            '<li><b>Project load</b> — How many Beyond-the-Text sections '
            'the chapter has. Each one adds to the score as these activities '
            'need extra planning time.</li>'
            '</ul>'
            '<p class="about-ei-close">The four scores are combined with '
            'fixed weights to give the effort index. Only relative values '
            'matter — it is used to share your available periods across '
            'chapters in proportion to their load.</p>'
            '</div>'
        )
    elif _subject == "The World Around Us":
        _fn1_text = (
            '<div class="about-ei">'
            '<h4>About the Effort Index</h4>'
            '<p>The effort index tells you how much classroom time a chapter typically '
            'needs compared to other chapters. Formula: '
            '<b>(Conceptual demand × 3) + (Task load × 2) + (Project load × 1.5) + Map work</b>. '
            'Chapters with a higher effort index get more periods.</p>'
            '<ul>'
            '<li><b>Conceptual demand (×3)</b> — How abstract the chapter\'s reasoning is: '
            '1 = concrete/tangible (family, simple observations); '
            '2 = slight abstraction — comparisons, simple cause-and-effect, or categorisation of familiar objects; '
            '3 = classification or material properties; '
            '4 = multi-step reasoning or concepts requiring inference beyond direct observation (e.g. seasonal cycles, ecosystem interdependence); '
            '5 = geological/astronomical/cultural-history abstraction.</li>'
            '<li><b>Task load (×2)</b> — Discrete score (0–3) based on the total count of '
            'student tasks (Activities, Discuss, Write, Find out, Draw, Let us reflect). '
            '0 = fewer than 10 tasks; 1 = 10–20; 2 = 21–30; 3 = more than 30.</li>'
            '<li><b>Project load (×1.5)</b> — 0 = none; 1 = light (multi-day observation); '
            '2 = substantial (artefact construction or sustained build project).</li>'
            '<li><b>Map work (×1)</b> — 0 = no maps; 1 = map reading; 2 = map drawing or '
            'regional comparison.</li>'
            '</ul>'
            '<p class="about-ei-close">Only relative values matter — the effort index is used '
            'to share your available periods across chapters in proportion to their load.</p>'
            '</div>'
        )
    elif _subject == "Mathematics" and get_stage(st.session_state.grade or "") == "preparatory":
        _fn1_text = (
            '<div class="about-ei">'
            '<h4>About the Effort Index</h4>'
            '<p>The effort index tells you how much classroom time a chapter typically needs '
            'compared to other chapters. Formula: '
            '<b>(Conceptual demand × 2) + (Task load × 2) + (Exploration load × 1.5) + (Procedural load × 1.5)</b>. '
            'Chapters with a higher effort index get more periods.</p>'
            '<ul>'
            '<li><b>Conceptual demand (×2)</b> — How abstract the chapter\'s reasoning is on the '
            'concrete-to-symbolic path. 1 = fully concrete (counting, matching, sorting tangible '
            'objects); 2 = slight abstraction (place value, simple patterns, measurement with '
            'standard units); 3 = symbolic or multi-step reasoning (multi-digit operations, '
            'fraction concepts, area/perimeter reasoning).</li>'
            '<li><b>Task load (×2)</b> — Discrete tier from total task count. '
            '0 = fewer than 8 tasks; 1 = 8–15; 2 = 16–25; 3 = more than 25.</li>'
            '<li><b>Exploration load (×1.5)</b> — Share of hands-on, manipulative, or game-based '
            'tasks. 0 = none; 1 = a few; 2 = prominently exploratory.</li>'
            '<li><b>Procedural load (×1.5)</b> — Share of compute, convert, or drill tasks. '
            '0 = none; 1 = moderate; 2 = heavily procedural.</li>'
            '</ul>'
            '<p class="about-ei-close">Only relative values matter — the effort index is used '
            'to share your available periods across chapters in proportion to their load.</p>'
            '</div>'
        )
    elif _subject == "Mathematics":
        _fn1_text = (
            '<div class="about-ei">'
            '<h4>About the Effort Index</h4>'
            '<p>The effort index is a number that tells you how much classroom '
            'time a chapter typically needs compared to other chapters in the '
            'subject. Chapters with a higher effort index get more periods; '
            'chapters with a lower one get fewer. It is calculated from four '
            'signals read from the chapter content.</p>'
            '<p><b>effort_index = (Conceptual demand × 2) + (Activity load × 2) + '
            '(Demo load × 1.5) + (Exercise execution load × 2)</b></p>'
            '<ul>'
            '<li><b>Conceptual demand (×2)</b> — The cognitive complexity of exercises and questions '
            'in the chapter (1–3). High-order thinking or multi-step reasoning raises the score.</li>'
            '<li><b>Activity load (×2)</b> — A discrete 0–3 tier from the number of hands-on activities '
            'students perform themselves: 0 = none; 1 = 1–3 (light); 2 = 4–7 (standard); '
            '3 = 8 or more (activity-heavy).</li>'
            '<li><b>Demo load (×1.5)</b> — A discrete 0–2 tier from the number of teacher demonstrations: '
            '0 = none; 1 = 1–2; 2 = 3 or more. These need preparation and focused class attention.</li>'
            '<li><b>Exercise execution load (×2)</b> — The weight of multi-step calculation or diagram '
            'production in the exercises (0–2). A heavier execution load means more time for guided '
            'practice and assessment.</li>'
            '</ul>'
            '<p class="about-ei-close">All four signals are bounded discrete tiers combined with fixed '
            'weights, so no single signal can dominate. Only relative values matter — the index shares '
            'your available periods across chapters in proportion to their load.</p>'
            '</div>'
        )
    elif _subject == "Science":
        _fn1_text = (
            '<div class="about-ei">'
            '<h4>About the Effort Index</h4>'
            '<p>The effort index is a number that tells you how much classroom '
            'time a chapter typically needs compared to other chapters in the '
            'subject. Chapters with a higher effort index get more periods; '
            'chapters with a lower one get fewer. It is calculated from four '
            'signals read from the chapter content.</p>'
            '<p><b>effort_index = (Conceptual demand × 2) + (Activity load × 2) + '
            '(Demo load × 1.5) + (Exercise execution load × 2)</b></p>'
            '<ul>'
            '<li><b>Conceptual demand (×2)</b> — The cognitive complexity of exercises and questions '
            'in the chapter (1–3). High-order thinking or multi-step reasoning raises the score.</li>'
            '<li><b>Activity load (×2)</b> — A discrete 0–3 tier from the number of hands-on activities '
            'students perform themselves: 0 = none; 1 = 1–3 (light); 2 = 4–7 (standard); '
            '3 = 8 or more (activity-heavy).</li>'
            '<li><b>Demo load (×1.5)</b> — A discrete 0–2 tier from the number of teacher demonstrations: '
            '0 = none; 1 = 1–2; 2 = 3 or more. These need preparation and focused class attention.</li>'
            '<li><b>Exercise execution load (×2)</b> — The weight of multi-step calculation or diagram '
            'production in the exercises (0–2). A heavier execution load means more time for guided '
            'practice and assessment.</li>'
            '</ul>'
            '<p class="about-ei-close">All four signals are bounded discrete tiers combined with fixed '
            'weights, so no single signal can dominate. Only relative values matter — the index shares '
            'your available periods across chapters in proportion to their load.</p>'
            '</div>'
        )
    else:
        _fn1_text = (
            "Periods allocated using the Largest Remainder Method (LRM), "
            "weighted by chapter competency load (W3 × 3 +"
            " W2 × 2 + W1 × 1)."
        )
    _html = _html.replace('<p id="fn1"></p>', f'<p id="fn1">{_fn1_text}</p>')

    components.html(_html, height=950, scrolling=True)

else:
    # ═════════════════════════════════════════════════
    #  MY PLANS WORKSPACE
    # ═════════════════════════════════════════════════

    _sp_root = PROJECT_ROOT / "mirror" / "saved_plans"

    # NOTE: the deferred-assessment generation block has moved to the Generate
    # workspace (see the role=="Generate" branch). When "Generate Assessment"
    # is clicked here, role flips to "Generate" so the popup appears in the
    # same place as a normal LP/A run, then flips back to "My Plans" on
    # completion so the teacher lands on the now-PDF row.

    # ── Detail view — shown when a plan row's View button has been clicked ────
    if st.session_state.mp_viewing_plan is not None:
        _vp       = st.session_state.mp_viewing_plan
        _vgrade   = _vp.get("grade",   "")
        _vsubject = _vp.get("subject", "")
        _v_ch_num   = _vp.get("chapter_number", 0)
        _v_ch_title = _vp.get("chapter_title",  "")

        # ── Back button (top) ─────────────────────────────────────────────────
        if st.button("← Back to My Plans", key="mp_back_top"):
            st.session_state.mp_viewing_plan = None
            st.rerun()

        st.markdown('<div style="height:0.25rem;"></div>', unsafe_allow_html=True)

        # ── Resolve chapter export dict ───────────────────────────────────────
        _v_all_chapters = load_all_chapters(_vgrade, _vsubject) if (_vgrade and _vsubject) else []
        _v_chapter_export = next(
            (c for c in _v_all_chapters if c["chapter_number"] == _v_ch_num), None
        )
        if _v_chapter_export is None:
            _v_chapter_export = {
                "chapter_title":  _v_ch_title,
                "chapter_number": _v_ch_num,
                "chapter_weight": "",
                "primary":        [],
            }

        _vresult = dict(_vp["result"])

        # ── Competency descriptions ───────────────────────────────────────────
        _v_stage  = get_stage(_vgrade)          if _vgrade   else "middle"
        _v_subj_f = subject_to_folder(_vsubject) if _vsubject else "social_sciences"
        try:
            _v_comp_descs = json.loads(
                (PROJECT_ROOT / f"mirror/framework/{_v_subj_f}/{_v_stage}"
                 / f"competency_descriptions_{_v_stage}.json")
                .read_text(encoding="utf-8")
            )
        except Exception:
            _v_comp_descs = {}

        # ── Normalise to lpa_page.html-compatible shape (handles old + new JSON) ─
        _v_lo_handoff          = _normalise_lo_handoff(_vresult, _v_comp_descs)
        _v_assessment_sections = _normalise_assessment_sections(_vresult, _v_comp_descs)

        # ── Render LPA HTML page ──────────────────────────────────────────────
        try:
            _v_lpa_tpl = (PROJECT_ROOT / "lpa_page.html").read_text(encoding="utf-8")
        except Exception:
            _v_lpa_tpl = "<p>lpa_page.html not found.</p>"

        _v_lpa_data = {
            "chapter_title":       _v_ch_title,
            "chapter_number":      _v_ch_num,
            "grade":               _vgrade,
            "subject":             _vsubject,
            "period_schedule":     _vp.get("period_schedule_display", ""),
            "lo_handoff":          _v_lo_handoff,
            "assessment_sections": _v_assessment_sections,
        }
        _v_lpa_inject = "window.LPA_DATA = " + json.dumps(_v_lpa_data, ensure_ascii=False) + ";\n"
        _v_lpa_height_script = """
<script>
(function() {
  /* Track the largest height seen during initial render so an early/partial
     measurement can never shrink the iframe below already-rendered content.
     A "settle" flag flips on after the page has been stable for a beat,
     after which collapses are allowed to shrink the iframe. */
  var maxSeen = 0;
  var settled = false;
  function measure() {
    var lpa = document.querySelector('.lpa');
    var lpaH = lpa ? Math.ceil(lpa.getBoundingClientRect().height) : 0;
    return Math.max(
      lpaH,
      lpa ? lpa.scrollHeight : 0,
      lpa ? lpa.offsetHeight : 0,
      document.body ? document.body.scrollHeight : 0,
      document.documentElement ? document.documentElement.scrollHeight : 0,
      document.documentElement ? document.documentElement.offsetHeight : 0
    );
  }
  function fitIframe() {
    var h = measure();
    if (h < 100) return;
    h += 30;
    if (!settled) {
      if (h < maxSeen) return;       /* never shrink during initial render */
      maxSeen = h;
    }
    try {
      window.parent.postMessage(
        { isStreamlitMessage: true, type: 'streamlit:setFrameHeight', height: h }, '*'
      );
    } catch(e) {}
  }
  /* Aggressive cascade of timeouts to cover slow async rendering */
  [30, 100, 250, 500, 900, 1500, 2500, 4000].forEach(function(d) {
    setTimeout(fitIframe, d);
  });
  /* After 4.5s assume the page has settled — collapses may now shrink it. */
  setTimeout(function() { settled = true; fitIframe(); }, 4500);

  /* ResizeObserver is the reliable signal that content size has changed
     (handles font loads, image decodes, accordion toggles). */
  function attachObservers() {
    var lpa = document.querySelector('.lpa');
    if (window.ResizeObserver) {
      var ro = new ResizeObserver(fitIframe);
      if (lpa) ro.observe(lpa);
      if (document.body) ro.observe(document.body);
    }
    /* MutationObserver as a backstop for browsers without ResizeObserver
       and for attribute-only changes that don't alter size synchronously. */
    var debounceTimer = null;
    if (document.body) {
      new MutationObserver(function() {
        clearTimeout(debounceTimer);
        debounceTimer = setTimeout(fitIframe, 120);
      }).observe(document.body, { childList: true, subtree: true, attributes: true });
    }
  }
  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', attachObservers);
  } else {
    attachObservers();
  }
})();
</script>
"""
        _v_lpa_html = _v_lpa_tpl.replace("/* __LPA_DATA__ */", _v_lpa_inject) + _v_lpa_height_script
        components.html(
            _v_lpa_html,
            height=2200, scrolling=False,
        )

        st.markdown('<div style="height:1rem;"></div>', unsafe_allow_html=True)
        st.divider()

        # ── Back button (bottom) ──────────────────────────────────────────────
        if st.button("← Back to My Plans", key="mp_back_bottom"):
            st.session_state.mp_viewing_plan = None
            st.rerun()


    if st.session_state.mp_viewing_plan is None:
        # Load ALL saved plans across all grades/subjects
        _all_plans = []
        if _sp_root.exists():
            for _f in sorted(_sp_root.rglob("ch_*.json"), reverse=True):
                try:
                    _all_plans.append(json.loads(_f.read_text(encoding="utf-8")))
                except Exception:
                    pass

        # Filters come from the sidebar Grade / Subject / Saved selectboxes.
        # Grade & Subject map None → unfiltered ("All"). Saved filters on the
        # ISO date in saved_at: Today = today, Yesterday = today-1, This week
        # = last 7 days inclusive of today, This month = current calendar
        # month from day 1 through today.
        from datetime import date as _date, timedelta as _td
        _today = _date.today()
        _saved_choice = st.session_state.get("mp_saved_filter", "All")
        if _saved_choice == "Today":
            _date_lo, _date_hi = _today, _today
        elif _saved_choice == "Yesterday":
            _y = _today - _td(days=1)
            _date_lo, _date_hi = _y, _y
        elif _saved_choice == "This week":
            _date_lo, _date_hi = _today - _td(days=6), _today
        elif _saved_choice == "This month":
            _date_lo, _date_hi = _today.replace(day=1), _today
        else:
            _date_lo, _date_hi = None, None

        def _in_date_window(plan):
            if _date_lo is None:
                return True
            _ts = plan.get("saved_at", "")[:10]
            try:
                _d = _date.fromisoformat(_ts)
            except Exception:
                return False
            return _date_lo <= _d <= _date_hi

        _visible = [
            p for p in _all_plans
            if (st.session_state.grade   is None or p.get("grade")   == st.session_state.grade)
            and (st.session_state.subject is None or p.get("subject") == st.session_state.subject)
            and _in_date_window(p)
        ]

        st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)

        if not _visible:
            st.markdown(
                '<div class="ws-placeholder">No saved plans yet. '
                'Generate a plan and click Save to My Plans.</div>',
                unsafe_allow_html=True,
            )
        else:
            # ── Pure-Streamlit column table (header + inline buttons per row) ───────
            st.markdown("""
    <style>
    .mp-th       { font-size:0.65rem; font-weight:600; letter-spacing:0.08em;
                   text-transform:uppercase; color:#5a5754; padding-bottom:2px; }
    .mp-ch-title { font-size:0.88rem; font-weight:500; color:#1a1917; margin-bottom:2px; }
    .mp-ch-meta  { font-size:0.72rem; color:#9c9693; }
    .mp-cell     { font-size:0.82rem; color:#3d3b38; padding-top:6px; }
    </style>
    """, unsafe_allow_html=True)

            # Header row
            _hc = st.columns([3, 1, 1.5, 0.8, 1.2, 1.2])
            _hc[0].markdown('<div class="mp-th">Chapter</div>',       unsafe_allow_html=True)
            _hc[1].markdown('<div class="mp-th">Grade</div>',         unsafe_allow_html=True)
            _hc[2].markdown('<div class="mp-th">Saved</div>',         unsafe_allow_html=True)
            _hc[3].markdown('<div class="mp-th">Display</div>',       unsafe_allow_html=True)
            _hc[4].markdown('<div class="mp-th" style="text-align:left;">Lesson plan</div>',   unsafe_allow_html=True)
            _hc[5].markdown('<div class="mp-th" style="text-align:left;">Assessment</div>',    unsafe_allow_html=True)
            st.markdown(
                '<hr style="margin:4px 0 6px;border:none;border-top:1px solid #e8e5e0;">',
                unsafe_allow_html=True,
            )

            # One row per plan
            for _p in _visible:
                _ch_num   = _p.get("chapter_number", 0)
                _ch_title = _p.get("chapter_title", "")
                _grade    = _p.get("grade", "")
                _subject  = _p.get("subject", "")
                _saved_at = _p.get("saved_at", "")[:10]
                _filename = _p.get("filename", "")
                _safe_fn  = re.sub(r"[^a-zA-Z0-9_]", "_", _filename)
                try:
                    from datetime import datetime as _dt
                    _saved_disp = _dt.fromisoformat(_saved_at).strftime("%-d %b %Y")
                except Exception:
                    _saved_disp = _saved_at
                _ch_for_pdf = next(
                    (c for c in chapters if c["chapter_number"] == _ch_num),
                    {"chapter_title": _ch_title, "chapter_weight": "",
                     "chapter_number": _ch_num, "primary": []}
                )
                # LP PDF via lp_pdf_generator (new ReportLab format)
                try:
                    from lp_pdf_generator import build_lp_pdf_bytes as _blpb_mp
                    _mp_lp_payload = {
                        "saved_at":       _p.get("saved_at", datetime.now().isoformat(timespec="seconds")),
                        "grade":          _grade,
                        "subject":        _subject,
                        "chapter_number": _ch_num,
                        "chapter_title":  _ch_title,
                        "result":         {"lesson_plan": _p["result"].get("lesson_plan", {})},
                    }
                    _mp_lp_bytes = _blpb_mp(_mp_lp_payload)
                except Exception:
                    _mp_lp_bytes = b""
                # Assessment PDF — new ReportLab format
                try:
                    from assessment_pdf_generator import build_assessment_pdf_bytes as _bapb_mp
                    _mp_assess_bytes = _bapb_mp(_p)
                except Exception:
                    _mp_assess_bytes = b""
                _safe_t = re.sub(r"[^\w\s-]", "", _ch_title).strip().replace(" ", "_")[:40]

                _rc = st.columns([3, 1, 1.5, 0.8, 1.2, 1.2])
                _rc[0].markdown(
                    f'<div class="mp-ch-title">{_ch_title}</div>'
                    f'<div class="mp-ch-meta">Ch {str(_ch_num).zfill(2)} · {_subject}</div>',
                    unsafe_allow_html=True,
                )
                _rc[1].markdown(f'<div class="mp-cell">{_grade}</div>',       unsafe_allow_html=True)
                _rc[2].markdown(f'<div class="mp-cell">{_saved_disp}</div>',  unsafe_allow_html=True)
                with _rc[3]:
                    if st.button("View", key=f"view_{_safe_fn}", use_container_width=True):
                        st.session_state.mp_viewing_plan = _p
                        st.rerun()
                with _rc[4]:
                    st.download_button(
                        label="PDF ⬇",
                        data=_mp_lp_bytes,
                        file_name=f"Aruvi_{_safe_t}_LP.pdf",
                        mime="application/pdf",
                        key=f"mp_lp_{_safe_fn}",
                        type="primary",
                    )
                with _rc[5]:
                    # LP/A split — assessment column branches on plan_status.
                    # full_lpa  → PDF download (existing behaviour).
                    # lp_only   → "Generate Assessment" button that triggers
                    #             a deferred run handled by the workspace block
                    #             at the top of this tab.
                    # Old plans (no plan_status key) default to full_lpa so the
                    # existing My Plans rows continue to work unchanged.
                    _plan_status   = _p.get("plan_status", "full_lpa")
                    _has_assessment = (_plan_status == "full_lpa")
                    if _has_assessment:
                        st.download_button(
                            label="PDF ⬇",
                            data=_mp_assess_bytes if _mp_assess_bytes else b"",
                            file_name=f"Aruvi_{_safe_t}_Assessment.pdf",
                            mime="application/pdf",
                            key=f"mp_assess_{_safe_fn}",
                            type="primary",
                        )
                    else:
                        if st.button(
                            "Generate Assessment",
                            key=f"mp_gen_assess_{_safe_fn}",
                            use_container_width=True,
                            type="secondary",
                        ):
                            st.session_state.mp_deferred_assess_plan       = _p
                            st.session_state.mp_deferred_assess_generating = True
                            # Clear any stale Generate-tab display state so the
                            # Generate workspace renders cleanly (no shadow of a
                            # previous LP result showing behind the deferred
                            # assessment popup).
                            st.session_state.lpa_result        = None
                            st.session_state.lpa_generating    = False
                            st.session_state.show_gen_confirm  = False
                            st.session_state.show_save_prompt  = False
                            st.session_state.teacher_ch_idx    = None
                            # Switch to the Generate tab so the popup renders
                            # in the same place as a normal LP/A run. The
                            # deferred block below auto-switches back to
                            # My Plans once the run completes.
                            st.session_state.role = "Generate"
                            st.query_params["role"] = "Generate"
                            st.rerun()
                st.markdown(
                    '<hr style="margin:2px 0;border:none;border-top:0.5px solid #f0ede9;">',
                    unsafe_allow_html=True,
                )

# ── Ask Aruvi FAB + Bottom Drawer ────────────────────────────────────────────
CATEGORY_LABELS = {
    "cat_a": "How Aruvi plans your lessons",
    "cat_b": "How Aruvi builds your assessments",
    "cat_c": "How time is allocated across chapters",
    "cat_d": "Using the platform",
    "cat_e": "What Aruvi cannot do",
}
st.markdown("""
<style>
/* ── Popup card ── */
div[class*="st-key-ask_aruvi_popup"] {
    position: fixed !important;
    bottom: 90px !important;
    right: 28px !important;
    width: 320px !important;
    max-height: 75vh !important;
    background: #FFFFFF !important;
    border-radius: 16px !important;
    border: 1px solid #E0DDD8 !important;
    z-index: 99998 !important;
    box-shadow: 0 20px 60px rgba(0,0,0,0.20), 0 4px 16px rgba(0,0,0,0.10) !important;
    overflow-y: auto !important;
    padding: 0 !important;
}
/* ── Secondary popup (Q&A + Feedback) — taller than category popup ── */
div[class*="st-key-ask_aruvi_agent_popup"] {
    position: fixed !important;
    bottom: 90px !important;
    right: 28px !important;
    width: 320px !important;
    max-height: 83vh !important;
    background: #FFFFFF !important;
    border-radius: 16px !important;
    border: 1px solid #E0DDD8 !important;
    z-index: 99999 !important;
    box-shadow: 0 20px 60px rgba(0,0,0,0.20), 0 4px 16px rgba(0,0,0,0.10) !important;
    overflow-y: auto !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] [data-testid="stVerticalBlock"] {
    gap: 0px !important;
    row-gap: 0px !important;
    padding-bottom: 40px !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] [data-testid="element-container"],
div[class*="st-key-ask_aruvi_agent_popup"] [data-testid="stVerticalBlockBorderWrapper"] {
    margin: 0 !important;
    padding: 0 !important;
}
/* Agent panel — white pill textareas with light boundary */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_query_input"],
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_text"] {
    padding: 0 12px 0 12px !important;
    margin: 0 !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_query_input"] [data-baseweb="textarea"],
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_query_input"] > div,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_text"] [data-baseweb="textarea"],
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_text"] > div {
    overflow: visible !important;
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_query_input"] textarea,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_text"] textarea {
    height: 96px !important;
    min-height: 96px !important;
    font-size: 0.85rem !important;
    border-radius: 14px !important;
    border: 1px solid #D9D5CF !important;
    resize: none !important;
    line-height: 1.45 !important;
    padding: 10px 52px 10px 14px !important;
    background: #FFFFFF !important;
    color: #2A2826 !important;
    width: 100% !important;
    box-sizing: border-box !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_query_input"] textarea::placeholder,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_text"] textarea::placeholder {
    color: #B8B4AE !important;
}
/* Submit buttons (↑) overlaid inside the textareas, no border */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"],
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] {
    position: relative !important;
    height: 0 !important;
    width: 100% !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] > div,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] > div > div,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] > div,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] > div > div {
    height: 0 !important;
    width: 100% !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] button,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] button {
    position: absolute !important;
    top: -88px !important;
    right: 16px !important;
    width: 34px !important;
    height: 34px !important;
    min-height: 34px !important;
    min-width: 34px !important;
    max-width: 34px !important;
    border-radius: 8px !important;
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
    color: #1B2A3B !important;
    font-size: 1.3rem !important;
    line-height: 1 !important;
    padding: 0 !important;
    z-index: 20 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] button *,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] button *,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] button p,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] button p {
    color: #1B2A3B !important;
    font-size: 1.3rem !important;
    line-height: 1 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_submit"] button:hover,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_fb_submit"] button:hover {
    color: #2C3E50 !important;
    background: transparent !important;
}
/* Thumbs-up / thumbs-down buttons under the response — drop border, grey the emoji */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_up"] button,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_down"] button {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    padding: 4px 6px !important;
    min-height: unset !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_up"] button:hover,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_down"] button:hover {
    background: transparent !important;
    border: none !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_up"] button *,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_thumb_down"] button * {
    filter: grayscale(1) brightness(0.45) !important;
}
/* Follow-up textarea (shown after thumbs-down) — unfilled box with dark grey border */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_followup"] {
    padding: 0 16px !important;
    margin: 0 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_followup"] [data-baseweb="textarea"],
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_followup"] > div {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_followup"] textarea {
    background: transparent !important;
    border: 1px solid #4A4A4A !important;
    border-radius: 10px !important;
    font-size: 0.78rem !important;
    color: #2A2826 !important;
    line-height: 1.4 !important;
    padding: 10px 12px !important;
    resize: none !important;
    width: 100% !important;
    box-sizing: border-box !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_followup"] textarea::placeholder {
    color: #6B6760 !important;
    font-size: 0.78rem !important;
}
/* Submit / Skip pill buttons under the follow-up textarea */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_submit"] button,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_skip"] button {
    background: #3A3A3A !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 999px !important;
    font-size: 0.7rem !important;
    font-weight: 500 !important;
    padding: 4px 14px !important;
    min-height: unset !important;
    width: auto !important;
    white-space: nowrap !important;
    min-width: max-content !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_submit"] button *,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_skip"] button * {
    color: #FFFFFF !important;
    font-size: 0.7rem !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_submit"] button:hover,
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_skip"] button:hover {
    background: #2A2A2A !important;
    color: #FFFFFF !important;
}
/* Action row holding the Submit / Skip pills — same horizontal inset as the textarea */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-agent_fu_actions"] {
    padding: 8px 16px 0 16px !important;
}
/* Prevent narrow-popup column wrap — keep Submit & Skip side by side */
div[class*="st-key-agent_fu_actions"] [data-testid="stHorizontalBlock"] {
    flex-wrap: nowrap !important;
    gap: 8px !important;
}
/* Submit / Skip columns shrink-wrap to the pill's natural width */
div[class*="st-key-agent_fu_actions"] [data-testid="stColumn"]:nth-child(1),
div[class*="st-key-agent_fu_actions"] [data-testid="stColumn"]:nth-child(2) {
    flex: 0 0 auto !important;
    width: auto !important;
    min-width: max-content !important;
}
/* Character counter under each textarea */
div[class*="st-key-ask_aruvi_agent_popup"] .aa-char-count {
    text-align: right !important;
    font-size: 0.72rem !important;
    color: #9A968F !important;
    padding: 4px 16px 0 0 !important;
    margin: 0 !important;
    line-height: 1 !important;
}
/* Back + close buttons inside agent panel */
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-aa_agent_back_btn"] button {
    background: transparent !important;
    border: none !important;
    color: #2C7A7B !important;
    font-size: 0.7rem !important;
    font-weight: 600 !important;
    padding: 12px 16px 10px 16px !important;
    min-height: unset !important;
    width: 100% !important;
    text-align: left !important;
    justify-content: flex-start !important;
    border-bottom: 1px solid #F0EDE9 !important;
    border-radius: 0 !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_close"] button {
    background: transparent !important;
    border: none !important;
    border-bottom: 1px solid #F0EDE9 !important;
    border-radius: 0 !important;
    color: #2C7A7B !important;
    font-size: 0.7rem !important;
    font-weight: 600 !important;
    padding: 12px 16px 10px 16px !important;
    min-height: unset !important;
    width: 100% !important;
    text-align: left !important;
    justify-content: flex-start !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] div[class*="st-key-ask_aruvi_agent_close"] button:hover {
    background: #F5F9F9 !important;
    color: #1B2A3B !important;
}
/* Entry-point buttons in main popup — style as subdued link rows */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_open_question"] button,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_open_feedback"] button {
    background: transparent !important;
    border: none !important;
    border-top: 1px solid #F0EDE9 !important;
    border-radius: 0 !important;
    color: #2C7A7B !important;
    font-size: 0.72rem !important;
    font-weight: 500 !important;
    text-align: left !important;
    justify-content: flex-start !important;
    width: 100% !important;
    padding: 12px 16px !important;
    min-height: unset !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_open_question"] button:hover,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_open_feedback"] button:hover {
    background: #F5F9F9 !important;
    color: #1B2A3B !important;
}
/* Inline scope disclaimer shown above the question textarea */
div[class*="st-key-ask_aruvi_agent_popup"] .aa-scope-note {
    font-size: 0.78rem !important;
    color: #5C5852 !important;
    padding: 8px 16px 10px 16px !important;
    margin: 0 !important;
    line-height: 1.45 !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] .aa-scope-note .aa-scope-head {
    color: #2C3E50 !important;
    font-weight: 600 !important;
    display: block !important;
    margin-bottom: 3px !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] .aa-scope-note .aa-scope-block + .aa-scope-block {
    margin-top: 12px !important;
}
div[class*="st-key-ask_aruvi_agent_popup"] .aa-scope-note p {
    margin: 0 !important;
    padding: 0 !important;
}
/* Kill all Streamlit internal spacing */
div[class*="st-key-ask_aruvi_popup"] [data-testid="stVerticalBlock"] {
    gap: 0px !important;
    row-gap: 0px !important;
}
div[class*="st-key-ask_aruvi_popup"] [data-testid="element-container"],
div[class*="st-key-ask_aruvi_popup"] [data-testid="stVerticalBlockBorderWrapper"] {
    margin: 0 !important;
    padding: 0 !important;
}
/* Header bar */
.aa-header {
    padding: 14px 16px 10px 16px;
    border-bottom: 1px solid #F0EDE9;
}
.aa-header-title {
    font-size: 0.62rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #2C7A7B;
    margin: 0;
}
.aa-header-sub {
    font-size: 0.7rem;
    color: #9C9693;
    margin-top: 2px;
}
/* Category cards */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] {
    margin: 0 !important;
    padding: 0 8px !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] button,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] [data-testid="stBaseButton-secondary"],
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] [data-testid="stBaseButton-primary"] {
    background: #FAFAF8 !important;
    border: none !important;
    border-bottom: 1px solid #F0EDE9 !important;
    border-radius: 0 !important;
    color: #2C2A27 !important;
    font-size: 0.75rem !important;
    font-weight: 400 !important;
    padding: 11px 12px !important;
    width: 100% !important;
    min-height: 40px !important;
    height: auto !important;
    text-align: left !important;
    justify-content: space-between !important;
    letter-spacing: 0 !important;
    line-height: 1.3 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] button *,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] [data-testid] * {
    font-size: 0.75rem !important;
    color: #2C2A27 !important;
    line-height: 1.3 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] button:hover,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-chip_"] button:hover * {
    background: #F0F7F7 !important;
    color: #1B2A3B !important;
}
div[class*="st-key-ask_aruvi_popup"] [data-testid="stBaseButton-primary"],
div[class*="st-key-ask_aruvi_popup"] [data-testid="stBaseButton-primary"] * {
    background: #EAF4F4 !important;
    color: #2C7A7B !important;
    font-weight: 600 !important;
    border: none !important;
    border-bottom: 1px solid #C8E8E8 !important;
}
/* Query box area */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] {
    padding: 0 12px 0 12px !important;
    margin: 0 !important;
    overflow: visible !important;
}
/* BaseWeb textarea wrapper has overflow:hidden by default — must override or bottom border clips */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] [data-baseweb="textarea"],
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] > div {
    overflow: visible !important;
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] textarea {
    height: 80px !important;
    min-height: 80px !important;
    font-size: 0.75rem !important;
    border-radius: 10px !important;
    border: 1px solid #E0DDD8 !important;
    resize: none !important;
    line-height: 1.5 !important;
    padding: 10px 36px 10px 12px !important;
    background: #FAFAF8 !important;
    color: #2C2A27 !important;
    width: 100% !important;
    box-sizing: border-box !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] textarea::placeholder {
    color: #C0BCB8 !important;
    opacity: 1 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] textarea:focus {
    border-color: #2C7A7B !important;
    outline: none !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] p,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_query_input"] small {
    font-size: 0.58rem !important;
    color: #B8B4B0 !important;
}
/* Send button — CRITICAL: width:100% prevents fit-content from breaking right: anchor */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_submit"] {
    position: relative !important;
    height: 0 !important;
    width: 100% !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_submit"] > div,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_submit"] > div > div {
    height: 0 !important;
    width: 100% !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_submit"] button {
    position: absolute !important;
    top: 4px !important;
    right: 14px !important;
    width: 26px !important;
    height: 26px !important;
    min-height: 26px !important;
    min-width: 26px !important;
    max-width: 26px !important;
    border-radius: 50% !important;
    background: #E8682A !important;
    border: none !important;
    color: #FFFFFF !important;
    font-size: 0.85rem !important;
    padding: 0 !important;
    line-height: 1 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    z-index: 20 !important;
    pointer-events: all !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_submit"] button:hover {
    background: #C95820 !important;
}
/* Response box */
.aa-response-wrap {
    padding: 0 12px 8px 12px;
}
.aruvi-response-box {
    background: #F5F9F9;
    border-left: 3px solid #2C7A7B;
    padding: 10px 13px;
    font-size: 0.75rem;
    color: #2C2A27;
    margin-top: 6px;
    border-radius: 0 8px 8px 0;
    line-height: 1.6;
}
/* Clear button */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_clear"] button {
    background: transparent !important;
    border: none !important;
    color: #B8B4B0 !important;
    font-size: 0.62rem !important;
    padding: 2px 0 0 0 !important;
    min-height: unset !important;
    text-decoration: underline !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_clear"] button:hover {
    color: #5A5754 !important;
}
/* Divider — hidden to reduce gap */
.aa-divider {
    display: none !important;
    margin: 0 !important;
}
/* Feedback section label */
.aa-fb-label {
    padding: 4px 16px 0 16px;
    font-size: 0.58rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: #2C7A7B;
}
/* Feedback textarea */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] {
    padding: 0 12px 0 12px !important;
    margin: 0 !important;
    overflow: visible !important;
}
/* BaseWeb textarea wrapper has overflow:hidden by default — must override or bottom border clips */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] [data-baseweb="textarea"],
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] > div {
    overflow: visible !important;
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] textarea {
    height: 80px !important;
    min-height: 80px !important;
    font-size: 0.75rem !important;
    border-radius: 10px !important;
    border: 1px solid #E0DDD8 !important;
    resize: none !important;
    line-height: 1.5 !important;
    padding: 10px 36px 10px 12px !important;
    background: #FAFAF8 !important;
    color: #2C2A27 !important;
    width: 100% !important;
    box-sizing: border-box !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] textarea::placeholder {
    color: #C0BCB8 !important;
    opacity: 1 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_text"] textarea:focus {
    border-color: #2C7A7B !important;
    outline: none !important;
}
/* Feedback send button — CRITICAL: width:100% prevents fit-content from breaking right: anchor */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_submit"] {
    position: relative !important;
    height: 0 !important;
    width: 100% !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_submit"] > div,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_submit"] > div > div {
    height: 0 !important;
    width: 100% !important;
    overflow: visible !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_submit"] button {
    position: absolute !important;
    top: 4px !important;
    right: 14px !important;
    width: 26px !important;
    height: 26px !important;
    min-height: 26px !important;
    min-width: 26px !important;
    max-width: 26px !important;
    border-radius: 50% !important;
    background: #E8682A !important;
    border: none !important;
    color: #FFFFFF !important;
    font-size: 0.85rem !important;
    padding: 0 !important;
    line-height: 1 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    z-index: 20 !important;
    pointer-events: all !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-ask_aruvi_fb_submit"] button:hover {
    background: #C95820 !important;
}
/* Thumbs buttons */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-thumb_"] button {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
    border-radius: 50% !important;
    font-size: 1rem !important;
    width: 28px !important;
    height: 28px !important;
    min-height: 28px !important;
    padding: 0 !important;
    color: #444444 !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-thumb_"] button *,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-thumb_"] button p,
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-thumb_"] button div {
    color: #444444 !important;
    font-size: 1rem !important;
}
/* FAB */
div[class*="st-key-ask_aruvi_fab"] button {
    position: fixed !important;
    bottom: 28px !important;
    right: 28px !important;
    width: 48px !important;
    height: 48px !important;
    border-radius: 50% !important;
    background: #1B2A3B !important;
    color: #ffffff !important;
    font-size: 1.2rem !important;
    border: none !important;
    z-index: 99999 !important;
    box-shadow: 0 4px 16px rgba(0,0,0,0.20) !important;
    min-height: unset !important;
    padding: 0 !important;
}
div[class*="st-key-ask_aruvi_fab"] button:hover {
    background: #2C7A7B !important;
}
html { overflow-y: scroll !important; }
/* ── Detail view panel ── */
.aa-detail-panel {
    padding: 0;
}
.aa-detail-back {
    display: flex;
    align-items: center;
    gap: 6px;
    padding: 12px 16px 10px 16px;
    border-bottom: 1px solid #F0EDE9;
    cursor: pointer;
}
.aa-detail-back-arrow {
    font-size: 0.85rem;
    color: #2C7A7B;
}
.aa-detail-back-label {
    font-size: 0.62rem;
    font-weight: 700;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: #2C7A7B;
}
.aa-detail-cat-title {
    font-size: 0.82rem;
    font-weight: 600;
    color: #1B2A3B;
    padding: 14px 16px 4px 16px;
}
.aa-detail-cat-desc {
    font-size: 0.68rem;
    color: #9C9693;
    padding: 0 16px 12px 16px;
    border-bottom: 1px solid #F0EDE9;
}
.aa-qa-pair {
    padding: 12px 16px;
    border-bottom: 1px solid #F5F3EF;
}
.aa-qa-q {
    font-size: 0.75rem;
    font-weight: 600;
    color: #1B2A3B;
    margin-bottom: 5px;
    line-height: 1.4;
}
.aa-qa-a {
    font-size: 0.72rem;
    color: #5A5754;
    line-height: 1.6;
}
/* Back button inside popup */
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-aa_back_btn"] button {
    background: transparent !important;
    border: none !important;
    color: #2C7A7B !important;
    font-size: 0.7rem !important;
    font-weight: 600 !important;
    padding: 12px 16px 10px 16px !important;
    min-height: unset !important;
    width: 100% !important;
    text-align: left !important;
    justify-content: flex-start !important;
    border-bottom: 1px solid #F0EDE9 !important;
    border-radius: 0 !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
}
div[class*="st-key-ask_aruvi_popup"] div[class*="st-key-aa_back_btn"] button:hover {
    background: #F5F9F9 !important;
    color: #1B2A3B !important;
}
/* Hide CMD+Enter hint — keep 0/140 counter (it's the last child) */
div[class*="st-key-ask_aruvi_popup"] [data-testid="InputInstructions"] > *:first-child {
    display: none !important;
}
/* Follow-up textarea font size */
div[class*="st-key-ask_aruvi_followup"] textarea {
    font-size: 0.65rem !important;
}
/* Follow-up Submit and Skip buttons */
div[class*="st-key-fu_submit"] button,
div[class*="st-key-fu_skip"] button {
    background-color: #4A4A4A !important;
    color: #FFFFFF !important;
    font-size: 0.62rem !important;
    border-radius: 6px !important;
    border: none !important;
}
/* Feedback confirmation div — matches the inline intro paragraph */
.aruvi-fb-confirm {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0 16px;
    margin-top: 22px;
    font-family: inherit;
    font-weight: 400;
    font-size: 0.78rem;
    letter-spacing: normal;
    line-height: 1.45;
    color: #5C5852;
    background: transparent;
}
.aruvi-fb-confirm-text {
    flex: 1;
}
/* Character counter padding — query input */
div[class*="st-key-ask_aruvi_query_input"] small,
div[class*="st-key-ask_aruvi_query_input"] p {
    padding-right: 10px !important;
    margin-right: 10px !important;
    width: calc(100% - 12px) !important;
}
/* Character counter padding — feedback textarea */
div[class*="st-key-ask_aruvi_fb_text"] small,
div[class*="st-key-ask_aruvi_fb_text"] p {
    padding-right: 10px !important;
    margin-right: 10px !important;
    width: calc(100% - 12px) !important;
}
</style>
""", unsafe_allow_html=True)

# FAB button — Streamlit button styled to look like FAB
_fab_label = "✕" if st.session_state.ask_aruvi_open else "💬"
_fab_col = st.container()
with _fab_col:
    if st.button(_fab_label, key="ask_aruvi_fab"):
        st.session_state.ask_aruvi_open = not st.session_state.ask_aruvi_open
        # Always reset to main menu when closing/reopening
        st.session_state.ask_aruvi_detail_cat = None
        st.session_state.ask_aruvi_category = None
        st.session_state.ask_aruvi_response = ""
        st.session_state.ask_aruvi_last_query = ""
        st.session_state.ask_aruvi_show_thumbs = False
        st.session_state.ask_aruvi_thumb_done = False
        st.session_state.ask_aruvi_show_followup = False
        st.rerun()

if st.session_state.ask_aruvi_open:
    with st.container(key="ask_aruvi_popup"):

        # ── Load Q&A knowledge base for detail view ───────────────────────────
        import json as _json
        _qa_kb_path = PROJECT_ROOT / "mirror/ask_aruvi/qa_knowledge_base.json"
        try:
            _qa_kb = _json.loads(_qa_kb_path.read_text(encoding="utf-8"))
        except Exception:
            _qa_kb = {"categories": {}}

        # ── DETAIL VIEW — show Q&A pairs for selected category ────────────────
        if st.session_state.ask_aruvi_detail_cat is not None:
            _dcat_key = st.session_state.ask_aruvi_detail_cat
            _dcat     = _qa_kb.get("categories", {}).get(_dcat_key, {})
            _dcat_label = CATEGORY_LABELS.get(_dcat_key, "")
            _dcat_desc  = _dcat.get("description", "")
            _dpairs     = _dcat.get("pairs", [])

            # Back button
            if st.button("‹  Back to Ask Aruvi", key="aa_back_btn",
                          use_container_width=True):
                st.session_state.ask_aruvi_detail_cat = None
                st.rerun()

            # Category title and description
            st.markdown(
                f'<div class="aa-detail-cat-title">{_dcat_label}</div>'
                f'<div class="aa-detail-cat-desc">{_dcat_desc}</div>',
                unsafe_allow_html=True,
            )

            # Q&A pairs
            for _pair in _dpairs:
                st.markdown(
                    f'<div class="aa-qa-pair">'
                    f'<div class="aa-qa-q">{_pair.get("q", "")}</div>'
                    f'<div class="aa-qa-a">{_pair.get("a", "")}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)

        # ── RESPONSE VIEW — Q&A detail after submitting a question ───────────
        elif st.session_state.ask_aruvi_response:

            # Back button — resets to the state before the question was typed
            if st.button("‹  Back to Ask Aruvi", key="aa_back_btn",
                          use_container_width=True):
                st.session_state.ask_aruvi_response      = ""
                st.session_state.ask_aruvi_last_query    = ""
                st.session_state.ask_aruvi_show_thumbs   = False
                st.session_state.ask_aruvi_thumb_done    = False
                st.session_state.ask_aruvi_show_followup = False
                st.rerun()

            # Question + answer in the same Q&A pair format as category detail
            st.markdown(
                f'<div class="aa-qa-pair">'
                f'<div class="aa-qa-q">{st.session_state.ask_aruvi_last_query}</div>'
                f'<div class="aa-qa-a">{st.session_state.ask_aruvi_response}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

            # Thumbs
            if st.session_state.ask_aruvi_show_thumbs and \
                    not st.session_state.ask_aruvi_thumb_done:
                _t1, _t2, _t3 = st.columns([1, 1, 8])
                with _t1:
                    if st.button("👍︎", key="thumb_up"):
                        write_thumbs_feedback(
                            session_id=st.session_state.ask_aruvi_session_id,
                            rating="up",
                            query=st.session_state.ask_aruvi_last_query,
                            response_excerpt=st.session_state.ask_aruvi_response[:200],
                            category_selected=st.session_state.ask_aruvi_category or "",
                        )
                        st.session_state.ask_aruvi_thumb_done = True
                        st.rerun()
                with _t2:
                    if st.button("👎︎", key="thumb_down"):
                        st.session_state.ask_aruvi_show_followup = True
                        st.rerun()

            if st.session_state.ask_aruvi_show_followup and \
                    not st.session_state.ask_aruvi_thumb_done:
                _fu_text = st.text_area(
                    "followup",
                    placeholder="Please provide feedback on what is missing?",
                    label_visibility="collapsed",
                    key="ask_aruvi_followup",
                    max_chars=140,
                    height=90,
                )
                _fu1, _fu2 = st.columns([1, 1])
                with _fu1:
                    if st.button("Submit", key="fu_submit"):
                        write_thumbs_feedback(
                            session_id=st.session_state.ask_aruvi_session_id,
                            rating="down",
                            query=st.session_state.ask_aruvi_last_query,
                            response_excerpt=st.session_state.ask_aruvi_response[:200],
                            category_selected=st.session_state.ask_aruvi_category or "",
                            follow_up_text=_fu_text or None,
                        )
                        st.session_state.ask_aruvi_thumb_done = True
                        st.rerun()
                with _fu2:
                    if st.button("Skip", key="fu_skip"):
                        write_thumbs_feedback(
                            session_id=st.session_state.ask_aruvi_session_id,
                            rating="down",
                            query=st.session_state.ask_aruvi_last_query,
                            response_excerpt=st.session_state.ask_aruvi_response[:200],
                            category_selected=st.session_state.ask_aruvi_category or "",
                        )
                        st.session_state.ask_aruvi_thumb_done = True
                        st.rerun()

            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)

        # ── MAIN VIEW ─────────────────────────────────────────────────────────
        else:

            # Header
            st.markdown(
                '<div class="aa-header">'
                '<div class="aa-header-title">Ask Aruvi</div>'
                '</div>',
                unsafe_allow_html=True,
            )

            # Category cards — click opens detail view
            for _i, (_key, _label) in enumerate(CATEGORY_LABELS.items(), start=1):
                _active = st.session_state.ask_aruvi_category == _key
                _chip_text = f"{_i}. {_label}  ›"
                if st.button(
                    ("✓  " if _active else "") + _chip_text,
                    key=f"chip_{_key}",
                    type="primary" if _active else "secondary",
                    use_container_width=True,
                ):
                    st.session_state.ask_aruvi_detail_cat = _key
                    st.rerun()

            # ── Entry-point buttons — always visible, open secondary panel ────
            # Two stacked rows: question vs. feedback. The mode flag controls
            # which section the secondary panel renders.
            if st.button(
                "💬  Ask a specific question  ›",
                key="ask_aruvi_open_question",
                use_container_width=True,
            ):
                st.session_state.ask_aruvi_agent_open = True
                st.session_state.ask_aruvi_agent_mode = "question"
                st.session_state.ask_aruvi_open = False   # hide category popup
                st.rerun()
            if st.button(
                "📝  Share feedback  ›",
                key="ask_aruvi_open_feedback",
                use_container_width=True,
            ):
                st.session_state.ask_aruvi_agent_open = True
                st.session_state.ask_aruvi_agent_mode = "feedback"
                st.session_state.ask_aruvi_open = False   # hide category popup
                st.rerun()

# ── Ask Aruvi — secondary panel (Q&A + Feedback) ────────────────────────────
# Opens when the teacher clicks "Ask a specific question or share feedback".
# Uses aruvi_ask (Haiku when USE_MANAGED_AGENT=False, managed agent when True).
# Sits pixel-perfect on top of the category popup via matching CSS geometry.
if st.session_state.ask_aruvi_agent_open:
    with st.container(key="ask_aruvi_agent_popup"):

        # ── RESPONSE VIEW ─────────────────────────────────────────────────────
        if st.session_state.ask_aruvi_agent_response:

            # ── Thanks view (after thumbs-down + Submit/Skip) ─────────────────
            if st.session_state.ask_aruvi_agent_fu_done:
                if st.button("‹  Back", key="aa_agent_back_btn",
                              use_container_width=True):
                    # Reset all agent state and return to the main Ask Aruvi menu
                    st.session_state.ask_aruvi_agent_open           = False
                    st.session_state.ask_aruvi_open                 = True
                    st.session_state.ask_aruvi_agent_response       = ""
                    st.session_state.ask_aruvi_agent_last_query     = ""
                    st.session_state.ask_aruvi_agent_show_thumbs    = False
                    st.session_state.ask_aruvi_agent_thumb_done     = False
                    st.session_state.ask_aruvi_agent_show_followup  = False
                    st.session_state.ask_aruvi_agent_fu_done        = False
                    st.session_state.ask_aruvi_agent_fu_skipped     = False
                    st.rerun()
                if st.session_state.ask_aruvi_agent_fu_skipped:
                    _thanks_body = '<p>Thank you.</p>'
                else:
                    _thanks_body = (
                        '<p>Thanks for letting us know. Your feedback will help us '
                        'improve the service.</p>'
                    )
                st.markdown(
                    '<div class="aa-scope-note">'
                    '<div class="aa-scope-block">'
                    f'{_thanks_body}'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True,
                )

            # ── Standard answer + thumbs + follow-up view ─────────────────────
            else:
                if st.button("‹  Back", key="aa_agent_back_btn",
                              use_container_width=True):
                    st.session_state.ask_aruvi_agent_response      = ""
                    st.session_state.ask_aruvi_agent_last_query    = ""
                    st.session_state.ask_aruvi_agent_show_thumbs   = False
                    st.session_state.ask_aruvi_agent_thumb_done    = False
                    st.session_state.ask_aruvi_agent_show_followup = False
                    st.rerun()

                st.markdown(
                    f'<div class="aa-qa-pair">'
                    f'<div class="aa-qa-q">{st.session_state.ask_aruvi_agent_last_query}</div>'
                    f'<div class="aa-qa-a">{st.session_state.ask_aruvi_agent_response}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                # Thumbs feedback
                if st.session_state.ask_aruvi_agent_show_thumbs and \
                        not st.session_state.ask_aruvi_agent_thumb_done:
                    _at1, _at2, _at3 = st.columns([1, 1, 8])
                    with _at1:
                        if st.button("👍", key="agent_thumb_up"):
                            write_thumbs_feedback(
                                session_id        = st.session_state.ask_aruvi_session_id,
                                rating            = "up",
                                query             = st.session_state.ask_aruvi_agent_last_query,
                                response_excerpt  = st.session_state.ask_aruvi_agent_response[:200],
                                category_selected = "managed_agent",
                            )
                            # Reset all agent state and return to the main Ask Aruvi menu
                            st.session_state.ask_aruvi_agent_open           = False
                            st.session_state.ask_aruvi_open                 = True
                            st.session_state.ask_aruvi_agent_response       = ""
                            st.session_state.ask_aruvi_agent_last_query     = ""
                            st.session_state.ask_aruvi_agent_show_thumbs    = False
                            st.session_state.ask_aruvi_agent_thumb_done     = False
                            st.session_state.ask_aruvi_agent_show_followup  = False
                            st.session_state.ask_aruvi_agent_fu_done        = False
                            st.session_state.ask_aruvi_agent_fu_skipped     = False
                            st.rerun()
                    with _at2:
                        if st.button("👎", key="agent_thumb_down"):
                            st.session_state.ask_aruvi_agent_show_followup = True
                            st.rerun()

                if st.session_state.ask_aruvi_agent_show_followup and \
                        not st.session_state.ask_aruvi_agent_thumb_done:
                    _afu_text = st.text_area(
                        "agent_followup",
                        placeholder="Please provide feedback on what is missing?",
                        label_visibility="collapsed",
                        key="ask_aruvi_agent_followup",
                        max_chars=140,
                        height=90,
                    )
                    with st.container(key="agent_fu_actions"):
                        _afu1, _afu2, _afu3 = st.columns([1, 1, 6])
                        with _afu1:
                            if st.button("Submit", key="agent_fu_submit"):
                                write_thumbs_feedback(
                                    session_id        = st.session_state.ask_aruvi_session_id,
                                    rating            = "down",
                                    query             = st.session_state.ask_aruvi_agent_last_query,
                                    response_excerpt  = st.session_state.ask_aruvi_agent_response[:200],
                                    category_selected = "managed_agent",
                                    follow_up_text    = _afu_text or None,
                                )
                                st.session_state.ask_aruvi_agent_thumb_done = True
                                st.session_state.ask_aruvi_agent_fu_done    = True
                                st.rerun()
                        with _afu2:
                            if st.button("Skip", key="agent_fu_skip"):
                                write_thumbs_feedback(
                                    session_id        = st.session_state.ask_aruvi_session_id,
                                    rating            = "down",
                                    query             = st.session_state.ask_aruvi_agent_last_query,
                                    response_excerpt  = st.session_state.ask_aruvi_agent_response[:200],
                                    category_selected = "managed_agent",
                                )
                                st.session_state.ask_aruvi_agent_thumb_done = True
                                st.session_state.ask_aruvi_agent_fu_done    = True
                                st.session_state.ask_aruvi_agent_fu_skipped = True
                                st.rerun()

            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)

        # ── MAIN VIEW — ask + feedback ────────────────────────────────────────
        else:
            # Back button — mirrors the category detail back button in style
            if st.button("‹  Back to Ask Aruvi", key="ask_aruvi_agent_close",
                          use_container_width=True):
                st.session_state.ask_aruvi_agent_open = False
                st.session_state.ask_aruvi_open = True    # restore category popup
                st.session_state.ask_aruvi_agent_fb_sent = False  # clear confirmation
                st.rerun()

            # Q&A pill — shown only in question mode
            if st.session_state.ask_aruvi_agent_mode == "question":
                # Inline scope disclaimer — sets expectations before tokens are spent
                st.markdown(
                    '<div class="aa-scope-note">'
                    '<div class="aa-scope-block">'
                    '<span class="aa-scope-head">What I can help with</span>'
                    '<p>Ask me anything about how Aruvi works — how chapters are '
                    'allocated time, how your lesson plans and assessments are built, '
                    'how the NCF shapes what\'s taught, or how to find your way around '
                    'the app.</p>'
                    '</div>'
                    '<div class="aa-scope-block">'
                    '<span class="aa-scope-head">What I can\'t help with</span>'
                    '<p>I\'m not built to teach the lesson itself — explaining '
                    'chapter concepts, solving textbook questions, or summarising a '
                    'lesson isn\'t something I can do. Your textbook is the better '
                    'place for those.</p>'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True,
                )
                _agent_query = st.text_area(
                    "agent_query",
                    placeholder="Ask a specific question…",
                    label_visibility="collapsed",
                    key="ask_aruvi_agent_query_input",
                    height=96,
                    max_chars=140,
                )
                _agent_ask_clicked = st.button("➤", key="ask_aruvi_agent_submit",
                                                use_container_width=False)
                st.markdown(
                    f'<div class="aa-char-count">{len(_agent_query or "")}/140</div>',
                    unsafe_allow_html=True,
                )

                if _agent_ask_clicked and _agent_query.strip():
                    with st.spinner(""):
                        _agent_result = aruvi_ask(
                            query      = _agent_query.strip(),
                            session_id = st.session_state.ask_aruvi_session_id,
                            tab        = st.session_state.role,
                            subject    = st.session_state.get("subject", ""),
                            grade      = st.session_state.get("grade", ""),
                        )
                    st.session_state.ask_aruvi_agent_response      = _agent_result["response"]
                    st.session_state.ask_aruvi_agent_last_query    = _agent_query.strip()
                    st.session_state.ask_aruvi_agent_show_thumbs   = True
                    st.session_state.ask_aruvi_agent_thumb_done    = False
                    st.session_state.ask_aruvi_agent_show_followup = False
                    log_ask_aruvi_tokens(
                        session_id    = st.session_state.ask_aruvi_session_id,
                        query         = _agent_query.strip(),
                        category      = "managed_agent",
                        tab           = st.session_state.role,
                        subject       = st.session_state.get("subject", ""),
                        grade         = st.session_state.get("grade", ""),
                        input_tokens  = _agent_result.get("input_tokens", 0),
                        output_tokens = _agent_result.get("output_tokens", 0),
                    )
                    st.rerun()

            # Feedback pill — shown only in feedback mode
            elif st.session_state.ask_aruvi_agent_mode == "feedback":
                if not st.session_state.ask_aruvi_agent_fb_sent:
                    # Inline intro — warms up the ask and sets the follow-up expectation
                    st.markdown(
                        '<div class="aa-scope-note">'
                        '<div class="aa-scope-block">'
                        '<p>Tell me what\'s working, what\'s not, or anything you\'d '
                        'like Aruvi to do better. We read every note and will get back '
                        'to you soon.</p>'
                        '</div>'
                        '</div>',
                        unsafe_allow_html=True,
                    )
                    _agent_fb_text = st.text_area(
                        "agent_feedback",
                        placeholder="Share feedback on Aruvi…",
                        label_visibility="collapsed",
                        key=f"ask_aruvi_agent_fb_text_{st.session_state.ask_aruvi_agent_fb_reset}",
                        height=96,
                        max_chars=140,
                    )
                    _agent_fb_clicked = st.button("➤", key="ask_aruvi_agent_fb_submit")
                    if _agent_fb_text:
                        st.markdown(
                            f'<div class="aa-char-count">{len(_agent_fb_text)}/140</div>',
                            unsafe_allow_html=True,
                        )
                    if _agent_fb_clicked and _agent_fb_text.strip():
                        write_general_feedback(
                            session_id    = st.session_state.ask_aruvi_session_id,
                            feedback_text = _agent_fb_text.strip(),
                            tab           = st.session_state.role,
                            subject       = st.session_state.get("subject", ""),
                            grade         = st.session_state.get("grade", ""),
                        )
                        st.session_state.ask_aruvi_agent_fb_sent  = True
                        st.session_state.ask_aruvi_agent_fb_reset += 1
                        st.rerun()
                else:
                    # Box is hidden after submit — user must close & reopen feedback
                    # to leave another note (avoids confusion about double-submit).
                    st.markdown(
                        '<div class="aruvi-fb-confirm">'
                        '<span class="aruvi-fb-confirm-text">Thank you. We\'ve received your note and will get back to you soon.</span>'
                        '</div>',
                        unsafe_allow_html=True,
                    )

            st.markdown('<div style="height:12px"></div>', unsafe_allow_html=True)
