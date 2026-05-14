"""Build Aruvi_Project_Report_V7.docx from a structured outline.

One-off generator. Mirrors V6's title-page style (centered title, subtitle,
version line) and uses Heading 1 / Heading 2 / Normal styles for the body.
No tables, no images.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT = (
    Path(__file__).resolve().parents[1]
    / "knowledge_commons"
    / "other commons"
    / "Aruvi_Project_Report_V7.docx"
)


def add_centered(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_para(doc, text):
    return doc.add_paragraph(text)


def add_h1(doc, text):
    return doc.add_heading(text, level=1)


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def build():
    doc = Document()

    # Title page
    add_centered(doc, "ARUVI · அருவி", bold=True, size=24)
    add_centered(doc, "An NCF 2023-Aligned Pedagogical Platform for Indian Schools", size=12)
    add_centered(doc, "Project Report · Version 7 · Internal Document", size=11)
    doc.add_paragraph()

    # ---------------------------------------------------------------- Section 1
    add_h1(doc, "1 · What Aruvi Is")

    add_para(
        doc,
        "NCF 2023 marks a real shift in Indian schooling — from finishing chapters "
        "to building competencies. Schools accept this shift in principle. They "
        "struggle with it in practice, because there is no translation layer between "
        "a policy document and a teacher’s next class. The gap is not intent. It is "
        "the absence of working tools.",
    )

    add_para(
        doc,
        "Aruvi is that translation layer. It maps NCERT chapters to NCF competencies, "
        "generates lesson plans that fit the teacher’s available periods and are "
        "rooted in actual textbook content, builds chapter assessments aligned to the "
        "same plan, and produces a whole-year allocation of teaching time across "
        "chapters. Every output is traceable to a versioned rule document — the "
        "subject’s “constitution” — so a reader can see why the AI made each "
        "decision.",
    )

    add_para(
        doc,
        "Aruvi serves four kinds of users. Teachers get a complete lesson plan and "
        "assessment for any chapter, scaled to the periods they actually have. "
        "Administrators and principals get a whole-school annual plan that splits "
        "teaching time proportionately across chapters and surfaces curricular gaps "
        "before they show up in results. Home schoolers get a structured path through "
        "a grade’s textbooks with ready activities and check-points. Self-taught "
        "students get assessments and a worked progression to test their own "
        "understanding, chapter by chapter.",
    )

    # ---------------------------------------------------------------- Section 2
    add_h1(doc, "2 · The Approach")

    add_para(
        doc,
        "Aruvi’s central insight is that the four subjects — Social Sciences, "
        "Science, English, and Mathematics — need genuinely different organising "
        "logics, because their NCERT textbooks are built differently. Forcing all "
        "four through one universal recipe would flatten real differences and weaken "
        "the result. The lesson plan and the assessment for each subject flow from "
        "the same organising idea, so each strand below covers them together.",
    )

    # 2.1 Social Sciences
    add_h2(doc, "2.1 Social Sciences — Competency-based, with an implied-LO handoff")
    add_para(
        doc,
        "Organising idea: each chapter is rich enough to develop several NCF "
        "competencies across its sections, and the section is the natural unit of "
        "work.",
    )
    add_para(
        doc,
        "The lesson plan walks the chapter section by section. For every activity "
        "the plan emits one observable, implied learning outcome — a single sentence "
        "describing what a student who completed the activity can now do. The plan "
        "is built within the teacher’s actual period budget, with each activity "
        "anchored to a named section of the textbook so nothing is invented.",
    )
    add_para(
        doc,
        "Those implied learning outcomes are the only bridge between the lesson "
        "plan and the assessment. The assessment generator reads them, ignores "
        "everything else, and writes one set of questions per outcome. The depth "
        "and type of questions are set by how heavily the chapter develops each "
        "competency — Central competencies get extended responses and an open task, "
        "Substantive competencies get short constructed responses, and Present "
        "competencies get diagnostic multiple-choice questions.",
    )
    add_para(
        doc,
        "This fits Social Sciences because the chapters genuinely sprawl. A single "
        "chapter on the rise of empires can develop chronological reasoning, "
        "source analysis, and geographical interpretation in parallel. The section-"
        "by-section walk respects that sprawl instead of forcing one competency to "
        "stand in for all.",
    )

    # 2.2 Science
    add_h2(doc, "2.2 Science — Progression stage based")
    add_para(
        doc,
        "Organising idea: each Science chapter builds one central cognitive "
        "operation through clear stages — observe, infer, generalise, apply. The "
        "stage, not the section, is the natural unit of work.",
    )
    add_para(
        doc,
        "The lesson plan reads the chapter as a whole, identifies the progression "
        "the chapter itself sets up, and lays activities along those stages. Each "
        "stage ends with one implied learning outcome that names the cognitive step "
        "the student has just taken. Activities are drawn from the chapter’s own "
        "embedded experiments, exercise questions, and exploratory prompts — Aruvi "
        "does not invent activities the textbook does not support.",
    )
    add_para(
        doc,
        "The assessment receives the implied learning outcomes in order. Question "
        "type and depth scale with the stage’s position in the progression: "
        "foundational outcomes get diagnostic multiple-choice questions, middle "
        "outcomes add short constructed responses, and the final outcome gets an "
        "open task or extended investigation that exercises the chapter’s full "
        "central operation.",
    )
    add_para(
        doc,
        "This fits Science because Science chapters are self-contained competency "
        "units. The chapter is built to produce one cognitive ability; the right "
        "scaffolding is the staircase that gets the student there, not a set of "
        "parallel competencies.",
    )

    # 2.3 English
    add_h2(doc, "2.3 English — Spine-oriented")
    add_para(
        doc,
        "Organising idea: language teaching is built on six skill spines — Reading "
        "for Comprehension, Listening, Speaking, Writing, Vocabulary and Grammar, "
        "and Beyond-the-Text. Within a chapter, every main section (poem, prose, "
        "dialogue) is walked across all six spines.",
    )
    add_para(
        doc,
        "The lesson plan treats each (section × spine) pair as a cell and emits one "
        "implied learning outcome per cell — for example, an outcome for the poem’s "
        "Reading work, another for its Speaking work, another for its Writing work. "
        "NCF competency codes are present underneath, but the teacher never has to "
        "engage with them: the spine structure carries the alignment.",
    )
    add_para(
        doc,
        "The assessment generates one originally-written question per implied "
        "outcome, grounded in the section’s actual text. A Reading outcome gets a "
        "comprehension question on the actual passage. A Writing outcome gets a "
        "writing task tied to the same passage. The teacher receives a guide for "
        "every question explaining what to look for.",
    )
    add_para(
        doc,
        "This fits English because language pedagogy is the spines. Splitting them "
        "out and walking through them is more faithful than forcing English into a "
        "competency-mapping mould borrowed from another subject.",
    )

    # 2.4 Mathematics
    add_h2(
        doc,
        "2.4 Mathematics — Textbook-based worked examples, activities, and exercises",
    )
    add_para(
        doc,
        "Organising idea: every Mathematics chapter has its own internal progression "
        "— worked examples introduce a method, activities make the method tangible, "
        "and exercises drill it. Each chapter also builds on the chapter before it.",
    )
    add_para(
        doc,
        "The lesson plan respects this innate progression. Each period is anchored "
        "to specific worked examples, activities, and exercise questions from the "
        "textbook in order. Prior chapter knowledge is treated as the entry "
        "condition — the plan begins where the previous chapter left off, not at "
        "an abstract competency.",
    )
    add_para(
        doc,
        "The assessment pairs every implied learning outcome with two artefacts: a "
        "freshly written question that tests transferable understanding, and a "
        "companion exercise drawn from the textbook’s own set. The student is "
        "tested both on whether the method generalises and on whether the textbook’s "
        "own progression has been internalised.",
    )
    add_para(
        doc,
        "This fits Mathematics because Mathematics is the only subject where the "
        "textbook itself already encodes a rigorous progression. Aruvi co-opts the "
        "textbook rather than competing with it.",
    )

    # Section 2 close
    add_para(
        doc,
        "Across all four strands the mechanism is the same: chapter content drives "
        "a lesson plan, the lesson plan emits implied learning outcomes, and only "
        "those outcomes pass through to the assessment. The organising principle "
        "differs by subject because the textbook structure differs by subject.",
    )

    # ---------------------------------------------------------------- Section 3
    add_h1(doc, "3 · The Planning Tool")

    add_para(
        doc,
        "The Planning Tool answers a single question: given a fixed total of "
        "available periods in a term or a year, how should they be split across "
        "chapters so that chapters needing more classroom effort get more time? "
        "The answer depends on the subject, because what makes a chapter heavy is "
        "different in each subject.",
    )

    add_h2(doc, "3.1 Social Sciences — Competency weight")
    add_para(
        doc,
        "Each chapter’s competencies are weighted Central (3), Substantive (2), or "
        "Present (1). The chapter’s weight is the sum of these primary weights. "
        "The Planning Tool allocates periods in proportion to chapter weight — "
        "chapters that develop more competencies, or develop them more deeply, get "
        "more teaching time.",
    )

    add_h2(doc, "3.2 Science — Effort index")
    add_para(
        doc,
        "Periods are allocated on a composite effort index. Four signals are read "
        "directly from the chapter: how conceptually demanding the exercises are, "
        "how many student-run activities the chapter includes, how many teacher "
        "demonstrations it requires, and how heavy the exercise execution load is. "
        "Each signal has a multiplier that reflects its true classroom time cost — "
        "teacher demonstrations carry a higher multiplier than student activities, "
        "for example, because they include preparation overhead that the period "
        "count does not see. Click any chapter and the breakdown is visible. The "
        "number is never opaque.",
    )

    add_h2(doc, "3.3 English — Effort index, spine-based signals")
    add_para(
        doc,
        "The same idea as Science, with signals chosen for language work: spine "
        "load (how many of the six spines the chapter actively develops), task "
        "density (how many graded tasks the chapter sets), writing demand (the "
        "weight of long-form writing), and project load (the size of any "
        "Beyond-the-Text projects). These combine into an effort index that drives "
        "proportional allocation. A principal can see at a glance which chapters "
        "are heavy on writing or speaking, and plan time around them.",
    )

    add_h2(doc, "3.4 Mathematics — Effort index, textbook-architecture signals")
    add_para(
        doc,
        "Again the same idea, with signals drawn from the textbook’s own "
        "architecture: the size of the worked-example set, the count of in-chapter "
        "activities, the exercise load, and the conceptual demand of the methods "
        "introduced. Chapters with denser worked-example progressions naturally "
        "get more periods.",
    )

    add_para(
        doc,
        "In every subject the Planning Tool offers a defensible default — not a "
        "final answer. The teacher or principal owns the final number and can "
        "override any chapter’s allocation before generating a plan. The tool’s "
        "job is to make the trade-off visible, not to remove it.",
    )

    # Footer line
    doc.add_paragraph()
    add_centered(
        doc,
        "Aruvi · NCF 2023-Aligned Pedagogical Platform · "
        "Project Report Version 7 · Internal Document",
        size=9,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
