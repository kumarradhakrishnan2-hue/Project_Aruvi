"""
refresh_kb.py — Ask Aruvi Q&A Knowledge Base Refresh Script
Two modes: --propose (Stage 1) and --write --proposals N,M,P (Stage 2)
"""

import argparse
import json
import os
import sys
from datetime import date
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).parent.parent
KB_JSON      = PROJECT_ROOT / "mirror" / "ask_aruvi" / "qa_knowledge_base.json"
PENDING_JSON = PROJECT_ROOT / "mirror" / "ask_aruvi" / "kb_proposals_pending.json"
OUTPUT_DIR   = PROJECT_ROOT / "knowledge_commons" / "other commons"
OUTPUT_DOCX  = OUTPUT_DIR / "qa_knowledge_base.docx"
FWD_QUERIES  = PROJECT_ROOT / "mirror" / "ask_aruvi" / "forwarded_queries"

KNOWLEDGE_SOURCES = {
    "cat_a": [
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Optimizer_PeriodAllocation_V1.docx",
    ],
    "cat_b": [
        PROJECT_ROOT / "knowledge_commons" / "constitutions" / "assessment" / "science" / "Assessment_Constitution_Science.docx",
        PROJECT_ROOT / "knowledge_commons" / "constitutions" / "assessment" / "social_sciences" / "Assessment_Constitution_Social_Sciences.docx",
    ],
    "cat_c": [
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Project_Report_V7.docx",
        PROJECT_ROOT / "knowledge_commons" / "constitutions" / "competency_mapping" / "sciences" / "Aruvi_Competency_Mapping_Constitution_science.docx",
    ],
    "cat_d": [
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Storage_Protocol_V3.docx",
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Ask_Aruvi_Design_V1_1.docx",
    ],
    "cat_e": [
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Ask_Aruvi_Design_V1_1.docx",
        PROJECT_ROOT / "knowledge_commons" / "other commons" / "Aruvi_Project_Report_V7.docx",
    ],
}

CATEGORY_ORDER = ["cat_a", "cat_b", "cat_c", "cat_d", "cat_e"]

# ── Helpers ────────────────────────────────────────────────────────────────

def load_kb():
    with open(KB_JSON, "r", encoding="utf-8") as f:
        return json.load(f)


def extract_docx_text(path: Path, max_chars: int = 6000) -> str:
    """Extract plain text from a .docx file, up to max_chars."""
    if not path.exists():
        return f"[Source not found: {path.name}]"
    try:
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text[:max_chars]
    except Exception as e:
        return f"[Error reading {path.name}: {e}]"


def count_words(text: str) -> int:
    return len(text.split())


def truncate_to_words(text: str, limit: int = 80) -> str:
    words = text.split()
    if len(words) <= limit:
        return text
    return " ".join(words[:limit]) + "…"


def call_haiku(client, source_excerpt: str, q: str, a: str) -> dict:
    """Call Haiku to judge a single Q&A pair. Returns verdict dict."""
    prompt = (
        f"KNOWLEDGE SOURCE EXCERPT:\n{source_excerpt}\n\n"
        f"Q&A PAIR:\nQ: {q}\nA: {a}\n\n"
        'Return JSON: {"verdict": "keep"|"amend"|"delete", '
        '"reason": "one sentence", '
        '"revised_answer": "amended answer text if verdict is amend, else null"}'
    )
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=400,
        system=(
            "You are an editorial assistant for the Ask Aruvi platform helpline. "
            "Judge whether the given Q&A pair is accurate and current against "
            "the provided knowledge source excerpt. "
            "Respond ONLY in valid JSON. No preamble, no markdown fences."
        ),
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.content[0].text.strip()
    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"verdict": "keep", "reason": "Could not parse Haiku response.", "revised_answer": None}


def load_forwarded_queries() -> list:
    """Return list of unmatched query strings from forwarded_queries/."""
    queries = []
    if not FWD_QUERIES.exists():
        return queries
    for f in sorted(FWD_QUERIES.glob("*.json")):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            if isinstance(data, list):
                queries.extend([str(q) for q in data])
            elif isinstance(data, dict) and "queries" in data:
                queries.extend([str(q) for q in data["queries"]])
        except Exception:
            pass
    return queries


# ── Stage 1: Propose ───────────────────────────────────────────────────────

def run_propose():
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: ANTHROPIC_API_KEY not set.", file=sys.stderr)
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)
    kb = load_kb()

    proposals = []
    proposal_num = 0
    report_lines = []

    report_lines.append("=" * 64)
    report_lines.append("ASK ARUVI — Q&A KNOWLEDGE BASE: PROPOSALS REPORT")
    report_lines.append(f"Generated: {date.today().isoformat()}")
    report_lines.append("=" * 64)

    for cat_key in CATEGORY_ORDER:
        cat = kb["categories"][cat_key]
        label = cat["label"].upper()
        report_lines.append(f"\nCATEGORY {cat_key[-1].upper()} — {label}")

        # Build source excerpt for this category
        excerpts = []
        for src_path in KNOWLEDGE_SOURCES.get(cat_key, []):
            text = extract_docx_text(src_path, max_chars=4000)
            excerpts.append(f"[{src_path.name}]\n{text}")
        source_excerpt = "\n\n".join(excerpts)[:7000]

        unchanged_count = 0
        cat_proposals = []

        for pair in cat["pairs"]:
            q = pair["q"]
            a = pair["a"]
            pair_id = pair.get("id", "?")

            result = call_haiku(client, source_excerpt, q, a)
            verdict = result.get("verdict", "keep")
            reason = result.get("reason", "")
            revised = result.get("revised_answer")

            if verdict == "keep":
                unchanged_count += 1
                continue

            proposal_num += 1
            proposal = {
                "number": proposal_num,
                "type": verdict.upper(),
                "cat_key": cat_key,
                "pair_id": pair_id,
                "q": q,
                "current_a": a,
                "proposed_a": revised if verdict == "amend" else None,
                "reason": reason,
            }
            proposals.append(proposal)
            cat_proposals.append(proposal)

        report_lines.append(f"  Unchanged: {unchanged_count} pairs")

        for p in cat_proposals:
            report_lines.append(f"\n  Proposal {p['number']} — {p['type']}")
            report_lines.append(f"    Q: {p['q']}")
            if p["type"] == "AMEND":
                report_lines.append(f"    Current A: {p['current_a']}")
                report_lines.append(f"    Proposed A: {p['proposed_a']}")
            elif p["type"] == "DELETE":
                report_lines.append(f"    Current A: {p['current_a']}")
            report_lines.append(f"    Reason: {p['reason']}")

        if not cat_proposals:
            report_lines.append("  No proposals.")

        report_lines.append("-" * 64)

    # Check forwarded queries for ADD proposals
    fwd_queries = load_forwarded_queries()
    if fwd_queries:
        report_lines.append("\nFORWARDED QUERIES (candidate ADD pairs)")
        for i, q_text in enumerate(fwd_queries[:5], 1):
            report_lines.append(f"  [{i}] {q_text}")

    # Summary
    amend_count = sum(1 for p in proposals if p["type"] == "AMEND")
    delete_count = sum(1 for p in proposals if p["type"] == "DELETE")
    add_count    = sum(1 for p in proposals if p["type"] == "ADD")
    unchanged_total = kb["metadata"]["total_pairs"] - len(proposals)

    report_lines.append("\n" + "=" * 64)
    report_lines.append(f"SUMMARY: {len(proposals)} proposals across {len(set(p['cat_key'] for p in proposals))} categories")
    report_lines.append(f"  AMEND: {amend_count}   DELETE: {delete_count}   ADD: {add_count}   Unchanged pairs: {unchanged_total}")
    report_lines.append("")
    report_lines.append('Awaiting clearance. Reply with: "Clear all", "Clear N,M,P",')
    report_lines.append('"Clear all except N,M", or inline amendments per proposal.')
    report_lines.append("=" * 64)

    report = "\n".join(report_lines)
    print(report)

    # Save pending proposals
    pending = {
        "generated": date.today().isoformat(),
        "proposals": proposals,
    }
    PENDING_JSON.write_text(json.dumps(pending, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\n[Saved proposals to {PENDING_JSON}]")

    if not proposals:
        print("\nNo proposals generated — KB is up to date. No action needed.")


# ── Stage 2: Write ─────────────────────────────────────────────────────────

def run_write(proposal_ids: str):
    if not PENDING_JSON.exists():
        print("ERROR: No pending proposals file found. Run --propose first.", file=sys.stderr)
        sys.exit(1)

    pending = json.loads(PENDING_JSON.read_text(encoding="utf-8"))
    all_proposals = pending["proposals"]

    # Parse which proposals to apply
    if proposal_ids.strip().lower() == "all":
        cleared = all_proposals
    else:
        nums = {int(x.strip()) for x in proposal_ids.split(",") if x.strip().isdigit()}
        cleared = [p for p in all_proposals if p["number"] in nums]

    if not cleared:
        print("No matching proposals to apply. Nothing written.")
        return

    kb = load_kb()
    warnings = []

    for proposal in cleared:
        cat_key   = proposal["cat_key"]
        pair_id   = proposal["pair_id"]
        ptype     = proposal["type"]
        pairs     = kb["categories"][cat_key]["pairs"]

        if ptype == "AMEND":
            new_a = proposal.get("override_a") or proposal["proposed_a"] or proposal["current_a"]
            if count_words(new_a) > 80:
                new_a = truncate_to_words(new_a, 80)
                warnings.append(f"Proposal {proposal['number']}: answer truncated to 80 words.")
            for pair in pairs:
                if pair.get("id") == pair_id:
                    pair["a"] = new_a
                    break

        elif ptype == "DELETE":
            for pair in pairs:
                if pair.get("id") == pair_id:
                    pair["flagged_for_deletion"] = True
                    break

        elif ptype == "ADD":
            new_q = proposal["q"]
            new_a = proposal.get("override_a") or proposal.get("proposed_a", "")
            if count_words(new_a) > 80:
                new_a = truncate_to_words(new_a, 80)
                warnings.append(f"Proposal {proposal['number']}: answer truncated to 80 words.")
            # Generate a new ID
            existing_ids = {p.get("id", "") for p in pairs}
            letter = cat_key[-1]
            for n in range(1, 200):
                cand = f"{letter}{n:02d}"
                if cand not in existing_ids:
                    new_id = cand
                    break
            pairs.append({"id": new_id, "q": new_q, "a": new_a})

    # Update metadata
    total = sum(len(v["pairs"]) for v in kb["categories"].values())
    kb["metadata"]["total_pairs"] = total
    kb["metadata"]["last_refreshed"] = date.today().isoformat()

    # Write JSON
    KB_JSON.write_text(json.dumps(kb, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"✓ Updated JSON written to {KB_JSON}")

    # Write Word doc
    write_word_doc(kb)
    print(f"✓ Word doc written to {OUTPUT_DOCX}")

    # Clean up pending
    PENDING_JSON.unlink()
    print("✓ Pending proposals file deleted.")

    if warnings:
        print("\nWarnings:")
        for w in warnings:
            print(f"  ⚠ {w}")

    print(f"\nApplied {len(cleared)} proposal(s). KB now has {total} pairs.")


# ── Word document writer ───────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Add a thin horizontal line between Q&A pairs."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def write_word_doc(kb: dict):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_heading("Ask Aruvi — Q&A Knowledge Base", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Version {kb['metadata'].get('version', 'V1.1')}  |  "
        f"Last refreshed: {kb['metadata']['last_refreshed']}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    for cat_key in CATEGORY_ORDER:
        cat = kb["categories"][cat_key]
        # Category heading
        heading = doc.add_heading(cat["label"], level=1)

        # Category description (italic)
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(cat.get("description", ""))
        desc_run.italic = True
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph()

        active_pairs = [p for p in cat["pairs"] if not p.get("flagged_for_deletion")]

        for i, pair in enumerate(active_pairs):
            # Q line
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q: {pair['q']}")
            q_run.bold = True
            q_run.font.size = Pt(10.5)

            # A line
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"A: {pair['a']}")
            a_run.font.size = Pt(10.5)

            if i < len(active_pairs) - 1:
                add_horizontal_rule(doc)

        doc.add_page_break()

    # Footer via core properties isn't supported directly — add footer text as last para
    footer_para = doc.add_paragraph("Aruvi Knowledge Commons | Internal Document")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(str(OUTPUT_DOCX))


# ── CLI ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Ask Aruvi KB Refresh Script")
    parser.add_argument("--propose", action="store_true", help="Stage 1: generate proposals")
    parser.add_argument("--write",   action="store_true", help="Stage 2: write cleared proposals")
    parser.add_argument("--proposals", type=str, default="all",
                        help="Comma-separated proposal numbers to apply, or 'all'")
    args = parser.parse_args()

    if args.propose:
        run_propose()
    elif args.write:
        run_write(args.proposals)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
